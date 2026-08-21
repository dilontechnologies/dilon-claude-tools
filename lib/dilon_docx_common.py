# -*- coding: utf-8 -*-
"""
Shared Pandoc-conversion and Word-styling helpers used by both
dilon-document-compiler and dilon-document-form-compiler.

Moved out of generate_dilon_doc.py as a pure refactor - see
docs/superpowers/specs/2026-08-17-document-extraction-and-form-tooling-design.md
(dilon-claude-tools repo) for why these two skills share this module.
"""

import math
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docxcompose.composer import Composer
from jinja2 import Environment

# Pandoc's own {#id}/{#fig:label} attribute syntax (documented in
# MARKDOWN_STYLING_GUIDE.md for figure labels and heading ids) collides
# with Jinja2's default {# ... #} comment syntax - {#fig:label} alone (no
# closing #}) raises TemplateSyntaxError: Missing end of comment tag.
# Body rendering only ever needs {{field}} substitution and
# {% raw %}...{% endraw %} escaping, never comments, so the comment
# delimiter is disabled outright (set to a string that can't occur in
# authored markdown) rather than worked around per-document.
_JINJA_ENV = Environment(
    comment_start_string='\x00JINJA_COMMENT_DISABLED_START\x00',
    comment_end_string='\x00JINJA_COMMENT_DISABLED_END\x00',
)


def apply_table_style_to_object(table, style_name):
    """
    Apply a custom table style to a single table object.

    Args:
        table: python-docx Table object
        style_name: Name of the table style to apply
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Apply the table style
    table.style = style_name

    # Get or create tblPr element
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)

    # Get or create tblLook element
    tbl_look = tbl_pr.find(qn('w:tblLook'))
    if tbl_look is None:
        tbl_look = OxmlElement('w:tblLook')
        tbl_pr.append(tbl_look)

    # Set table style options based on the style name
    if style_name == 'DilonTable_Chart':
        # Enable both header row and first column
        tbl_look.set(qn('w:val'), '04A0')
        tbl_look.set(qn('w:firstRow'), '1')
        tbl_look.set(qn('w:firstColumn'), '1')
        tbl_look.set(qn('w:lastRow'), '0')
        tbl_look.set(qn('w:lastColumn'), '0')
        tbl_look.set(qn('w:noHBand'), '0')
        tbl_look.set(qn('w:noVBand'), '1')
    elif style_name == 'DilonTable_List':
        # Enable header row only
        tbl_look.set(qn('w:val'), '0400')
        tbl_look.set(qn('w:firstRow'), '1')
        tbl_look.set(qn('w:firstColumn'), '0')
        tbl_look.set(qn('w:lastRow'), '0')
        tbl_look.set(qn('w:lastColumn'), '0')
        tbl_look.set(qn('w:noHBand'), '0')
        tbl_look.set(qn('w:noVBand'), '1')


def apply_paragraph_style_to_objects(paragraphs, style_name):
    """
    Apply a custom paragraph style to a list of paragraph objects.

    Args:
        paragraphs: List of python-docx Paragraph objects
        style_name: Name of the paragraph style to apply
    """
    for para in paragraphs:
        para.style = style_name


def parse_column_widths(spec_text, num_columns):
    """
    Parse a @@@TABLE_COLUMNS:...@@@ marker's contents into a per-column
    width spec.

    Returns a list of length num_columns where each entry is a positive
    float (inches) or the string 'x' (flex column), or None if the spec
    is invalid: wrong entry count, 2+ 'x' entries, or an entry
    that isn't a positive number or 'x'/'X'.
    """
    entries = [entry.strip() for entry in spec_text.split(',')]
    if len(entries) != num_columns:
        return None

    parsed = []
    flex_count = 0
    for entry in entries:
        if entry.lower() == 'x':
            flex_count += 1
            parsed.append('x')
            continue

        try:
            value = float(entry)
        except ValueError:
            return None
        if value <= 0 or not math.isfinite(value):
            return None
        parsed.append(value)

    if flex_count > 1:
        return None

    return parsed


def apply_table_column_widths(table, widths, available_width):
    """
    Set explicit per-column widths (inches) on a table.

    widths: list as returned by parse_column_widths() - one entry per
    column, each a positive float (inches) or 'x' (flex column).
    available_width: total content width in inches (page width minus
    margins) the table's columns must not exceed.

    Raises ValueError if the fixed widths leave no room for a flex
    column, or if there's no flex column and the fixed widths alone
    exceed available_width.
    """
    fixed_total = sum(width for width in widths if width != 'x')
    has_flex = 'x' in widths

    if has_flex:
        flex_width = available_width - fixed_total
        if flex_width <= 0:
            raise ValueError(
                f"fixed column widths ({fixed_total}in) leave no room for the "
                f"flex column within the available width ({available_width}in)"
            )
    elif fixed_total > available_width:
        raise ValueError(
            f"column widths ({fixed_total}in) exceed the available width ({available_width}in)"
        )

    table.autofit = False
    for idx, width in enumerate(widths):
        value = flex_width if width == 'x' else width
        emu_width = Inches(value)
        table.columns[idx].width = emu_width
        for cell in table.columns[idx].cells:
            cell.width = emu_width


def apply_styles(docx_file):
    """
    Apply custom styles to tables and paragraphs based on @@@ markers in the Word document.

    Uses a state machine to scan the document once:
    - NO_MARKER: Default state, looking for @@@ markers
    - PARAGRAPH_MARKER: Found paragraph start marker, searching for END_STYLE

    States:
        NO_MARKER -> NO_MARKER (found @@@TABLE_STYLE:...@@@ - handled immediately)
        NO_MARKER -> PARAGRAPH_MARKER (found @@@STYLE:...@@@ without END on same line)
        NO_MARKER -> NO_MARKER (found @@@STYLE:...@@@ with END on same line)
        PARAGRAPH_MARKER -> NO_MARKER (found @@@END_STYLE@@@)

    Args:
        docx_file: Path to the Word document
    """
    doc = Document(docx_file)

    # State machine states
    NO_MARKER = 0
    PARAGRAPH_MARKER = 1

    state = NO_MARKER

    # State variables
    saved_style_name = None
    paragraph_start = None

    # Operation collections
    table_operations = []  # List of (table_object, style_name)
    styled_table_elements = set()  # Set of XML elements for tables that received explicit styles
    column_width_operations = []  # List of (table_object, columns_spec_text)
    paragraph_operations = []  # List of (start_idx, end_idx, style_name)
    paragraphs_to_trim = []  # List of (idx, cleaned_text)

    print(f"  Scanning document with state machine...")

    # PHASE 1: Single-pass scan with state machine
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if state == NO_MARKER:
            # Early exit: skip paragraphs without markers
            if '@@@' not in text:
                continue

            # Skip code blocks to avoid processing documentation examples
            if para.runs and para.runs[0].style and 'Verbatim' in para.runs[0].style.name:
                continue

            # Handle table style / column-width markers (must start the
            # paragraph; either marker may lead when both are stacked,
            # since Pandoc merges adjacent no-blank-line marker lines into
            # a single paragraph)
            if text.startswith('@@@TABLE_STYLE:') or text.startswith('@@@TABLE_COLUMNS:'):
                style_match = re.search(r'@@@TABLE_STYLE:(\w+)@@@', text)
                columns_match = re.search(r'@@@TABLE_COLUMNS:([\w.,\s]+)@@@', text)

                if style_match or columns_match:
                    next_element = para._element.getnext()

                    # Check if immediate next element is a table
                    table_obj = None
                    if next_element is not None and next_element.tag.endswith('tbl'):
                        # Find the Table object that wraps this XML element
                        for table in doc.tables:
                            if table._element == next_element:
                                table_obj = table
                                break

                    if table_obj is not None:
                        if style_match:
                            style_name = style_match.group(1)
                            table_operations.append((table_obj, style_name))
                            styled_table_elements.add(table_obj._element)
                        if columns_match:
                            column_width_operations.append((table_obj, columns_match.group(1)))
                    else:
                        print(f"  Table marker '{text[:60]}' at paragraph {i} not followed by table")

                    # Mark marker paragraph for deletion
                    paragraphs_to_trim.append((i, ''))

            # Handle paragraph style markers (must start the paragraph)
            elif text.startswith('@@@STYLE:'):
                style_match = re.search(r'@@@STYLE:(\w+)@@@', text)
                if style_match:
                    saved_style_name = style_match.group(1)

                    # Check for inline case (both markers in same paragraph)
                    if text.endswith('@@@END_STYLE@@@'):
                        cleaned_text = re.sub(r'@@@STYLE:\w+@@@|@@@END_STYLE@@@', '', text).strip()
                        paragraph_operations.append((i, i, saved_style_name))
                        paragraphs_to_trim.append((i, cleaned_text))
                    else:
                        # Multi-paragraph case: save start position and transition
                        paragraph_start = i
                        state = PARAGRAPH_MARKER

        elif state == PARAGRAPH_MARKER:
            # Skip code blocks while searching for END marker
            if para.runs and para.runs[0].style and 'Verbatim' in para.runs[0].style.name:
                continue

            # Look for END marker
            if text.endswith('@@@END_STYLE@@@'):
                paragraph_operations.append((paragraph_start, i, saved_style_name))

                # Mark all paragraphs in range for trimming
                for k in range(paragraph_start, i + 1):
                    para_to_trim = doc.paragraphs[k]

                    if para_to_trim.runs and para_to_trim.runs[0].style and 'Verbatim' in para_to_trim.runs[0].style.name:
                        continue

                    original_text = para_to_trim.text
                    cleaned_text = re.sub(r'@@@STYLE:\w+@@@|@@@END_STYLE@@@', '', original_text).strip()

                    if original_text.strip() != cleaned_text:
                        paragraphs_to_trim.append((k, cleaned_text))

                # Transition back to NO_MARKER
                state = NO_MARKER
                saved_style_name = None
                paragraph_start = None

    # Check for unclosed blocks
    if state == PARAGRAPH_MARKER:
        print(f"  Found START marker at paragraph {paragraph_start} but no valid END marker")

    # Add default styles to unstyled tables
    for table in doc.tables:
        if table._element not in styled_table_elements:
            table_operations.append((table, 'DilonTable_List'))

    # PHASE 2: Execute all collected operations
    # Apply table styles
    for table, style_name in table_operations:
        try:
            apply_table_style_to_object(table, style_name)
        except Exception as e:
            print(f"  Could not apply style '{style_name}' to table: {e}")

    # Apply column widths
    if column_width_operations:
        section = doc.sections[0]
        available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches

        for table, spec_text in column_width_operations:
            widths = parse_column_widths(spec_text, len(table.columns))
            if widths is None:
                print(f"  Invalid @@@TABLE_COLUMNS@@@ spec '{spec_text}' for a {len(table.columns)}-column table; leaving default widths")
                continue
            try:
                apply_table_column_widths(table, widths, available_width)
            except ValueError as e:
                print(f"  Could not apply column widths '{spec_text}': {e}")

    # Apply paragraph styles
    for start_idx, end_idx, style_name in paragraph_operations:
        try:
            paras = [doc.paragraphs[k] for k in range(start_idx, end_idx + 1)]
            apply_paragraph_style_to_objects(paras, style_name)
        except Exception as e:
            print(f"  Could not apply style '{style_name}' to paragraphs {start_idx}-{end_idx}: {e}")

    # Clean up marker paragraphs (reverse order to maintain indices)
    paragraphs_to_delete = []
    for idx, cleaned_text in sorted(paragraphs_to_trim, reverse=True):
        para = doc.paragraphs[idx]

        if cleaned_text:
            # Trim markers but keep content
            for run in para.runs:
                run._element.getparent().remove(run._element)
            para.add_run(cleaned_text)
        else:
            p_element = para._element
            prev_element = p_element.getprevious()
            next_element = p_element.getnext()
            between_two_tables = (
                prev_element is not None and prev_element.tag.endswith('tbl')
                and next_element is not None and next_element.tag.endswith('tbl')
            )

            if between_two_tables:
                # Deleting this paragraph would leave the two tables
                # directly adjacent in the XML, which Word merges into a
                # single visual table on open. Empty it instead so a
                # blank paragraph keeps separating them.
                for run in para.runs:
                    run._element.getparent().remove(run._element)
            else:
                # Delete paragraph entirely (marker only)
                paragraphs_to_delete.append(idx)
                p_element.getparent().remove(p_element)

    if paragraphs_to_delete:
        print(f"  Removed {len(paragraphs_to_delete)} marker paragraph(s)")

    doc.save(docx_file)


def add_field_simple_run(paragraph, instr, cached_text):
    """
    Append a Word fldSimple field (e.g. STYLEREF/SEQ) to a paragraph, with
    cached_text as the field's placeholder display value until Word
    recalculates it (Word shows this cached result until the field is
    updated - see set_update_fields_on_open()).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr)
    run_el = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    r_pr.append(OxmlElement('w:noProof'))
    run_el.append(r_pr)
    t_el = OxmlElement('w:t')
    t_el.text = cached_text
    run_el.append(t_el)
    fld.append(run_el)
    paragraph._p.append(fld)


def _narrow_bookmark(body, bookmark_start_el, new_start_anchor_el, new_end_anchor_el):
    """
    Removes bookmark_start_el and its matching bookmarkEnd (same
    w:id - bookmarkEnd carries no name, only bookmarkStart does, so
    they're paired by id) from wherever they currently sit, then
    re-inserts a fresh bookmarkStart/bookmarkEnd pair of the same name
    immediately before new_start_anchor_el and immediately after
    new_end_anchor_el. Works whether the anchors are paragraph-level
    (narrowing a whole-section bookmark down to one heading paragraph)
    or run-level (narrowing a figure's image+caption bookmark down to
    just its number's field runs) - lxml sibling insertion doesn't
    care which level it operates at, only that both anchors share the
    same parent as each other.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bookmark_id = bookmark_start_el.get(qn('w:id'))
    name = bookmark_start_el.get(qn('w:name'))
    bookmark_end_el = None
    for el in body.iter(qn('w:bookmarkEnd')):
        if el.get(qn('w:id')) == bookmark_id:
            bookmark_end_el = el
            break

    bookmark_start_el.getparent().remove(bookmark_start_el)
    if bookmark_end_el is not None:
        bookmark_end_el.getparent().remove(bookmark_end_el)

    new_start = OxmlElement('w:bookmarkStart')
    new_start.set(qn('w:id'), bookmark_id)
    new_start.set(qn('w:name'), name)
    new_end = OxmlElement('w:bookmarkEnd')
    new_end.set(qn('w:id'), bookmark_id)

    new_start_anchor_el.addprevious(new_start)
    new_end_anchor_el.addnext(new_end)
    return name


def _find_bookmark_start_before(element, prefix):
    """Walks backward through consecutive <w:bookmarkStart> siblings
    immediately preceding element, returning the first whose name
    starts with prefix (e.g. 'fig:'), or None if none match before a
    non-bookmark sibling is hit."""
    from docx.oxml.ns import qn

    prev = element.getprevious()
    while prev is not None and prev.tag == qn('w:bookmarkStart'):
        name = prev.get(qn('w:name'))
        if name and name.startswith(prefix):
            return prev
        prev = prev.getprevious()
    return None


def apply_figure_captions(docx_file):
    """
    Turn Pandoc's implicit-figure caption paragraphs into Word-native,
    auto-numbered figure captions.

    MARKDOWN_STYLING_GUIDE.md documents the authoring convention:
    `![Description.](path.png){#fig:label}` - the caption text lives in
    the image's alt-text brackets, with no manually-typed figure number.
    Because the "Captioned Figure" / "Image Caption" paragraph styles are
    defined in the reference template (see TEMPLATE_Word_Base.docx),
    Pandoc routes an image-with-alt-text into a distinct "Image Caption"
    paragraph - a reliable, unambiguous signal that this is a real figure
    caption, not just a plain paragraph that happens to follow an
    intentionally uncaptioned image (`![](path.png)`, no alt text, no
    caption paragraph emitted at all).

    For each match, this replaces the caption paragraph's content with:
        Figure {STYLEREF 2 \\s}.{SEQ Figure \\* ARABIC \\s 2} - Description.
    styled "Caption", so the chapter number is tied live to the nearest
    Heading 2 and the running figure count resets at each Heading 2
    boundary - both computed by Word itself, not this script. The exact
    field-code syntax here was extracted from a caption Word's own
    Insert Caption feature generated (see FIGURE_CAPTION_TEST.docx),
    rather than hand-derived, since a subtly wrong field switch would
    silently fail inside Word.
    """
    doc = Document(docx_file)

    count = 0
    for para in doc.paragraphs:
        if para.style is None or para.style.name != 'Image Caption':
            continue

        description = para.text.strip()

        for run in list(para.runs):
            run._element.getparent().remove(run._element)

        image_para_el = para._p.getprevious()
        bookmark_start_el = _find_bookmark_start_before(image_para_el, 'fig:') if image_para_el is not None else None

        para.style = doc.styles['Caption']
        para.add_run('Figure ')
        figure_start_el = para._p[-1]
        add_field_simple_run(para, ' STYLEREF 2 \\s ', '1')
        para.add_run('.')
        add_field_simple_run(para, ' SEQ Figure \\* ARABIC \\s 2 ', '1')
        figure_end_el = para._p[-1]
        if description:
            para.add_run(f' - {description}')
        count += 1

        if bookmark_start_el is not None:
            _narrow_bookmark(doc.element.body, bookmark_start_el, figure_start_el, figure_end_el)

    if count:
        doc.save(docx_file)
        print(f"  Converted {count} figure caption(s) to auto-numbered Word captions")

    return count


class ListNestingError(ValueError):
    """Raised when an ordered list (plain #. or, later, a step list)
    nests more than three levels deep - MARKDOWN_STYLING_GUIDE.md's
    documented cap. Bullet lists are not subject to this cap - only
    lists backed by a decimal-numbered abstractNum (Pandoc's #.
    auto-number marker, or the Dilon Step List numbering added in a
    later plan)."""


def _num_id_to_abstract_map(numbering_element):
    """Maps every <w:num numId="..."> in numbering_element to the
    <w:abstractNum abstractNumId="..."> it references."""
    from docx.oxml.ns import qn

    mapping = {}
    for num_el in numbering_element.findall(qn('w:num')):
        abstract_ref = num_el.find(qn('w:abstractNumId'))
        if abstract_ref is not None:
            mapping[num_el.get(qn('w:numId'))] = abstract_ref.get(qn('w:val'))
    return mapping


def _decimal_abstract_num_ids(numbering_element):
    """Returns the set of abstractNumIds whose level-0 numFmt is
    "decimal" - Pandoc's own native #. auto-number marker (verified
    against TEMPLATE_Word_Base.docx: a #. list gets numFmt="decimal",
    a - bullet list gets numFmt="bullet", both under the same "Compact"
    paragraph style - style alone can't distinguish them, numFmt can)."""
    from docx.oxml.ns import qn

    decimal_ids = set()
    for abstract_el in numbering_element.findall(qn('w:abstractNum')):
        abstract_id = abstract_el.get(qn('w:abstractNumId'))
        for lvl_el in abstract_el.findall(qn('w:lvl')):
            if lvl_el.get(qn('w:ilvl')) != '0':
                continue
            num_fmt_el = lvl_el.find(qn('w:numFmt'))
            if num_fmt_el is not None and num_fmt_el.get(qn('w:val')) == 'decimal':
                decimal_ids.add(abstract_id)
            break
    return decimal_ids


def _paragraph_num_id_and_ilvl(p_element):
    """Returns (numId, ilvl) as strings for a <w:p> oxml element carrying
    list numbering, or (None, None) if it carries none."""
    from docx.oxml.ns import qn

    p_pr = p_element.find(qn('w:pPr'))
    if p_pr is None:
        return None, None
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is None:
        return None, None
    num_id_el = num_pr.find(qn('w:numId'))
    ilvl_el = num_pr.find(qn('w:ilvl'))
    num_id = num_id_el.get(qn('w:val')) if num_id_el is not None else None
    ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else None
    return num_id, ilvl


def validate_list_nesting_depth(docx_file, max_ilvl=2):
    """
    Fails loudly if any ordered-list paragraph (Pandoc's native #.
    marker, or - in a later plan - a Dilon step list) in docx_file
    nests deeper than max_ilvl (0-indexed; the default 2 means three
    levels total: 0, 1, 2). Bullet lists are exempt - the three-level
    cap is a step/ordered-list authoring rule
    (MARKDOWN_STYLING_GUIDE.md SS5), not a general Word-document rule.

    Raises ListNestingError listing (up to 5 of) the offending items'
    own text, rather than letting Word render a 4th-level item with
    undefined/incorrect numbering.
    """
    doc = Document(docx_file)
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return
    numbering_element = numbering_part.element

    decimal_ids = _decimal_abstract_num_ids(numbering_element)
    num_id_to_abstract = _num_id_to_abstract_map(numbering_element)

    offenders = []
    for para in doc.paragraphs:
        num_id, ilvl = _paragraph_num_id_and_ilvl(para._p)
        if num_id is None or ilvl is None:
            continue
        if num_id_to_abstract.get(num_id) not in decimal_ids:
            continue
        if int(ilvl) > max_ilvl:
            offenders.append(para.text.strip() or '(empty)')

    if offenders:
        shown = "; ".join(repr(t) for t in offenders[:5])
        more = f" (+{len(offenders) - 5} more)" if len(offenders) > 5 else ""
        raise ListNestingError(
            f"Ordered list nested more than {max_ilvl + 1} levels deep at: "
            f"{shown}{more}. Reformat as separate lists or sub-bullets - "
            f"see MARKDOWN_STYLING_GUIDE.md Section 5."
        )


def remap_ordered_lists_to_dilon_step_list(docx_file):
    """
    Restyles every paragraph carrying Pandoc's native #. ordered-list
    numbering (MARKDOWN_STYLING_GUIDE.md SS5.2) onto the "Dilon Step
    List" Word style, leaving Pandoc's own numId/ilvl completely
    untouched - Pandoc's default ordered-list numbering is already
    correct dotted-decimal and correctly continues counting through
    nested sub-lists (verified against TEMPLATE_Word_Base.docx), so
    this only needs to change how the paragraph *looks*, not
    reimplement how it's numbered.

    Distinguishes an ordered list from a bullet list by numFmt (see
    _decimal_abstract_num_ids) rather than by paragraph style, since
    Pandoc gives both the same "Compact" style - style alone can't
    tell them apart.
    """
    doc = Document(docx_file)
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return 0
    numbering_element = numbering_part.element

    decimal_ids = _decimal_abstract_num_ids(numbering_element)
    num_id_to_abstract = _num_id_to_abstract_map(numbering_element)

    count = 0
    for para in doc.paragraphs:
        num_id, _ = _paragraph_num_id_and_ilvl(para._p)
        if num_id is None:
            continue
        if num_id_to_abstract.get(num_id) not in decimal_ids:
            continue
        para.style = doc.styles['Dilon Step List']
        count += 1

    if count:
        doc.save(docx_file)
        print(f"  Applied 'Dilon Step List' style to {count} ordered-list paragraph(s)")
    return count


_LIST_CONTINUE_MARKER_RE = re.compile(r'(?m)^([ \t]*@@@CONTINUE:#list:[\w-]+@@@[ \t]*)\n(?!\n)')


def ensure_blank_line_after_list_continue_markers(markdown_text):
    """
    Ensures a blank line separates a @@@CONTINUE:#list:name@@@ marker
    line from the ordered list that follows it, mirroring
    ensure_blank_line_after_table_markers()'s reasoning: without a
    blank line, Pandoc may fold the marker line into the list's own
    first item (or otherwise fail to treat the list as a fresh block
    with its own numId) instead of leaving it as a separate plain-text
    paragraph resolve_list_continuations() can find on its own.
    """
    return _LIST_CONTINUE_MARKER_RE.sub(lambda m: m.group(1) + '\n\n', markdown_text)


_LIST_CONTINUE_RE = re.compile(r'^@@@CONTINUE:#list:([\w-]+)@@@$')


class ListContinuationError(ValueError):
    """Raised when @@@CONTINUE:#list:name@@@ references a {#list:name}
    anchor that doesn't exist anywhere in the document, or when the
    same {#list:name} label is declared more than once."""


def resolve_list_continuations(docx_file):
    """
    Finds every @@@CONTINUE:#list:name@@@ marker paragraph in docx_file
    and makes the very next ordered-list paragraph after it continue
    counting from the {#list:name}-tagged paragraph's own list, by
    rewriting every paragraph currently sharing the *new* block's numId
    onto the *tagged* block's numId - native Word numbering then
    continues correctly because both blocks share one numId (same
    mechanism step_numbering.py's apply_step_numbering() already uses
    for its own 'continue' support, keyed here by a bookmark instead of
    a sentinel manifest).

    {#list:name} is the same bracketed-span-with-id syntax already used
    for {#fig:x}/{#step:x} - Pandoc turns it into a real bookmark on
    its own; @@@CONTINUE:#list:name@@@ matches no markdown syntax, so
    Pandoc passes it through untouched as a plain paragraph. Both
    halves are resolved here, entirely post-conversion - no markdown
    preprocessing needed beyond the blank-line normalization above.

    Raises ListContinuationError - listing every marker with no
    matching anchor, and every label declared more than once - rather
    than degrading silently, matching this codebase's halt-on-broken-
    reference convention for every other anchor type.
    """
    from docx.oxml.ns import qn

    doc = Document(docx_file)
    body = doc.element.body

    bookmark_owner = {}
    duplicates = set()
    for bookmark_start in body.iter(qn('w:bookmarkStart')):
        name = bookmark_start.get(qn('w:name'))
        if not (name and name.startswith('list:')):
            continue
        if name in bookmark_owner:
            duplicates.add(name)
            continue
        owner = bookmark_start
        while owner is not None and owner.tag != qn('w:p'):
            owner = owner.getparent()
        if owner is not None:
            bookmark_owner[name] = owner

    paragraphs = list(doc.paragraphs)
    missing = []
    resolved = 0
    markers_to_remove = []

    for i, para in enumerate(paragraphs):
        match = _LIST_CONTINUE_RE.match(para.text.strip())
        if not match:
            continue

        label = match.group(1)
        bookmark_name = f'list:{label}'
        owner_p = bookmark_owner.get(bookmark_name)
        if bookmark_name in duplicates or owner_p is None:
            missing.append(label)
            continue

        source_num_id, _ = _paragraph_num_id_and_ilvl(owner_p)
        if source_num_id is None:
            missing.append(label)
            continue

        target_num_id = None
        for later in paragraphs[i + 1:]:
            target_num_id, _ = _paragraph_num_id_and_ilvl(later._p)
            if target_num_id is not None:
                break

        if target_num_id is not None and target_num_id != source_num_id:
            for p in doc.paragraphs:
                num_id, _ = _paragraph_num_id_and_ilvl(p._p)
                if num_id != target_num_id:
                    continue
                num_id_el = p._p.find(qn('w:pPr')).find(qn('w:numPr')).find(qn('w:numId'))
                num_id_el.set(qn('w:val'), source_num_id)

        markers_to_remove.append(para._p)
        resolved += 1

    if missing or duplicates:
        parts = []
        if missing:
            parts.append("no matching {#list:...} anchor for: " + ", ".join(repr(m) for m in missing))
        if duplicates:
            parts.append("duplicate {#list:...} anchor(s): " + ", ".join(repr(d.split(':', 1)[1]) for d in sorted(duplicates)))
        raise ListContinuationError("; ".join(parts))

    for p_el in markers_to_remove:
        p_el.getparent().remove(p_el)

    if resolved:
        doc.save(docx_file)
        print(f"  Resolved {resolved} list continuation marker(s)")
    return resolved


def narrow_section_bookmarks(docx_file):
    """
    Narrows every sec:label bookmark (from `## Heading {#sec:label}`)
    from Pandoc's default whole-section span down to wrapping only the
    heading paragraph itself - verified empirically against
    TEMPLATE_Word_Base.docx that Pandoc's own bookmark for a heading id
    spans from the heading through the end of that section's content,
    which would pull an entire section's body text into a REF field
    otherwise. See spec SS7.
    """
    from docx.oxml.ns import qn

    doc = Document(docx_file)
    body = doc.element.body

    count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else None
        if not (style_name and style_name.startswith('Heading')):
            continue

        bookmark_start_el = _find_bookmark_start_before(para._p, 'sec:')
        if bookmark_start_el is None:
            continue

        _narrow_bookmark(body, bookmark_start_el, para._p, para._p)
        count += 1

    if count:
        doc.save(docx_file)
        print(f"  Narrowed {count} section bookmark(s) to their heading paragraph")
    return count


_REFERENCE_MARKER_RE = re.compile(r'\[\]\(#(fig|sec|step):([\w-]+)\)')


def preprocess_reference_markers(markdown_text):
    """
    Replaces every [](#TYPE:label) - TYPE one of fig/sec/step - with an
    XREF:TYPE:label sentinel, so Pandoc never sees an empty-text link
    (which would otherwise become a hyperlink with nothing visible to
    click). A link with real text, e.g. [see figure](#fig:label), and
    an id whose prefix isn't one of the three registered types (e.g. an
    arbitrary heading auto-id used for an ordinary hyperlink) are both
    left completely untouched - only the empty-text, registered-type
    form is intercepted.
    """
    return _REFERENCE_MARKER_RE.sub(lambda m: f'XREF:{m.group(1)}:{m.group(2)}', markdown_text)


_XREF_SENTINEL_RE = re.compile(r'XREF:(fig|sec|step):([\w-]+)')


class ReferenceResolutionError(ValueError):
    """Raised when a [](#TYPE:label) reference has no matching
    {#TYPE:label} anchor, or when the same {#TYPE:label} label is
    declared more than once, for any type passed to
    resolve_reference_markers()."""


def resolve_reference_markers(docx_file, type_resolvers):
    """
    Finds every XREF:TYPE:label sentinel (from
    preprocess_reference_markers()) in docx_file and replaces it with
    whatever type_resolvers[TYPE](paragraph, bookmark_name) appends to
    the paragraph at that point - one callback per reference type,
    letting each type own its own field-building logic (fig/sec below;
    step's lives in step_numbering.py and is passed in by the caller,
    since steps are compiler-only) while sharing one sentinel scan, one
    bookmark-duplicate check, and one halt-on-error path.

    Raises ReferenceResolutionError listing every sentinel with no
    matching bookmark and every bookmark name declared more than once,
    across every type present in type_resolvers - not narrowed to one
    type at a time - rather than resolving what it can and leaving the
    rest silently broken.
    """
    from docx.oxml.ns import qn

    doc = Document(docx_file)
    body = doc.element.body

    prefixes = tuple(f'{t}:' for t in type_resolvers)
    bookmark_names = {}
    duplicates = set()
    for el in body.iter(qn('w:bookmarkStart')):
        name = el.get(qn('w:name'))
        if not (name and name.startswith(prefixes)):
            continue
        if name in bookmark_names:
            duplicates.add(name)
        bookmark_names[name] = True

    if duplicates:
        raise ReferenceResolutionError(
            "duplicate anchor(s): " + ", ".join(repr(d) for d in sorted(duplicates))
        )

    missing = []
    resolved = 0

    for para in doc.paragraphs:
        matches = list(_XREF_SENTINEL_RE.finditer(para.text))
        if not matches:
            continue

        original_text = para.text
        for run in list(para.runs):
            run._element.getparent().remove(run._element)

        cursor = 0
        for match in matches:
            before_text = original_text[cursor:match.start()]
            if before_text:
                para.add_run(before_text)

            ref_type, label = match.group(1), match.group(2)
            bookmark_name = f'{ref_type}:{label}'
            if bookmark_name in bookmark_names and ref_type in type_resolvers:
                type_resolvers[ref_type](para, bookmark_name)
                resolved += 1
            else:
                missing.append(bookmark_name)

            cursor = match.end()

        trailing_text = original_text[cursor:]
        if trailing_text:
            para.add_run(trailing_text)

    if missing:
        raise ReferenceResolutionError(
            "no matching anchor for: " + ", ".join(repr(m) for m in missing)
        )

    if resolved:
        doc.save(docx_file)
        print(f"  Resolved {resolved} cross-reference(s)")
    return resolved


def resolve_fig_reference(para, bookmark_name):
    """type_resolvers['fig'] callback: a plain hyperlinked REF against
    the narrowed "Figure N.M" bookmark (apply_figure_captions()) - no
    \\r needed, since the bookmark wraps literal text and our own
    STYLEREF/SEQ field runs directly, not a native list item's own
    rendered marker."""
    add_complex_field(para, f'REF {bookmark_name} \\h', '1')


def resolve_sec_reference(para, bookmark_name):
    """type_resolvers['sec'] callback: literal "Section " + a
    hyperlinked REF \\r against the narrowed heading-paragraph bookmark
    (narrow_section_bookmarks()) - \\r extracts the heading's own live,
    natively-rendered outline number."""
    para.add_run('Section ')
    add_complex_field(para, f'REF {bookmark_name} \\r \\h', '1')


def center_image_paragraphs(docx_file):
    """
    Center every paragraph that contains an inline image, captioned or
    not.

    A captioned figure's image lands in a "Captioned Figure"-styled
    paragraph, but an UNcaptioned image (`![](path.png)`, no alt text -
    see apply_figure_captions()) lands in a plain "Normal"-styled
    paragraph indistinguishable, by style alone, from an ordinary text
    paragraph. So this detects "contains an image" structurally (a
    DrawingML <a:blip> descendant) rather than by paragraph style, which
    covers both cases with one pass.
    """
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(docx_file)

    count = 0
    for para in doc.paragraphs:
        if para._p.findall('.//' + qn('a:blip')):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            count += 1

    if count:
        doc.save(docx_file)
        print(f"  Centered {count} image paragraph(s)")

    return count


def set_update_fields_on_open(docx_file):
    """
    Force Word to recalculate all fields (STYLEREF/SEQ figure numbers, TOC
    page numbers, etc.) the moment the document is opened, rather than
    showing this script's cached placeholder values until the user
    manually selects-all-and-presses-F9.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(docx_file)
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        el = OxmlElement('w:updateFields')
        el.set(qn('w:val'), 'true')
        settings.insert(0, el)
        doc.save(docx_file)


_LOGO_PATH = Path(__file__).resolve().parent.parent / "templates" / "assets" / "dilon_logo.png"


def _clear_container(container):
    """Remove every paragraph/table from a python-docx header/footer (or
    body) container, leaving it empty so populate_header()/
    populate_footer() can build fresh content into it - rather than
    appending after whatever placeholder paragraph Word/python-docx
    always creates by default."""
    for child in list(container._element):
        container._element.remove(child)


def add_complex_field(paragraph, instr, cached_text):
    """
    Append a complex Word field (begin/instrText/separate/cached-result/
    end run sequence) to a paragraph. Used for PAGE/NUMPAGES fields, which
    - unlike the STYLEREF/SEQ fields add_field_simple_run() handles -
    are conventionally authored as complex fields (that's what Word's own
    Insert Page Number feature emits), not fldSimple.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _fld_char(kind):
        run_el = OxmlElement('w:r')
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), kind)
        run_el.append(fld)
        return run_el

    paragraph._p.append(_fld_char('begin'))

    instr_run = OxmlElement('w:r')
    instr_el = OxmlElement('w:instrText')
    instr_el.set(qn('xml:space'), 'preserve')
    instr_el.text = instr
    instr_run.append(instr_el)
    paragraph._p.append(instr_run)

    paragraph._p.append(_fld_char('separate'))

    text_run = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = cached_text
    text_run.append(t_el)
    paragraph._p.append(text_run)

    paragraph._p.append(_fld_char('end'))


def strip_leading_empty_paragraphs(document):
    """
    Remove empty paragraphs from the very start of `document`'s body,
    stopping at the first table, non-empty paragraph, or sectPr.

    A docx's body always needs at least one paragraph if it has no other
    content, so TEMPLATE_Word_Base.docx - despite being header/footer/
    styles-only by design - carries one empty paragraph before its
    sectPr as a structural artifact of being a real Word file.
    compose_documents() builds the merged document as
    Composer(Document(doc_paths[0])), which keeps doc_paths[0]'s body
    content as-is - so leaving that empty paragraph in Part A produced a
    stray leading blank line at the top of every compiled document. Call
    this on Part A right after populate_header()/populate_footer(),
    before saving it for compose_documents() to use as its base.
    """
    from docx.oxml.ns import qn

    body = document._element.body
    for child in list(body):
        if child.tag != qn('w:p'):
            break
        text = ''.join(node.text or '' for node in child.iter(qn('w:t')))
        if text.strip():
            break
        body.remove(child)


def populate_header(document, metadata):
    """
    Build the running header (logo | Title/Number | Rev | Page N of M)
    directly via python-docx, in place of a docxtpl-rendered Jinja
    version baked into the template. Shared by dilon-document-compiler
    and dilon-document-form-compiler, since both use the same
    TEMPLATE_Word_Base.docx.

    Column widths are sized to fit "999 of 999" in the page-number column
    without wrapping/truncating (regulatory documents can run well past
    single or double-digit page counts), while keeping the table's total
    width within the page's available content width (previously
    overflowed it) - see the base template's page setup for margins.
    """
    from docx.shared import Inches, Pt, Twips, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    _clear_container(header)

    # Logo/Rev/Page column widths match an approved, hand-tuned Dilon Word
    # document's own header table; the Title/Number column absorbs
    # whatever width is left so the table always fills the page's full
    # content width.
    available_width = section.page_width - section.left_margin - section.right_margin
    logo_width = Twips(1525)
    rev_width = Twips(1255)
    page_num_width = Twips(1625)
    title_width = Emu(available_width - logo_width - rev_width - page_num_width)
    col_widths = [logo_width, title_width, rev_width, page_num_width]
    table_width = Emu(available_width)
    table = header.add_table(rows=1, cols=4, width=table_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._element.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), 'auto')
        borders.append(edge_el)
    tbl_pr.append(borders)
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        side_el = OxmlElement(f'w:{side}')
        side_el.set(qn('w:w'), '29')
        side_el.set(qn('w:type'), 'dxa')
        cell_mar.append(side_el)
    tbl_pr.append(cell_mar)

    row = table.rows[0]
    row.height = Inches(945 / 1440)
    for idx, width in enumerate(col_widths):
        table.columns[idx].width = width
        row.cells[idx].width = width
        row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Cell 0: logo
    logo_cell = row.cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _LOGO_PATH.exists():
        logo_para.add_run().add_picture(str(_LOGO_PATH), width=Inches(0.88), height=Inches(0.656))

    # Cell 1: Title / Number - both left-justified (paragraph default;
    # no indent/tab needed, unlike the hanging-indent + leading-tab
    # combo the original template used, which left the "Number:" line
    # landing at Word's default tab stop instead of the true left edge)
    title_cell = row.cells[1]
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.add_run(f"Title: {metadata.get('title', '')}")
    number_para = title_cell.add_paragraph()
    number_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    number_para.add_run(f"Number: {metadata.get('doc_number', '')}")

    # Cell 2: Rev
    rev_cell = row.cells[2]
    rev_para = rev_cell.paragraphs[0]
    rev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rev_run = rev_para.add_run(f"Rev {metadata.get('current_revision', '')}")
    rev_run.font.bold = True

    # Cell 3: Page N of M (live PAGE/NUMPAGES fields)
    page_cell = row.cells[3]
    page_label_para = page_cell.paragraphs[0]
    page_label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_label_run = page_label_para.add_run('Page')
    page_label_run.font.bold = True

    page_field_para = page_cell.add_paragraph()
    page_field_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_complex_field(page_field_para, 'page \\* arabic', '1')
    mid_run = page_field_para.add_run(' of ')
    mid_run.font.bold = True
    add_complex_field(page_field_para, 'numpages ', '1')
    for run in page_field_para.runs:
        run.font.bold = True


def populate_footer(document, metadata):
    """
    Build the running footer as a 3-column/2-row table (doc_number/rev |
    ECO # | revision date, then a single full-width cell for the
    confidentiality boilerplate) directly via python-docx, in place of a
    docxtpl-rendered Jinja version. Shared by dilon-document-compiler and
    dilon-document-form-compiler.

    A table (not tab-stopped paragraphs, the previous approach) is used
    so the three ID-line pieces sit in genuinely equal-width columns
    regardless of their own text length, matching the header table's
    already-table-based layout.

    ECO number/date come from the latest (last) entry in the front
    matter's `revisions` list, since neither is its own flat top-level
    field.
    """
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_container(footer)

    revisions = metadata.get('revisions') or []
    latest_revision = revisions[-1] if revisions else {}

    # Divided in twips (not EMU): gridCol/tcW widths are themselves
    # written in twips, so splitting the EMU-based available_width three
    # ways and only then converting each third back to twips would round
    # each of the three shares independently and could overshoot the true
    # available width (observed: 1 twip over). Dividing in twips first
    # keeps the three columns' sum exactly equal to available_width.
    available_width_twips = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    available_width = Twips(available_width_twips)

    table = footer.add_table(rows=2, cols=3, width=available_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Equal thirds - available_width_twips isn't always evenly divisible
    # by 3, so any leftover twip(s) go to the last column so the three
    # still sum to exactly the full available width.
    third_twips = available_width_twips // 3
    col_widths = [
        Twips(third_twips),
        Twips(third_twips),
        Twips(available_width_twips - 2 * third_twips),
    ]
    for idx, width in enumerate(col_widths):
        table.columns[idx].width = width
        for cell in table.columns[idx].cells:
            cell.width = width

    tbl_pr = table._element.tblPr
    # Top/bottom rule lines framing the whole footer, no internal or side
    # borders - matches the look of the paragraph-border footer this
    # table replaces (a rule above the ID line, a rule below the notice).
    borders = OxmlElement('w:tblBorders')
    for edge, val in (
        ('top', 'single'), ('bottom', 'single'),
        ('left', 'nil'), ('right', 'nil'),
        ('insideH', 'nil'), ('insideV', 'nil'),
    ):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), val)
        if val == 'single':
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '4')
            edge_el.set(qn('w:color'), 'auto')
        borders.append(edge_el)
    tbl_pr.append(borders)

    # Tight cell margins (matches the header table) so row 1's height
    # tracks the font size instead of Word's default ~0.08in cell padding.
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        side_el = OxmlElement(f'w:{side}')
        side_el.set(qn('w:w'), '29')
        side_el.set(qn('w:type'), 'dxa')
        cell_mar.append(side_el)
    tbl_pr.append(cell_mar)

    FOOTER_FONT_SIZE = Pt(9)

    def _add_footer_run(paragraph, text):
        # The 'Normal' style's own default spacing (6pt before + 6pt
        # after every paragraph) would otherwise pad out each footer
        # line - zeroed here so the table hugs the text as tightly as
        # the row-height settings above intend.
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.font.size = FOOTER_FONT_SIZE
        return run

    # Row 1: doc_number/rev (left) | ECO # (center) | revision date (right)
    # - height pinned to the font size itself (AT_LEAST so a taller value
    # never gets clipped, but nothing forces it any taller than that).
    row1 = table.rows[0]
    row1.height = FOOTER_FONT_SIZE
    row1.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    id_cell, eco_cell, date_cell = row1.cells
    id_para = id_cell.paragraphs[0]
    id_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_footer_run(id_para, f"{metadata.get('doc_number', '')} Rev {metadata.get('current_revision', '')}")

    eco_para = eco_cell.paragraphs[0]
    eco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_footer_run(eco_para, latest_revision.get('eco_number', ''))

    date_para = date_cell.paragraphs[0]
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_footer_run(date_para, f"Revision Date: {latest_revision.get('eco_date', '')}")

    # Row 2: a single cell spanning all 3 columns for the confidentiality
    # boilerplate - height is left on Word's default AUTO rule, so it
    # grows only as much as its two lines of text need.
    notice_cell = table.rows[1].cells[0].merge(table.rows[1].cells[2])
    notice_lines = [
        "This document and information contained within is confidential and proprietary to Dilon Technologies.",
        "All unauthorized use and/or reproduction is prohibited.",
    ]
    notice_para = notice_cell.paragraphs[0]
    notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_footer_run(notice_para, notice_lines[0])
    second_para = notice_cell.add_paragraph()
    second_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_footer_run(second_para, notice_lines[1])


def extract_yaml_and_markdown(md_file):
    """
    Extract YAML front matter and Markdown body from a Markdown file.

    Returns:
        tuple: (yaml_data dict, markdown_body str)
    """
    import yaml

    with open(md_file, 'r', encoding='utf-8-sig') as f:
        content = f.read()

    # Check for YAML front matter
    yaml_pattern = r'^---\s*\n(.*?)\n---\s*\n(.*)$'
    match = re.match(yaml_pattern, content, re.DOTALL)

    if match:
        yaml_text = match.group(1)
        markdown_body = match.group(2)
        yaml_data = yaml.safe_load(yaml_text)
        return yaml_data, markdown_body
    else:
        return {}, content


_TABLE_MARKER_RUN = re.compile(
    r'^(?:[ \t]*@@@TABLE_(?:STYLE:\w+|COLUMNS:[\w.,\s]+)@@@[ \t]*\n)+',
    re.MULTILINE,
)


def _insert_blank_line_if_needed(match):
    already_blank = match.string[match.end():match.end() + 1] == '\n'
    return match.group(0) if already_blank else match.group(0) + '\n'


def ensure_blank_line_after_table_markers(markdown_text):
    """
    Ensure a blank line separates a run of one or more stacked
    @@@TABLE_STYLE:...@@@ / @@@TABLE_COLUMNS:...@@@ marker lines from the
    table that follows them.

    MARKDOWN_STYLING_GUIDE.md documents "no blank lines between the marker
    and the table" as the convention, but Pandoc's pipe-table parser only
    recognizes a table when it is preceded by a blank line - without one,
    the marker line(s) and the entire table are merged into a single
    garbled text paragraph and no table is produced at all. Normalizing
    here (before Pandoc ever sees the markdown) keeps the documented
    authoring convention working while still producing a real, styleable,
    width-settable table.

    Uses a callback (rather than a trailing negative-lookahead assertion)
    to decide whether a blank line already follows the marker run: a
    lookahead combined with the `+` repetition over marker lines can
    backtrack to a *partial* match when the full run is already followed
    by a blank line, inserting a spurious blank line in the middle of a
    stacked marker run instead of leaving it alone.
    """
    return _TABLE_MARKER_RUN.sub(_insert_blank_line_if_needed, markdown_text)


_IMAGE_ONLY_LINE = re.compile(
    r'^[ \t]*!\[[^\]\n]*\]\([^)\n]*\)(?:\{[^}\n]*\})?[ \t]*$'
)


def ensure_blank_line_between_images(markdown_text):
    """
    Ensure a blank line separates two consecutive image-only markdown
    lines (e.g. two figures declared back-to-back with no prose between
    them).

    Pandoc's implicit-figures extension only promotes an image to a
    captioned figure when its paragraph contains nothing but that one
    image (MARKDOWN_STYLING_GUIDE.md SS4.1's "caption in the alt text"
    convention relies on this). Two image lines with no blank line
    between them are, per ordinary Markdown paragraph rules, the SAME
    paragraph - so neither satisfies "just one image" and Pandoc silently
    drops figure treatment for BOTH: no caption, no {#fig:...} bookmark,
    no numbering. Normalizing here (before Pandoc ever sees the markdown)
    means the natural habit of stacking figures one after another keeps
    working, mirroring ensure_blank_line_after_table_markers() above for
    the equivalent problem with @@@TABLE_STYLE@@@/@@@TABLE_COLUMNS@@@
    markers.
    """
    lines = markdown_text.split('\n')
    result = []
    for i, line in enumerate(lines):
        result.append(line)
        if (
            i + 1 < len(lines)
            and _IMAGE_ONLY_LINE.match(line)
            and _IMAGE_ONLY_LINE.match(lines[i + 1])
        ):
            result.append('')
    return '\n'.join(result)


def render_jinja(markdown_text, metadata):
    """
    Pre-render a markdown body's {{field}} references against the same
    front-matter dict used for the header/footer/signature page, before
    Pandoc ever sees the text - so a document's body can reference e.g.
    {{doc_number}}/{{title}} and have it resolve from the one YAML source
    of truth instead of being hand-duplicated. A body with no {{...}} in
    it is returned unchanged. Escape literal {{ }} with Jinja2's own
    {% raw %}...{% endraw %} if a document must discuss the syntax itself.
    """
    return _JINJA_ENV.from_string(markdown_text).render(**metadata)


def markdown_to_docx(markdown_text, output_file, reference_doc=None, resource_dir=None, include_toc=True):
    """
    Convert Markdown to a Word document using Pandoc.

    Args:
        markdown_text: Markdown content as string
        output_file: Path to save the Word document
        reference_doc: Optional path to reference document for styles
        resource_dir: Optional directory relative image paths in the
            markdown are resolved against (MARKDOWN_STYLING_GUIDE.md
            documents "image path is relative to the markdown file" - this
            is what makes that true regardless of the caller's own cwd).
            Without it, Pandoc resolves relative image paths against the
            process's current working directory, silently embedding a
            placeholder instead of the real image if that happens to
            differ from the markdown file's own directory.
        include_toc: Whether to generate a table of contents (--toc). A
            form document has no heading structure to build a TOC from -
            the form compiler passes False.
    """
    markdown_text = ensure_blank_line_after_table_markers(markdown_text)
    markdown_text = ensure_blank_line_between_images(markdown_text)

    # Create temporary markdown file
    temp_md = Path(output_file).parent / "_temp_content.md"
    with open(temp_md, 'w', encoding='utf-8') as f:
        f.write(markdown_text)

    # Build Pandoc command
    pandoc_cmd = [
        'pandoc',
        str(temp_md),
        '-o', str(output_file),
        '--standalone',
        '--from=markdown+smart+backtick_code_blocks+fenced_code_attributes+raw_html',
        '--wrap=preserve'
    ]
    if include_toc:
        pandoc_cmd.extend(['--toc', '--toc-depth=6'])

    # Add reference document if provided
    if reference_doc:
        pandoc_cmd.extend(['--reference-doc', str(reference_doc)])

    if resource_dir:
        pandoc_cmd.extend(['--resource-path', str(resource_dir)])

    # Use Pandoc to convert Markdown to Word
    try:
        result = subprocess.run(pandoc_cmd, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as e:
        print(f"Error running Pandoc: {e.stderr}")
        raise

    # Pandoc can exit 0 while still warning about problems (e.g. an
    # unresolved image path silently replaced with a placeholder) - surface
    # those warnings even on success instead of only printing on failure.
    if result.stderr:
        print(f"  Pandoc warnings:\n{result.stderr}")

    # Clean up temp file
    temp_md.unlink()


def compose_documents(*doc_paths):
    """
    Merge 2+ Word documents in order via docxcompose's Composer, appending
    each subsequent document to the first.

    Returns the Composer wrapping the merged document - call
    .save(output_path) on the result. Generalizes the compiler's original
    2-document merge_word_documents() to N documents so both the full A-B-C-D
    compile pipeline and the form compiler's 2-document pipeline share one
    implementation.
    """
    composer = Composer(Document(doc_paths[0]))
    for doc_path in doc_paths[1:]:
        composer.append(Document(doc_path))
    return composer
