"""
Form-specific markdown-authoring functions for dilon-document-form-compiler.

These live outside lib/dilon_docx_common.py deliberately: ordinary content
documents (dilon-document-compiler) never need them. Uses its own
@@@FORM_FIELD:Name@@@...@@@END_FORM_FIELD@@@ marker family so it never
collides with the shared module's @@@STYLE@@@/@@@TABLE_STYLE@@@ markers.
"""

import re

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

FORM_FIELD_RE = re.compile(
    r'@@@FORM_FIELD:(\w+)(?::([\d.]+in))?@@@(.*?)@@@END_FORM_FIELD@@@',
    re.DOTALL,
)

FORM_SECTION_HEADER_STYLE_NAME = "Form Section Header"

BRACKET_ANNOTATION_RE = re.compile(r'^(.*?)\[([^\[\]]+)\]\s*$')


def _parse_key_value_pairs(annotation_text):
    """Split a comma-separated `key=value,key=value` string into a dict.
    Entries without '=' are ignored, not an error."""
    annotations = {}
    for entry in annotation_text.split(','):
        if '=' not in entry:
            continue
        key, value = entry.split('=', 1)
        annotations[key.strip()] = value.strip()
    return annotations


def parse_bracket_annotations(text):
    """
    Split a trailing `[key=value,key=value,...]` annotation off `text`.

    Returns (cleaned_text, annotations) where annotations is a
    str -> str dict (values are not type-converted here - callers parse
    ints/floats themselves). Returns (text, {}) if there's no trailing
    bracket.
    """
    match = BRACKET_ANNOTATION_RE.match(text)
    if not match:
        return text, {}
    return match.group(1), _parse_key_value_pairs(match.group(2))


ROW_ANNOTATION_RE = re.compile(r'^(.*?)\s*\{([^{}]+)\}\s*$')


def parse_row_annotation(line):
    """
    Split a trailing ` {key=value,key=value,...}` row-level annotation
    off `line` - a distinct delimiter ({} instead of []) from the
    per-pair [...] annotations, so the two never collide. Returns
    (remaining_line, annotations_dict), or (line, {}) if there's no
    trailing brace group.
    """
    match = ROW_ANNOTATION_RE.match(line)
    if not match:
        return line, {}
    return match.group(1).rstrip(), _parse_key_value_pairs(match.group(2))


def parse_field_grid_block(block_text):
    """
    Parse a @@@FORM_FIELD:FieldGrid@@@ block's contents into a list of
    rows. Each row is a dict:
        {
            'pairs': [(label_text, pair_annotations_dict), ...],
            'annotations': row_annotations_dict,
        }

    Blank lines are skipped. A line that produces no non-empty pair
    label is skipped with a printed warning - never a hard failure.
    """
    rows = []
    for raw_line in block_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        remaining, row_annotations = parse_row_annotation(line)
        pair_tokens = [token.strip() for token in remaining.split('|')]
        pairs = []
        for token in pair_tokens:
            if not token:
                continue
            label_text, pair_annotations = parse_bracket_annotations(token)
            label_text = label_text.strip()
            if not label_text:
                continue
            pairs.append((label_text, pair_annotations))

        if not pairs:
            print(f"  FieldGrid: skipping unparseable row: {raw_line!r}")
            continue

        rows.append({'pairs': pairs, 'annotations': row_annotations})

    return rows


def resolve_row_settings(row_annotations):
    """Resolve a row's dir=/rows= annotations to (direction, rows),
    defaulting to ('h', 1) and warning-and-defaulting on invalid
    values."""
    direction = row_annotations.get('dir', 'h')
    if direction not in ('h', 'v'):
        print(f"  FieldGrid: invalid dir={direction!r} - using 'h'")
        direction = 'h'

    rows = 1
    if 'rows' in row_annotations:
        try:
            rows = int(row_annotations['rows'])
            if rows < 1:
                print(f"  FieldGrid: rows={row_annotations['rows']!r} must be >= 1 - using 1")
                rows = 1
        except ValueError:
            print(f"  FieldGrid: invalid rows={row_annotations['rows']!r} - using 1")

    return direction, rows


def resolve_pair_rows(pair_annotations, row_rows):
    """Resolve a pair's own rows= override, falling back to the row's
    default (row_rows) if absent or invalid."""
    if 'rows' not in pair_annotations:
        return row_rows
    try:
        rows = int(pair_annotations['rows'])
        if rows < 1:
            print(f"  FieldGrid: rows={pair_annotations['rows']!r} must be >= 1 - using {row_rows}")
            return row_rows
        return rows
    except ValueError:
        print(f"  FieldGrid: invalid rows={pair_annotations['rows']!r} - using {row_rows}")
        return row_rows


def resolve_pair_widths(pairs, row_width):
    """
    Resolve each pair's width (inches) from its optional pair=NN
    annotation.

    pairs: list of (label_text, annotations_dict) as returned by
        parse_field_grid_block().
    row_width: total available width (inches) for the row.

    Returns a list of floats (inches), one per pair, in the same order
    as `pairs`. Falls back to an even split across all pairs - with a
    printed warning - if declared percentages exceed 100, or leave a
    non-positive remainder for the undeclared pairs.
    """
    declared = {}
    for idx, (_, annotations) in enumerate(pairs):
        if 'pair' in annotations:
            try:
                declared[idx] = float(annotations['pair'])
            except ValueError:
                print(f"  FieldGrid: invalid pair={annotations['pair']!r} - ignoring")

    num_pairs = len(pairs)
    even_split = [row_width / num_pairs] * num_pairs

    if not declared:
        return even_split

    declared_total = sum(declared.values())
    undeclared_count = num_pairs - len(declared)
    remainder = 100.0 - declared_total

    if declared_total > 100 or (undeclared_count > 0 and remainder <= 0):
        print(
            f"  FieldGrid: pair= values on this row don't leave room for the "
            f"undeclared pairs (declared {declared_total}%) - falling back to an even split"
        )
        return even_split

    undeclared_share = remainder / undeclared_count if undeclared_count else 0.0
    widths = []
    for idx in range(num_pairs):
        pct = declared.get(idx, undeclared_share)
        widths.append(row_width * pct / 100.0)
    return widths


def resolve_label_width(pair_annotations, pair_width, row_dir):
    """
    Resolve a horizontal pair's label sub-cell width (inches) from its
    optional label=NN annotation. Ignored (with a warning) if the row is
    dir=v, since there's no left/right split to control. Returns
    (label_width, blank_width), or (None, None) for a vertical row.
    """
    if row_dir == 'v':
        if 'label' in pair_annotations:
            print("  FieldGrid: label= ignored on a dir=v row (no left/right split to control)")
        return None, None

    pct = 50.0
    if 'label' in pair_annotations:
        try:
            pct = float(pair_annotations['label'])
        except ValueError:
            print(f"  FieldGrid: invalid label={pair_annotations['label']!r} - using 50")
            pct = 50.0

    label_width = pair_width * pct / 100.0
    blank_width = pair_width - label_width
    return label_width, blank_width


def resolve_title_flag(row_annotations):
    """Resolve a row's title= annotation to a bool. Truthy values
    (case-insensitive): 'true', '1', 'yes'. Anything else, including
    absent, is not a title row."""
    return row_annotations.get('title', '').strip().lower() in ('true', '1', 'yes')


def _build_field_grid_title_row_table(document, pairs, row_rows, row_width):
    """
    Build a title row: a 1-row, 1-column Word table spanning the full
    row width, containing only the first pair's label as a bold run.
    Any additional `|`-separated pair tokens are dropped (warned). A
    `pair=`/`label=` annotation on the label token is ignored (warned)
    - there's no percentage split or label/blank split with only one
    column. `rows=` (row-level or the label's own override) still adds
    blank paragraphs below the bold label, inside the same cell.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if len(pairs) > 1:
        print(
            f"  FieldGrid: title row has {len(pairs)} pair tokens - "
            f"only the first ({pairs[0][0]!r}) is used, the rest are dropped"
        )

    label_text, pair_annotations = pairs[0]
    if 'pair' in pair_annotations:
        print("  FieldGrid: pair= ignored on a title row (nothing to split a percentage across with only one column)")
    if 'label' in pair_annotations:
        print("  FieldGrid: label= ignored on a title row (no label/blank split to control)")

    pair_rows = resolve_pair_rows(pair_annotations, row_rows)

    table = document.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.rows[0].cells[0]
    cell.width = Inches(row_width)
    table.columns[0].width = Inches(row_width)
    run = cell.paragraphs[0].add_run(label_text)
    run.bold = True
    for _ in range(pair_rows - 1):
        cell.add_paragraph()

    return table


def build_field_grid_row_table(document, row, row_width):
    """
    Build one FieldGrid row as a standalone 1-row Word table, added to
    `document` (so it inherits the document's available styles,
    including the built-in 'Table Grid' style).

    row: one entry from parse_field_grid_block()'s return value
        ({'pairs': [...], 'annotations': {...}}).
    row_width: total available width (inches) for this row.

    A row with a truthy title= annotation is delegated to
    _build_field_grid_title_row_table() instead - see its docstring.

    Returns the created Table object.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    pairs = row['pairs']
    direction, row_rows = resolve_row_settings(row['annotations'])

    if resolve_title_flag(row['annotations']):
        return _build_field_grid_title_row_table(document, pairs, row_rows, row_width)

    pair_widths = resolve_pair_widths(pairs, row_width)

    num_cols = len(pairs) if direction == 'v' else len(pairs) * 2
    table = document.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_idx = 0
    for (label_text, pair_annotations), pair_width in zip(pairs, pair_widths):
        pair_rows = resolve_pair_rows(pair_annotations, row_rows)

        if direction == 'v':
            cell = table.rows[0].cells[col_idx]
            cell.width = Inches(pair_width)
            table.columns[col_idx].width = Inches(pair_width)
            cell.paragraphs[0].add_run(label_text)
            for _ in range(pair_rows):
                cell.add_paragraph()
            col_idx += 1
        else:
            label_width, blank_width = resolve_label_width(pair_annotations, pair_width, direction)
            label_cell = table.rows[0].cells[col_idx]
            blank_cell = table.rows[0].cells[col_idx + 1]
            label_cell.width = Inches(label_width)
            blank_cell.width = Inches(blank_width)
            table.columns[col_idx].width = Inches(label_width)
            table.columns[col_idx + 1].width = Inches(blank_width)
            label_cell.paragraphs[0].add_run(label_text)
            for _ in range(pair_rows - 1):
                blank_cell.add_paragraph()
            col_idx += 2

    return table


def insert_field_grid(doc, paragraph, block_text, max_width_inches=None):
    """
    Replace `paragraph` (a @@@FORM_FIELD:FieldGrid@@@ marker paragraph)
    with one Word table per declared row, each built via
    build_field_grid_row_table(), stacked directly adjacent with no
    separating paragraph so they read as one continuous bordered grid.
    The marker paragraph itself is removed.
    """
    available_width = _paragraph_available_width_inches(paragraph)
    if max_width_inches is not None:
        if max_width_inches > available_width:
            print(
                f"  FieldGrid: max width {max_width_inches}in exceeds the available "
                f"width ({available_width}in) - clamped"
            )
            row_width = available_width
        else:
            row_width = max_width_inches
    else:
        row_width = available_width

    rows = parse_field_grid_block(block_text)
    if not rows:
        print("  FieldGrid: no valid rows found in block - leaving no content behind")

    marker_element = paragraph._p
    for row in rows:
        table = build_field_grid_row_table(doc, row, row_width)
        marker_element.addprevious(table._element)

    marker_element.getparent().remove(marker_element)


FIELD_GRID_BLOCK_RE = re.compile(
    r'(@@@FORM_FIELD:FieldGrid(?::[\d.]+in)?@@@)(.*?)(@@@END_FORM_FIELD@@@)',
    re.DOTALL,
)


def protect_field_grid_line_breaks(markdown_text):
    """
    Pandoc's markdown-to-docx conversion joins consecutive non-blank
    source lines into a single paragraph, separated by a space -
    collapsing every declared FieldGrid row onto one line before
    apply_form_fields() ever sees the compiled document. This appends a
    Markdown hard-line-break (two trailing spaces) to every non-blank
    line inside each @@@FORM_FIELD:FieldGrid@@@...@@@END_FORM_FIELD@@@
    block, so Pandoc preserves each row as a real line break -
    python-docx surfaces a resulting <w:br/> as '\\n' in the compiled
    paragraph's .text, exactly what parse_field_grid_block() expects to
    split rows on.

    Call this on the raw markdown body before it's handed to Pandoc.
    FillLine markers and text outside FieldGrid blocks are untouched.
    """
    def _protect(match):
        open_marker, block, close_marker = match.group(1), match.group(2), match.group(3)
        protected_lines = [
            line + '  ' if line.strip() else line
            for line in block.split('\n')
        ]
        return open_marker + '\n'.join(protected_lines) + close_marker

    return FIELD_GRID_BLOCK_RE.sub(_protect, markdown_text)


def _enclosing_table_cell(paragraph):
    """Return the nearest enclosing <w:tc> XML element, or None if
    `paragraph` is body-level (not inside a table cell)."""
    tc = paragraph._p.getparent()
    while tc is not None and tc.tag != qn('w:tc'):
        tc = tc.getparent()
    return tc


def _paragraph_available_width_inches(paragraph):
    """Return the available width (inches) for a fill-to-end-of-line tab
    stop: the enclosing table cell's width if the paragraph lives in a
    table cell, otherwise the page's content width (page width minus
    margins)."""
    tc = _enclosing_table_cell(paragraph)

    if tc is not None:
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is not None:
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is not None and tc_w.get(qn('w:w')):
                # tcW is in twentieths of a point (dxa); 1440 dxa = 1 inch
                return int(tc_w.get(qn('w:w'))) / 1440.0

    doc = paragraph.part.document
    section = doc.sections[0]
    return section.page_width.inches - section.left_margin.inches - section.right_margin.inches


def _iter_table_paragraphs(table):
    """Yield every paragraph inside `table`'s cells, recursing into any
    nested tables (a cell can itself contain a table)."""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_all_paragraphs(doc):
    """Yield every paragraph in `doc` - body-level paragraphs plus every
    paragraph inside every table cell (recursively). python-docx's
    `doc.paragraphs` alone only sees body-level paragraphs, silently
    skipping markers placed inside markdown table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def underscore_until_end_of_line(paragraph, width_override=None, num_lines=1):
    """
    Rewrite `paragraph` (currently just its label text, e.g. "Work
    Order:") into one or more right-aligned, underscore-leadered blank
    lines.

    width_override: explicit blank length in inches for the (single)
        label line. Ignored - with a warning - if num_lines > 1, since
        multi-line blanks always span the full available width. Clamped
        - with a warning - if it exceeds the available width in context.
    num_lines: total number of blank lines. 1 (default) keeps the
        original single "<label>\\t" paragraph. >1 appends num_lines - 1
        additional full-width blank paragraphs immediately after the
        label paragraph, each with its own tab stop.

    Word computes the resulting fill-to-the-edge blank(s) at render time;
    nothing here counts characters.
    """
    label = paragraph.text
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    available_width = _paragraph_available_width_inches(paragraph)

    if num_lines > 1:
        if width_override is not None:
            print(f"  FillLine: width= ignored because lines={num_lines} (multi-line blanks always span the full available width)")
        line_width = available_width
    elif width_override is not None:
        if width_override > available_width:
            print(f"  FillLine: width={width_override}in exceeds the available width ({available_width}in) - clamped")
            line_width = available_width
        else:
            line_width = width_override
    else:
        line_width = available_width

    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(line_width), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES
    )
    paragraph.add_run(f"{label}\t")

    previous_p = paragraph._p
    for _ in range(num_lines - 1):
        new_p = OxmlElement('w:p')
        previous_p.addnext(new_p)
        previous_p = new_p
        new_paragraph = Paragraph(new_p, paragraph._parent)
        new_paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(line_width), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES
        )
        new_paragraph.add_run("\t")


def apply_form_fields(docx_file):
    """
    Scan docx_file for @@@FORM_FIELD:Name@@@...@@@END_FORM_FIELD@@@
    markers and apply the matching form-specific function. Currently
    supports 'FillLine' (underscore_until_end_of_line). Unrecognized
    function names are left as plain text with a printed warning, matching
    the shared module's warn-and-degrade convention - never a hard
    failure.
    """
    doc = Document(docx_file)
    changed = False
    section_number = 0

    for para in list(_iter_all_paragraphs(doc)):
        match = FORM_FIELD_RE.search(para.text)
        if not match:
            continue

        function_name, block_width, label = match.group(1), match.group(2), match.group(3)
        if function_name == "FillLine":
            cleaned_label, annotations = parse_bracket_annotations(label)

            width_override = None
            if 'width' in annotations:
                width_match = re.match(r'^([\d.]+)in$', annotations['width'])
                if width_match:
                    width_override = float(width_match.group(1))
                else:
                    print(f"  FillLine: invalid width={annotations['width']!r} - ignoring")

            num_lines = 1
            if 'lines' in annotations:
                try:
                    num_lines = int(annotations['lines'])
                    if num_lines < 1:
                        print(f"  FillLine: lines={annotations['lines']!r} must be >= 1 - using 1")
                        num_lines = 1
                except ValueError:
                    print(f"  FillLine: invalid lines={annotations['lines']!r} - ignoring")

            for run in list(para.runs):
                run._element.getparent().remove(run._element)
            para.add_run(cleaned_label)
            underscore_until_end_of_line(para, width_override=width_override, num_lines=num_lines)
            changed = True

        elif function_name == "FieldGrid":
            if _enclosing_table_cell(para) is not None:
                print("  FieldGrid marker found inside a table cell - not supported, leaving marker text as-is")
                continue
            max_width_inches = float(block_width[:-2]) if block_width else None
            insert_field_grid(doc, para, label, max_width_inches)
            changed = True

        elif function_name == "Form_Section_Header":
            if _enclosing_table_cell(para) is not None:
                print("  Form_Section_Header marker found inside a table cell - not supported, leaving marker text as-is")
                continue
            section_number += 1
            for run in list(para.runs):
                run._element.getparent().remove(run._element)
            para.add_run(f"Section {section_number} - {label}")
            try:
                para.style = doc.styles[FORM_SECTION_HEADER_STYLE_NAME]
            except KeyError:
                print(f"  Form_Section_Header: style {FORM_SECTION_HEADER_STYLE_NAME!r} not found in the template - leaving default styling")
            changed = True

        else:
            print(f"  Unrecognized @@@FORM_FIELD:{function_name}@@@ - leaving marker text as-is")

    if changed:
        doc.save(docx_file)
