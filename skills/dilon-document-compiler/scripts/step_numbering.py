# -*- coding: utf-8 -*-
"""
Auto-numbered, cross-referenceable @@@STEPS@@@ list support for
dilon-document-compiler. See
docs/superpowers/specs/2026-08-20-work-instruction-step-numbering-design.md.

Scoped to dilon-document-compiler only - not shared with
dilon-document-form-compiler (forms have no procedural-steps concept).
"""

import re
import sys
import zipfile
from pathlib import Path

from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).resolve().parents[3] / "lib"))
from dilon_docx_common import (
    add_complex_field,
    add_field_simple_run,
    _decimal_abstract_num_ids,
    _num_id_to_abstract_map,
    _paragraph_num_id_and_ilvl,
    _narrow_bookmark,
)

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_step_clarification_abstract_num_id(template_path):
    """
    Reads template_path's numbering.xml/document.xml/styles.xml directly
    to find the abstractNumId backing the 'Dilon Step Clarification
    List' style - the lowerLetter list ordered clarifications inside
    @@@STEPS@@@ get relettered onto, one fresh instance per step. Same
    two-place lookup as the retired get_step_list_abstract_num_id():
    a sample paragraph using the style with a numPr override, or the
    style's own baked-in numPr.

    Returns the abstractNumId as a string, or None (with a printed
    warning) if the template doesn't have the style set up - callers
    treat None as "leave ordered clarifications alone for this
    compile" rather than failing the whole document.
    """
    with zipfile.ZipFile(template_path) as z:
        doc_xml = etree.fromstring(z.read('word/document.xml'))
        styles_xml = etree.fromstring(z.read('word/styles.xml'))
        if 'word/numbering.xml' not in z.namelist():
            print("  Warning: template has no numbering.xml; step clarification numbering skipped")
            return None
        num_xml = etree.fromstring(z.read('word/numbering.xml'))

    sample_num_id = None
    for p in doc_xml.iter(f'{{{W_NS}}}p'):
        p_style = p.find(f'.//{{{W_NS}}}pStyle')
        if p_style is None or p_style.get(f'{{{W_NS}}}val') != 'DilonStepClarificationList':
            continue
        num_id_el = p.find(f'.//{{{W_NS}}}numPr/{{{W_NS}}}numId')
        if num_id_el is not None:
            sample_num_id = num_id_el.get(f'{{{W_NS}}}val')
            break

    if sample_num_id is None:
        for s in styles_xml.iter(f'{{{W_NS}}}style'):
            if s.get(f'{{{W_NS}}}styleId') != 'DilonStepClarificationList':
                continue
            num_id_el = s.find(f'.//{{{W_NS}}}pPr/{{{W_NS}}}numPr/{{{W_NS}}}numId')
            if num_id_el is not None:
                sample_num_id = num_id_el.get(f'{{{W_NS}}}val')
            break

    if sample_num_id is None:
        print(
            "  Warning: no paragraph in the base template uses the 'Dilon "
            "Step Clarification List' style with numbering applied; "
            "step clarification numbering skipped. See "
            "docs/superpowers/specs/2026-08-24-steps-field-numbering-design.md"
        )
        return None

    for num in num_xml.iter(f'{{{W_NS}}}num'):
        if num.get(f'{{{W_NS}}}numId') == sample_num_id:
            return num.find(f'{{{W_NS}}}abstractNumId').get(f'{{{W_NS}}}val')

    print(f"  Warning: numId {sample_num_id} has no matching <w:num> entry; step clarification numbering skipped")
    return None


def create_num_instance(numbering_element, abstract_num_id):
    """
    Appends a new <w:num> instance to numbering_element (a docx's
    <w:numbering> root, e.g. doc.part.numbering_part.element) referencing
    abstract_num_id, with a numId guaranteed not to collide with any numId
    already present (Pandoc allocates its own numIds for ordinary
    markdown lists in the same document - this must never reuse one of
    those). Returns the new numId as an int.

    Also writes a <w:lvlOverride>/<w:startOverride val="1"> for every
    level 0-8, forcing this instance to start counting at 1 regardless of
    what else in the document already used abstract_num_id. Without this,
    Word continues a level's counter across separate numId instances that
    share the same abstractNumId (confirmed against the real template:
    the throwaway 'Dilon Step List' sample paragraph's own numId bled its
    count into a freshly-allocated numId sharing the same abstract list,
    rendering 2/3/4 instead of 1/2/3) - every NEW step-list sequence must
    start fresh regardless of prior usage elsewhere in the document.
    """
    existing_ids = [int(n.get(qn('w:numId'))) for n in numbering_element.findall(qn('w:num'))]
    new_num_id = max(existing_ids, default=0) + 1

    num_el = OxmlElement('w:num')
    num_el.set(qn('w:numId'), str(new_num_id))
    abstract_ref_el = OxmlElement('w:abstractNumId')
    abstract_ref_el.set(qn('w:val'), str(abstract_num_id))
    num_el.append(abstract_ref_el)

    for ilvl in range(9):
        lvl_override = OxmlElement('w:lvlOverride')
        lvl_override.set(qn('w:ilvl'), str(ilvl))
        start_override = OxmlElement('w:startOverride')
        start_override.set(qn('w:val'), '1')
        lvl_override.append(start_override)
        num_el.append(lvl_override)

    numbering_element.append(num_el)

    return new_num_id


class StepBlockError(ValueError):
    """Raised for a malformed @@@STEPS@@@/@@@END_STEPS@@@ pairing - an
    @@@STEPS@@@ with no matching @@@END_STEPS@@@ before the next
    @@@STEPS@@@ or the end of the document, or an @@@END_STEPS@@@ with
    no @@@STEPS@@@ open. Unlike the old markdown-level version of this
    error, this is now raised post-conversion (see
    apply_section_scoped_step_numbering()) and halts compilation - a
    malformed steps block is an authoring mistake worth surfacing
    clearly, not silently degrading."""


def ensure_blank_line_around_steps_markers(markdown_text):
    """
    Ensures a blank line separates @@@STEPS@@@ from the list that
    follows it, and separates @@@END_STEPS@@@ from the list item that
    precedes it. Both markers now pass through to Pandoc as literal
    paragraphs (no more markdown-level block extraction -
    apply_section_scoped_step_numbering() finds them post-conversion),
    so a marker directly adjacent to a list item with no blank line
    risks CommonMark treating it as a lazy-continuation of that list
    item's own paragraph text instead of a separate paragraph -
    silently corrupting both the step text and the marker detection.

    Implemented as a line-by-line scan rather than the single-regex
    style of ensure_blank_line_after_table_markers()/
    ensure_blank_line_after_list_continue_markers(), because this is
    the one marker in this codebase that needs a blank line inserted
    *before* it (for @@@END_STEPS@@@), not just after - a single
    regex substitution reads and inserts in one direction only.
    """
    lines = markdown_text.split('\n')
    result = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        is_open = stripped == '@@@STEPS@@@'
        is_close = stripped == '@@@END_STEPS@@@'

        if is_close and result and result[-1].strip() != '':
            result.append('')

        result.append(line)

        if is_open and i + 1 < len(lines) and lines[i + 1].strip() != '':
            result.append('')

    return '\n'.join(result)


def _strip_num_pr(p_element):
    """Removes <w:numPr> from p_element's <w:pPr>, if present - used to
    turn a step's paragraph from a list item into plain text before
    restyling it 'Dilon Step Heading'."""
    p_pr = p_element.find(qn('w:pPr'))
    if p_pr is None:
        return
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is not None:
        p_pr.remove(num_pr)


def _prepend_step_number_fields(para):
    """
    Builds 'Step N.M' fields (STYLEREF 3 \\s + '.' + a SEQ counter that
    resets at each Heading 3) and inserts them BEFORE para's existing
    text runs, so the number reads first and the author's own step text
    follows - unlike add_field_simple_run(), which always appends to
    the paragraph's end. Mirrors apply_figure_captions()'s field-code
    syntax, just reset at Heading 3 (the step's own numbering scope)
    instead of Heading 2.

    Returns (start_el, end_el): the field span's first and last child
    elements, so the caller can narrow a {#step:label} bookmark around
    exactly this span.
    """
    anchor = para._p.find(qn('w:r'))

    def _place(new_el):
        if anchor is not None:
            anchor.addprevious(new_el)
        else:
            para._p.append(new_el)

    add_field_simple_run(para, ' STYLEREF 3 \\s ', '1')
    start_el = para._p[-1]
    _place(start_el)

    dot_run = para.add_run('.')
    _place(dot_run._element)

    add_field_simple_run(para, ' SEQ DilonStep \\* ARABIC \\s 3 ', '1')
    end_el = para._p[-1]
    _place(end_el)

    tab_run = para.add_run()
    tab_run.add_tab()
    _place(tab_run._element)

    return start_el, end_el


def _find_step_bookmark_start_in(para_element):
    """Returns the first <w:bookmarkStart> inside para_element whose
    name starts with 'step:' (the author's []{#step:label} anchor,
    which Pandoc renders as a zero-width bookmark somewhere within the
    paragraph - not necessarily at the start), or None."""
    for el in para_element.iter(qn('w:bookmarkStart')):
        name = el.get(qn('w:name'))
        if name and name.startswith('step:'):
            return el
    return None


def apply_field_based_step_numbering(docx_file, clarification_abstract_num_id):
    """
    Walks docx_file's paragraphs in document order, tracking whether a
    @@@STEPS@@@ block is open and which Heading 3 is currently in
    scope. Every ilvl-0 #.-list paragraph inside an open block becomes
    a field-numbered 'Dilon Step Heading' paragraph (see
    _prepend_step_number_fields()); every ilvl>=1 #.-list paragraph
    (an ordered "clarification") gets relettered onto a fresh numId of
    the 'Dilon Step Clarification List' abstract list, one fresh
    instance per top-level step. A bullet-list paragraph inside a block
    is left completely alone.

    Raises StepBlockError for an @@@STEPS@@@ with no matching
    @@@END_STEPS@@@, an @@@END_STEPS@@@ with no @@@STEPS@@@ open, or a
    block left open across a Heading 3 boundary - all compilation-
    halting, matching the retired apply_section_scoped_step_numbering()'s
    conventions.
    """
    from docx import Document
    doc = Document(docx_file)
    numbering_part = doc.part.numbering_part if clarification_abstract_num_id is not None else None
    numbering_element = numbering_part.element if numbering_part is not None else None

    decimal_ids = _decimal_abstract_num_ids(numbering_element) if numbering_element is not None else set()
    num_id_to_abstract = _num_id_to_abstract_map(numbering_element) if numbering_element is not None else {}

    heading_style_available = 'Dilon Step Heading' in {s.name for s in doc.styles}
    if not heading_style_available:
        print("  Warning: template has no 'Dilon Step Heading' style; step numbering skipped")

    inside_steps = False
    marker_elements = []
    current_clarification_num_id = None
    numbered = 0

    for para in doc.paragraphs:
        stripped = para.text.strip()

        if para.style is not None and para.style.name and para.style.name.startswith('Heading 3'):
            if inside_steps:
                raise StepBlockError("@@@STEPS@@@ has no matching @@@END_STEPS@@@ before the next section heading")
            continue

        if stripped == '@@@STEPS@@@':
            if inside_steps:
                raise StepBlockError("@@@STEPS@@@ opened again before the previous block's @@@END_STEPS@@@")
            inside_steps = True
            marker_elements.append(para._p)
            continue

        if stripped == '@@@END_STEPS@@@':
            if not inside_steps:
                raise StepBlockError("@@@END_STEPS@@@ found with no matching @@@STEPS@@@ open")
            inside_steps = False
            marker_elements.append(para._p)
            continue

        if not inside_steps:
            continue

        num_id, ilvl = _paragraph_num_id_and_ilvl(para._p)
        if num_id is None or num_id_to_abstract.get(num_id) not in decimal_ids:
            continue  # bullet or non-list paragraph inside a block - left alone

        if ilvl in (None, '0'):
            current_clarification_num_id = None
            if not heading_style_available:
                continue
            bookmark_start_el = _find_step_bookmark_start_in(para._p)
            _strip_num_pr(para._p)
            para.style = doc.styles['Dilon Step Heading']
            start_el, end_el = _prepend_step_number_fields(para)
            if bookmark_start_el is not None:
                _narrow_bookmark(doc.element.body, bookmark_start_el, start_el, end_el)
            numbered += 1
        else:
            if clarification_abstract_num_id is None:
                continue
            if current_clarification_num_id is None:
                current_clarification_num_id = create_num_instance(numbering_element, clarification_abstract_num_id)
            para.style = doc.styles['Dilon Step Clarification List']
            num_pr = para._p.find(qn('w:pPr')).find(qn('w:numPr'))
            num_pr.find(qn('w:numId')).set(qn('w:val'), str(current_clarification_num_id))
            ilvl_el = num_pr.find(qn('w:ilvl'))
            if ilvl_el is None:
                ilvl_el = OxmlElement('w:ilvl')
                num_pr.insert(0, ilvl_el)
            ilvl_el.set(qn('w:val'), '0')
            numbered += 1

    if inside_steps:
        raise StepBlockError("@@@STEPS@@@ has no matching @@@END_STEPS@@@ before the end of the document")

    for p_el in marker_elements:
        p_el.getparent().remove(p_el)

    if marker_elements or numbered:
        doc.save(docx_file)
        if numbered:
            print(f"  Applied field-based step numbering to {numbered} paragraph(s)")
    return numbered


def resolve_step_reference(para, bookmark_name):
    """
    type_resolvers['step'] callback for
    dilon_docx_common.resolve_reference_markers(): literal "Step " +
    a plain REF against the step's own bookmark (Task 3's
    apply_field_based_step_numbering() already narrowed that bookmark
    down to wrap just the number-field span, so \\h alone reproduces
    the live "2.3.1" text - no \\r needed, since there's no native list
    marker to extract anymore). The in-place number itself carries no
    "Step " word (matches the reviewer's reference-document
    convention), so it's added here, at the reference site, instead.
    """
    para.add_run('Step ')
    add_complex_field(para, f'REF {bookmark_name} \\h', '1')


