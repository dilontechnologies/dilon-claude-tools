# -*- coding: utf-8 -*-
"""
Generate Document from Markdown in Dilon formatting.

This script:
1. Parses a Markdown file with YAML front matter
2. Builds the signature page (header/footer/signature table), revision
   table, and title page directly via python-docx from the front-matter
   dict - no Jinja/docxtpl involved for any of these parts
3. Converts the Markdown body to Rich Text via Pandoc, with body-level
   Jinja2 {{field}} substitution against the same front-matter dict
4. Composes all parts into the final Word document

Usage:
    python generate_dilon_doc.py <input.md> <output.docx>

Example:
    python generate_dilon_doc.py MAP-00001_Requirements.md MAP-00001_Requirements.docx

The Pandoc-conversion, Word-styling, and header/footer-building helpers
this script calls (apply_styles, apply_figure_captions, markdown_to_docx,
extract_yaml_and_markdown, set_update_fields_on_open, compose_documents,
populate_header, populate_footer) live in lib/dilon_docx_common.py (repo
root) - shared with dilon-document-form-compiler. See
docs/superpowers/specs/2026-08-17-document-extraction-and-form-tooling-design.md.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# Set UTF-8 encoding for Windows console
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

_REPO_ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(_REPO_ROOT / "lib"))
from dilon_docx_common import (  # noqa: E402
    apply_styles,
    apply_figure_captions,
    center_image_paragraphs,
    markdown_to_docx,
    extract_yaml_and_markdown,
    set_update_fields_on_open,
    compose_documents,
    ensure_blank_line_after_table_markers,
    ensure_blank_line_between_images,
    ensure_blank_line_after_list_continue_markers,
    remap_ordered_lists_to_dilon_step_list,
    resolve_list_continuations,
    validate_list_nesting_depth,
    ListContinuationError,
    ListNestingError,
    parse_column_widths,
    apply_table_column_widths,
    render_jinja,
    populate_header,
    populate_footer,
    strip_leading_empty_paragraphs,
)

sys.path.insert(0, str(Path(__file__).resolve().parent))
from step_numbering import (  # noqa: E402
    preprocess_step_references,
    ensure_blank_line_around_steps_markers,
    get_step_list_abstract_num_id,
    apply_section_scoped_step_numbering,
    resolve_step_references,
    StepBlockError,
    StepReferenceError,
)


def create_revision_table(revisions, available_width):
    """
    Create a formatted revision history table as a Word table object.

    Args:
        revisions: List of revision dictionaries with keys: number, description, eco_number, eco_date
        available_width: content width (page width minus margins), as a
            python-docx Length, to size the table to - see
            generate_requirements_document().

    Returns:
        python-docx Table object
    """
    from docx.shared import Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Create a temporary document to build the table
    temp_doc = Document()

    # Create table with title + header + data rows
    table = temp_doc.add_table(rows=2 + len(revisions), cols=4)
    table.style = 'Table Grid'

    # REV #/ECO #/DATE widths match an approved, hand-tuned Dilon Word
    # document's own revision table; DESCRIPTION absorbs whatever width
    # is left so the table always fills the page's full content width.
    # apply_table_column_widths() (not a plain table.columns[idx].width=
    # assignment, which only sets the tblGrid definition, not per-cell
    # width or autofit=False - Word's AutoFit then silently recalculates
    # the visible columns from cell content and ignores it) sets both, so
    # the widths actually render as specified.
    apply_table_column_widths(
        table,
        [Twips(805).inches, 'x', Twips(1620).inches, Twips(1535).inches],
        available_width.inches,
    )

    # Title row (row 0) - spans all columns
    title_cell = table.rows[0].cells[0]
    # Merge all cells in first row
    for i in range(1, 4):
        title_cell.merge(table.rows[0].cells[i])
    title_cell.text = "REVISION HISTORY"

    # Format title: bold, centered, gray background
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.bold = True
    title_run.font.size = Pt(11)

    # Set gray background for title
    title_shading = OxmlElement('w:shd')
    title_shading.set(qn('w:fill'), 'C0C0C0')
    title_cell._element.get_or_add_tcPr().append(title_shading)

    # Header row (row 1)
    header_cells = table.rows[1].cells
    headers = ['REV #', 'DESCRIPTION OF CHANGE', 'ECO #', 'DATE']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text

        # Format header: bold, centered, gray background
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(9)

        # Set gray background
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C0C0C0')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Data rows (starting at row 2)
    for idx, rev in enumerate(revisions):
        row_cells = table.rows[idx + 2].cells

        # REV #
        row_cells[0].text = rev.get('number', '')
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # DESCRIPTION
        row_cells[1].text = rev.get('description', '')

        # ECO #
        row_cells[2].text = rev.get('eco_number', '')
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # DATE
        row_cells[3].text = rev.get('eco_date', '')
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Center the table itself using XML
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)

    # Add table justification element to center the table
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tbl_pr.append(jc)

    return table


def create_signature_table(metadata, available_width):
    """
    Create a formatted signature-approval table as a Word table object.

    Previously baked directly into TEMPLATE_Word_Base.docx as a
    docxtpl-rendered table; built programmatically here instead so the
    template itself can be reduced to header/footer + styles only
    (matching dilon-document-form-compiler's TEMPLATE_Word_Form.docx).

    Args:
        metadata: front-matter dict with department, author,
            regulatory_rep, quality_rep, department_head
        available_width: content width (page width minus margins), as a
            python-docx Length, to size the table to - see
            generate_requirements_document().

    Returns:
        python-docx Table object
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Twips

    temp_doc = Document()
    table = temp_doc.add_table(rows=6, cols=3)
    table.style = 'Normal Table'

    # Group/Department and Signature widths match an approved, hand-tuned
    # Dilon Word document's own signature table; Preparer/Name absorbs
    # whatever width is left so the table always fills the page's full
    # content width. apply_table_column_widths() (not a plain
    # table.columns[idx].width= assignment, which only sets the tblGrid
    # definition, not per-cell width or autofit=False - Word's AutoFit
    # then silently recalculates the visible columns from cell content
    # and ignores it) sets both, so the widths actually render as
    # specified.
    apply_table_column_widths(
        table,
        [Twips(2330).inches, 'x', Twips(2440).inches],
        available_width.inches,
    )

    def set_cell(row_idx, col_idx, text, bold=False, center=True, fill=None):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = text
        paragraph = cell.paragraphs[0]
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.runs:
            paragraph.runs[0].font.bold = bold
        if fill:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), fill)
            cell._element.get_or_add_tcPr().append(shading)

    department = metadata.get('department', '')

    set_cell(0, 0, 'Group', bold=True, center=True, fill='BFBFBF')
    set_cell(0, 1, 'Preparer', bold=True, center=False, fill='BFBFBF')
    set_cell(0, 2, 'Signature', bold=True, center=True, fill='BFBFBF')

    set_cell(1, 0, department, center=True, fill='FFFFFF')
    set_cell(1, 1, metadata.get('author', ''), center=False, fill='FFFFFF')
    set_cell(1, 2, 'Electronic', center=True)

    set_cell(2, 0, 'Department', bold=True, center=True, fill='C0C0C0')
    set_cell(2, 1, 'Name', bold=True, center=False, fill='C0C0C0')
    set_cell(2, 2, 'Signature', bold=True, center=True, fill='C0C0C0')

    set_cell(3, 0, 'Regulatory', center=True)
    set_cell(3, 1, metadata.get('regulatory_rep', ''), center=False)
    set_cell(3, 2, 'Electronic', center=True)

    set_cell(4, 0, 'Quality', center=True)
    set_cell(4, 1, metadata.get('quality_rep', ''), center=False)
    set_cell(4, 2, 'Electronic', center=True)

    set_cell(5, 0, department, center=True)
    set_cell(5, 1, metadata.get('department_head', ''), center=False)
    set_cell(5, 2, 'Electronic', center=True)

    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Explicit grid borders (all edges + inside lines) - 'Normal Table'
    # alone carries none. Matches the DilonTable_List/DilonTable_Chart
    # custom styles' own border spec (single, sz=8 = 1pt, auto color) for
    # visual consistency with every other table in a compiled document.
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '8')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), 'auto')
        borders.append(edge_el)
    tbl_pr.append(borders)

    return table


def _insert_table_before_section_properties(document, table):
    """
    Insert `table` as the last body content of `document`, immediately
    before its sectPr - the correct position for content that must belong
    to the document's existing (only) section, as opposed to
    Document.body.append() which would place it after sectPr and produce
    an invalid section boundary.

    Also inserts a blank separator paragraph directly after the table
    (still before sectPr), so the table never ends up directly adjacent to
    whatever content compose_documents() splices in right after Part A
    (e.g. Part B's revision table) - Word silently merges two directly
    adjacent <w:tbl> elements into one visual table on open, the same
    hazard apply_styles()'s between-two-tables handling in
    lib/dilon_docx_common.py guards against for marker-adjacent tables.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body = document._element.body
    sect_pr = body.find(qn('w:sectPr'))
    separator = OxmlElement('w:p')
    if sect_pr is not None:
        sect_pr.addprevious(table._element)
        sect_pr.addprevious(separator)
    else:
        body.append(table._element)
        body.append(separator)


def _add_runs(paragraph, spans):
    """Append (text, bold, italic) tuples to `paragraph` as separate runs -
    used for title-page boilerplate that bolds/italicizes specific words
    mid-sentence."""
    for text, bold, italic in spans:
        run = paragraph.add_run(text)
        run.font.bold = bold
        run.font.italic = italic


def build_title_page(document, metadata):
    """
    Build the title page (title, author table, master-document/
    effectivity/approval boilerplate) directly via python-docx, in place
    of a docxtpl-rendered TEMPLATE_Word_Content.docx. Only {{title}} and
    {{author}} were ever dynamic in that template - everything else was
    static boilerplate, now inlined here as plain Python.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    title_para = document.add_paragraph(metadata.get('title', ''), style='Title')

    table = document.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(1.7326388888888888)
    table.columns[1].width = Inches(4.814583333333333)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    label_cell, value_cell = table.rows[0].cells
    label_run = label_cell.paragraphs[0].add_run('Author/Revised by:')
    label_run.font.bold = True
    value_cell.paragraphs[0].add_run(metadata.get('author', ''))

    # Explicit grid borders - 'Normal Table' alone carries none, which
    # made this table invisible (indistinguishable from plain side-by-side
    # text). Same border spec as create_signature_table()'s fix.
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '8')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), 'auto')
        borders.append(edge_el)
    tbl_pr.append(borders)

    for _ in range(4):
        document.add_paragraph()

    document.add_paragraph('Master Document', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    effectivity_heading = document.add_paragraph(style='Normal')
    effectivity_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    effectivity_run = effectivity_heading.add_run('Effectivity and Location:')
    effectivity_run.font.underline = True

    effectivity_body = document.add_paragraph(style='Normal')
    effectivity_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(effectivity_body, [
        ('This is an electronic document, the authoritative ', False, False),
        ('Item', True, False),
        (' ', False, True),
        ('in the ', False, False),
        ('Dilon Technologies', True, False),
        (' Production Workspace in the ', False, False),
        ('ARENA PLM system.', True, False),
    ])

    non_authoritative = document.add_paragraph(style='Normal')
    non_authoritative.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(non_authoritative, [
        ('\nAll other copies of this document, either in electronic or physical media, shall be considered as ', False, False),
        ('non-authoritative', True, False),
        (' copies.', False, False),
    ])

    document.add_paragraph(style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    approval_heading = document.add_paragraph(style='Normal')
    approval_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    approval_run = approval_heading.add_run('Approval, Release and Change History:')
    approval_run.font.underline = True

    approval_body = document.add_paragraph(style='Normal')
    approval_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(approval_body, [
        ('Records of the approvals and release of this document version and its full revision history are available in the History ', False, False),
        ('Subview', False, False),
        (' of the Revisions View of the related ', False, False),
        ('Item', True, False),
        (' in the ', False, False),
        ('Dilon Technologies Production Workspace', True, False),
        (' in the ', False, False),
        ('ARENA PLM system', True, False),
    ])

    for _ in range(3):
        document.add_paragraph()

    return title_para


def generate_requirements_document(markdown_path, output_path, signature_template_path=None):
    """
    Generate final requirements Word document.

    Args:
        markdown_path: Path to Markdown file with YAML front matter
        output_path: Path to save final Word document
        signature_template_path: Path to the base template (Part A) -
            header/footer/styles only; no Jinja fields, no docxtpl
            involved anymore. See populate_header()/populate_footer() in
            lib/dilon_docx_common.py.
    """
    if signature_template_path is None:
        signature_template_path = _REPO_ROOT / "templates" / "TEMPLATE_Word_Base.docx"
    else:
        signature_template_path = Path(signature_template_path)

    if not signature_template_path.exists():
        print(f"Error: Base template not found: {signature_template_path}")
        sys.exit(1)

    if not Path(markdown_path).exists():
        print(f"Error: Markdown file not found: {markdown_path}")
        sys.exit(1)

    print(f"Reading Markdown file: {markdown_path}")

    # Extract YAML metadata and Markdown body
    metadata, markdown_body = extract_yaml_and_markdown(markdown_path)
    markdown_body = render_jinja(markdown_body, metadata)
    markdown_body = preprocess_step_references(markdown_body)
    markdown_body = ensure_blank_line_around_steps_markers(markdown_body)
    markdown_body = ensure_blank_line_after_list_continue_markers(markdown_body)

    print(f"Metadata extracted: {list(metadata.keys())}")

    # Step 1: Build Part A (base template: header, footer, signature table)
    print(f"Building signature page (Part A) from: {signature_template_path}")
    doc_a = Document(signature_template_path)
    populate_header(doc_a, metadata)
    populate_footer(doc_a, metadata)
    strip_leading_empty_paragraphs(doc_a)

    section = doc_a.sections[0]
    available_width = Emu(section.page_width - section.left_margin - section.right_margin)

    # Build the signature-approval table programmatically (same pattern as
    # Part B's revision table) and insert it into Part A itself.
    signature_table = create_signature_table(metadata, available_width)
    _insert_table_before_section_properties(doc_a, signature_table)

    temp_part_a = Path(output_path).parent / "_temp_part_a.docx"
    doc_a.save(temp_part_a)
    print(f"Part A built")

    # Step 2: Generate Part B (Revision table)
    temp_part_b = Path(output_path).parent / "_temp_part_b.docx"
    if 'revisions' in metadata and metadata['revisions']:
        print("Building revision history table (Part B)...")
        revision_doc = Document()

        # Create the table
        table = create_revision_table(metadata['revisions'], available_width)

        # Center the table before adding to document
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add table to document
        revision_doc._element.body.append(table._element)

        revision_doc.save(temp_part_b)
        print(f"Part B generated")
    else:
        # Create empty document if no revisions
        revision_doc = Document()
        revision_doc.add_paragraph("No revision history available")
        revision_doc.save(temp_part_b)
        print(f"No revisions found, created placeholder")

    # Step 3: Build Part C (title page)
    print("Building title page (Part C)...")
    doc_c = Document()
    build_title_page(doc_c, metadata)
    temp_part_c = Path(output_path).parent / "_temp_part_c.docx"
    doc_c.save(temp_part_c)
    print(f"Part C built")

    # Step 4: Convert Markdown body to Word (Part D) using signature template as style reference
    print("Converting Markdown content to Word (Part D)...")
    temp_part_d = Path(output_path).parent / "_temp_part_d.docx"

    # Use signature template as reference to ensure consistent formatting across entire document
    markdown_to_docx(
        markdown_body, temp_part_d,
        reference_doc=signature_template_path,
        resource_dir=Path(markdown_path).resolve().parent,
    )

    # Apply all styles (tables and paragraphs) - scans Word document for @@@ markers, applies styles, removes markers
    print("Applying custom styles...")
    apply_styles(temp_part_d)

    print("Applying ordered-list styling...")
    remap_ordered_lists_to_dilon_step_list(temp_part_d)
    try:
        resolve_list_continuations(temp_part_d)
    except ListContinuationError as exc:
        print(f"Error: {exc}")
        sys.exit(1)
    try:
        validate_list_nesting_depth(temp_part_d)
    except ListNestingError as exc:
        print(f"Error: {exc}")
        sys.exit(1)

    # Convert Pandoc's implicit-figure captions into auto-numbered Word captions
    print("Applying figure caption numbering...")
    apply_figure_captions(temp_part_d)

    print("Centering image paragraphs...")
    center_image_paragraphs(temp_part_d)

    try:
        step_list_abstract_num_id = get_step_list_abstract_num_id(signature_template_path)
        print("Applying step-list numbering...")
        apply_section_scoped_step_numbering(temp_part_d, step_list_abstract_num_id)
        resolve_step_references(temp_part_d)
    except (StepBlockError, StepReferenceError) as exc:
        print(f"Error: {exc}")
        sys.exit(1)

    print(f"Part D converted")

    # Step 5: Merge all documents in order: A -> B -> C -> D
    print("Merging all parts (A -> B -> C -> D)...")
    composer = compose_documents(temp_part_a, temp_part_b, temp_part_c, temp_part_d)
    composer.save(output_path)

    # Ensure figure numbers / TOC page numbers are correct the moment the
    # document is opened, rather than showing cached placeholder text
    set_update_fields_on_open(output_path)

    # Clean up temporary files
    temp_part_a.unlink()
    temp_part_b.unlink()
    temp_part_c.unlink()
    temp_part_d.unlink()

    print(f"\nDocument generated successfully!")
    print(f"Output: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_requirements_doc.py <input.md> <output.docx> [base_template.docx]")
        print("\nExample:")
        print("  python generate_requirements_doc.py MAP-00001_Requirements.md MAP-00001_Requirements.docx")
        print("  python generate_requirements_doc.py MAP-00001_Requirements.md MAP-00001_Requirements.docx custom_base.docx")
        sys.exit(1)

    markdown_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    signature_template_path = Path(sys.argv[3]) if len(sys.argv) > 3 else None

    try:
        generate_requirements_document(markdown_path, output_path, signature_template_path)
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
