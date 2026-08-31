"""Extract a Dilon-style .docx into a best-effort Dilon markdown draft.

See docs/superpowers/specs/2026-08-17-document-extraction-and-form-tooling-design.md
(dilon-claude-tools repo) for the design this implements.
"""

import re
import sys
from pathlib import Path

import yaml

WORD_HEADING_RE = re.compile(r'^Heading (\d)$')
SUSPICIOUS_WORD_COUNT_THRESHOLD = 12


def word_heading_level(style_name):
    """Return the numeric Word heading level (1-9) for a paragraph style
    name like 'Heading 2', or None if style_name isn't a heading style."""
    if not style_name:
        return None
    m = WORD_HEADING_RE.match(style_name)
    return int(m.group(1)) if m else None


def compute_heading_shift(doc):
    """Scan every paragraph in doc and return the shift to add to a Word
    heading level to get its markdown '#' count, such that the shallowest
    heading level present maps to 2 (markdown '##', matching the
    compiler's H2-is-top convention). Returns 2 (no-op shift for a
    document with no headings, since level + 2 would be wrong - callers
    only use the shift when a heading is actually present)."""
    levels = [
        word_heading_level(p.style.name if p.style else None)
        for p in doc.paragraphs
    ]
    levels = [lvl for lvl in levels if lvl is not None]
    if not levels:
        return 2
    return 2 - min(levels)


def markdown_heading_prefix(word_level, shift):
    """Return the '#'*N markdown heading prefix for a Word heading level,
    given the shift computed by compute_heading_shift(). Clamped to
    markdown's practical range (## through ######)."""
    n = max(2, min(word_level + shift, 6))
    return '#' * n


def is_suspicious_heading_text(text):
    """True if a heading-styled paragraph reads like body text rather than
    a real heading (ends in a period, or unusually long) - a strong signal
    it's actually a numbered procedure step (MARKDOWN_STYLING_GUIDE.md SS6)
    that the source document authored with Word heading styles instead of
    a real list. Flagged for human review, never silently rewritten."""
    text = text.strip()
    if not text:
        return False
    if text.endswith('.'):
        return True
    return len(text.split()) > SUSPICIOUS_WORD_COUNT_THRESHOLD


TITLECASE_WORD_RE = re.compile(r"[A-Za-z']+")


def titlecase_heading(text):
    """Title-case heading text: capitalize the first letter of every word
    and lowercase the rest, regardless of the source document's own
    capitalization (real Dilon headings are inconsistently ALL CAPS /
    Sentence case / Mixed case)."""
    return TITLECASE_WORD_RE.sub(lambda m: m.group(0)[:1].upper() + m.group(0)[1:].lower(), text)


TOC_STYLE_RE = re.compile(r'(?i)^toc\s*\d+$')


def is_toc_paragraph(style_name, text):
    """True for a Word-generated table-of-contents entry ('toc 1'/'toc 2'/
    etc. paragraph styles, or the 'TOC Heading' style) or the literal
    'TABLE OF CONTENTS' heading text that precedes them. The compiler
    regenerates its own TOC from headings, so the source document's baked-
    in TOC would just be redundant, stale noise in the body."""
    if style_name and (TOC_STYLE_RE.match(style_name) or style_name.strip().lower() == 'toc heading'):
        return True
    return text.strip().lower() == 'table of contents'


def paragraph_is_list_item(paragraph):
    """True if paragraph is part of a numbered/bulleted list - either via
    the 'List Paragraph' style, or via direct w:numPr list formatting on a
    paragraph using a different style. Real Dilon documents mix both: a
    single-item "list" is sometimes left styled 'Normal' with manual list
    formatting rather than switched to the 'List Paragraph' style."""
    from docx.oxml.ns import qn
    style_name = paragraph.style.name if paragraph.style else None
    if style_name == "List Paragraph":
        return True
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


def paragraph_list_ilvl(paragraph):
    """Return the 0-based numbering indent level for a list-item
    paragraph's w:numPr formatting, or 0 if absent/unspecified. Used to
    preserve nested-list structure (e.g. Word's ilvl-1 sub-bullets) as
    2-space-indented markdown nesting (MARKDOWN_STYLING_GUIDE.md SS5.1)."""
    from docx.oxml.ns import qn
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is None:
        return 0
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return 0
    ilvl_el = numPr.find(qn('w:ilvl'))
    if ilvl_el is None:
        return 0
    try:
        return int(ilvl_el.get(qn('w:val')))
    except (TypeError, ValueError):
        return 0


def classify_table(table):
    """Return 'signature', 'revision', or 'content' for a python-docx
    Table, based on its first two rows' text (case-insensitive)."""
    header_text = " ".join(
        cell.text.strip().lower()
        for row in table.rows[:2]
        for cell in row.cells
    )
    if "revision history" in header_text or ("rev #" in header_text and "eco #" in header_text):
        return "revision"
    if "signature" in header_text and ("preparer" in header_text or "name" in header_text):
        return "signature"
    return "content"


def extract_signature_fields(table):
    """table: a Table classified as 'signature'. Returns (fields, warnings).
    fields may include 'author', 'department', 'department_head', and
    'signature_fields' (a list of {'department', 'name'} dicts for any
    rows beyond the fixed department_head row). Matches
    create_signature_table()'s canonical shape by position: row 1 is
    department/author, row 3 is department_head, and every row from 4
    onward is a signature_fields entry - a warning (not a failure) is
    returned when the department_head row's label doesn't match the top
    department field, since real source documents (e.g. WI-00077) use
    inconsistent role labels ("R&D / Eng", "Manufacturing") in these
    slots."""
    fields = {}
    warnings = []
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    if len(rows) < 4:
        warnings.append(
            f"signature table has {len(rows)} rows, expected at least 4 "
            "(canonical Group/Preparer/Signature + Department/Name/Signature "
            "+ Department Head shape) - extracted nothing, fill approvers in manually"
        )
        return fields, warnings

    fields["department"] = rows[1][0]
    fields["author"] = rows[1][1]
    fields["department_head"] = rows[3][1]

    department_head_label = rows[3][0]
    if department_head_label.strip().lower() != fields["department"].strip().lower():
        warnings.append(
            f"department head row labeled '{department_head_label}' doesn't "
            f"match the preparer's department '{fields['department']}' - "
            "verify by position"
        )

    fields["signature_fields"] = [
        {"department": row[0], "name": row[1]} for row in rows[4:] if len(row) >= 2
    ]

    return fields, warnings


def extract_revisions(table):
    """table: a Table classified as 'revision'. Returns a list of dicts
    with keys 'number', 'description', 'eco_number', 'eco_date', skipping
    a leading merged 'REVISION HISTORY' title row and the column-header
    row if present."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    idx = 0
    if rows and rows[0][0].strip().upper() == "REVISION HISTORY":
        idx = 1
    if idx < len(rows) and rows[idx][:2] == ["REV #", "DESCRIPTION OF CHANGE"]:
        idx += 1

    revisions = []
    for row in rows[idx:]:
        if len(row) < 4 or not row[0].strip():
            continue
        revisions.append({
            "number": row[0].strip(),
            "description": row[1].strip(),
            "eco_number": row[2].strip(),
            "eco_date": row[3].strip(),
        })
    return revisions


DOC_NUMBER_RE = re.compile(r'Number:\s*([A-Za-z]{2,}-\d+)')
# Revision values aren't digit-only: prototype revisions extend the
# format to e.g. "02-A" (major number + alphabetic prototype suffix), so
# the capture group accepts alphanumerics plus internal '.'/'-'.
REV_RE = re.compile(r'Rev\s+([A-Za-z0-9][A-Za-z0-9.\-]*)')
FOOTER_LINE_RE = re.compile(
    r'([A-Za-z]{2,}-\d+)\s+Rev\s+([A-Za-z0-9][A-Za-z0-9.\-]*)\s+(ECO-\d+)\s+Revision Date:\s*([\d/]+)'
)
FIGURE_PREFIX_RE = re.compile(r'^Figure\s+[\d.]+\s*[:\-]\s*', re.IGNORECASE)
HEADER_LABEL_VALUE_RE = re.compile(r'^([A-Za-z][A-Za-z \-]{0,20}):\s*(.*)$', re.DOTALL)
HEADER_TITLE_SKIP_LABELS = {'number', 'page'}
HEADER_TITLE_TRAILING_REV_RE = re.compile(r'\s*Rev\s+\d+.*$', re.DOTALL)


def extract_header_footer_metadata(doc):
    """Returns a dict that may include 'title', 'doc_number',
    'current_revision', 'footer_eco_number', 'footer_eco_date', parsed from
    the document's running header table and footer text - the most
    structured, reliable source for these fields (see the spec's empirical
    findings)."""
    fields = {}
    section = doc.sections[0]

    for table in section.header.tables:
        for row in table.rows:
            cell_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            joined = " ".join(cell_texts)
            m = DOC_NUMBER_RE.search(joined)
            if m:
                fields["doc_number"] = m.group(1)
            m = REV_RE.search(joined)
            if m:
                fields["current_revision"] = m.group(1)

            if cell_texts and "title" not in fields:
                label_match = HEADER_LABEL_VALUE_RE.match(cell_texts[0])
                if label_match and label_match.group(1).strip().lower() not in HEADER_TITLE_SKIP_LABELS:
                    value = label_match.group(2).strip()
                    if not value and len(cell_texts) > 1:
                        value = cell_texts[1]
                    value = HEADER_TITLE_TRAILING_REV_RE.sub('', value).strip()
                    value = re.sub(r'\s+', ' ', value)
                    if value:
                        fields["title"] = value

    # Two footer shapes exist across real documents: older/hand-authored
    # ones carry the ID line as tab-separated text in a single paragraph;
    # newer compiler output (see populate_footer() in
    # lib/dilon_docx_common.py) carries it as the first row of a 3-column
    # table instead - table cell text isn't part of footer.paragraphs, so
    # both are checked.
    footer_text_parts = [p.text for p in section.footer.paragraphs if p.text.strip()]
    for table in section.footer.tables:
        id_row_texts = [c.text.strip() for c in table.rows[0].cells if c.text.strip()]
        if id_row_texts:
            footer_text_parts.append(" ".join(id_row_texts))
    footer_text = "\n".join(footer_text_parts)
    m = FOOTER_LINE_RE.search(footer_text)
    if m:
        fields.setdefault("doc_number", m.group(1))
        fields.setdefault("current_revision", m.group(2))
        fields["footer_eco_number"] = m.group(3)
        fields["footer_eco_date"] = m.group(4)

    return fields


def strip_figure_prefix(text):
    """Strip a leading 'Figure N:'/'Figure N.M -' prefix from caption text,
    since the compiler generates that prefix itself from the image's alt
    text (MARKDOWN_STYLING_GUIDE.md SS4.1)."""
    return FIGURE_PREFIX_RE.sub('', text).strip()


def slugify(text, existing=None):
    """Lowercase, hyphenate text into a Pandoc-safe identifier fragment,
    deduplicated against `existing` (a set this function mutates) with a
    -2, -3, ... suffix on collision."""
    slug = re.sub(r'[^a-z0-9]+', '-', text.lower()).strip('-') or "figure"
    existing = existing if existing is not None else set()
    candidate = slug
    n = 2
    while candidate in existing:
        candidate = f"{slug}-{n}"
        n += 1
    existing.add(candidate)
    return candidate


def paragraph_image_rids(paragraph):
    """Return the r:embed relationship IDs of any inline images in a
    paragraph, in document order."""
    from docx.oxml.ns import qn
    return [
        blip.get(qn('r:embed'))
        for blip in paragraph._p.findall('.//' + qn('a:blip'))
        if blip.get(qn('r:embed'))
    ]


def save_image(doc, rid, images_dir, index):
    """Write the image identified by relationship id `rid` to
    images_dir/imageNN.<ext>, returning the filename (not full path)."""
    part = doc.part.related_parts[rid]
    ext = part.content_type.split('/')[-1].replace('jpeg', 'jpg')
    filename = f"image{index:02d}.{ext}"
    (Path(images_dir) / filename).write_bytes(part.blob)
    return filename


def iter_block_items(doc):
    """Yield each top-level child of the document body as a
    docx.text.paragraph.Paragraph or docx.table.Table, in document order -
    doc.paragraphs/doc.tables lose that relative order."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def _is_blank_block(block):
    """True if a block contributes no visible content to the extracted
    body - a text-less paragraph with no embedded image, or a table-of-
    contents artifact. Tables, images, captions, and heading paragraphs
    always count as real content."""
    from docx.table import Table
    if isinstance(block, Table):
        return False
    style_name = block.style.name if block.style else None
    text = block.text.strip()
    if is_toc_paragraph(style_name, text):
        return True
    if text:
        return False
    return not paragraph_image_rids(block)


def _nearest_real_block(blocks, index, step):
    """Scan from `index` in direction `step` (+1 or -1), skipping blank
    blocks (see _is_blank_block), and return the nearest non-blank block,
    or None if the scan runs off either end."""
    j = index + step
    while 0 <= j < len(blocks):
        if not _is_blank_block(blocks[j]):
            return blocks[j]
        j += step
    return None


def heading_is_empty_leaf(blocks, index, level):
    """True if the heading-styled paragraph at blocks[index] (Word heading
    `level`) is a content-less leaf being used to label a single fact
    rather than introduce a real subsection: it is the sole child of a
    parent heading exactly one level shallower (immediately before it,
    skipping blanks), AND has no body/list/table/image content of its own
    before the next heading of equal-or-shallower level or the end of the
    document. A heading followed by a *deeper* heading (a real parent) or
    by any real content is never an empty leaf."""
    from docx.text.paragraph import Paragraph

    prev = _nearest_real_block(blocks, index, -1)
    if not isinstance(prev, Paragraph):
        return False
    prev_level = word_heading_level(prev.style.name if prev.style else None)
    if prev_level != level - 1:
        return False

    nxt = _nearest_real_block(blocks, index, 1)
    if nxt is None:
        return True
    if not isinstance(nxt, Paragraph):
        return False
    nxt_level = word_heading_level(nxt.style.name if nxt.style else None)
    return nxt_level is not None and nxt_level <= level


def cell_has_multiple_paragraphs(cell):
    return len(cell.paragraphs) > 1 and any(p.text.strip() for p in cell.paragraphs[1:])


def table_to_grid_markdown(table):
    """Renders a table with any multi-paragraph cell as a Pandoc grid
    table (MARKDOWN_STYLING_GUIDE.md SS14). Each cell's paragraphs are
    collapsed to one line (space-joined) - faithful multi-line box
    rendering is out of scope for a mechanical draft; a human reviewing
    the extracted draft restores paragraph breaks where they matter."""
    rows = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
    ncols = max(len(row) for row in rows)
    rows = [row + [''] * (ncols - len(row)) for row in rows]
    widths = [max(len(rows[r][c]) for r in range(len(rows))) for c in range(ncols)]

    def border(char):
        return "+" + "+".join(char * (w + 2) for w in widths) + "+"

    def row_line(row):
        return "|" + "|".join(f" {cell.ljust(w)} " for cell, w in zip(row, widths)) + "|"

    lines = [border("-"), row_line(rows[0]), border("=")]
    for row in rows[1:]:
        lines.append(row_line(row))
        lines.append(border("-"))
    return "\n".join(lines)


def table_to_markdown(table):
    """Renders a 'content'-classified table as a pipe table, or a grid
    table if any cell has multiple paragraphs of text."""
    if any(cell_has_multiple_paragraphs(cell) for row in table.rows for cell in row.cells):
        return table_to_grid_markdown(table)

    rows = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
    ncols = max(len(row) for row in rows)
    rows = [row + [''] * (ncols - len(row)) for row in rows]
    lines = ["| " + " | ".join(rows[0]) + " |", "|" + "|".join(["---"] * ncols) + "|"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def build_markdown_body(doc, blocks, shift, images_dir, front_matter):
    """Walks `blocks` (from iter_block_items) in order, appending markdown
    lines and mutating `front_matter` in place with any signature/revision
    table data encountered. Returns (body_text, warnings)."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    warnings = []
    lines = []
    in_list = False
    last_list_ilvl = 0
    steps_open = False
    steps_base_level = None
    existing_slugs = set()
    image_index = 1

    def flush_list():
        nonlocal in_list
        if in_list:
            lines.append("")
            in_list = False

    def flush_steps():
        nonlocal steps_open, steps_base_level
        if steps_open:
            lines.append("")
            lines.append("@@@END_STEPS@@@")
            lines.append("")
            steps_open = False
            steps_base_level = None

    def flush_all():
        flush_list()
        flush_steps()

    for i, block in enumerate(blocks):
        if isinstance(block, Table):
            flush_all()
            kind = classify_table(block)
            if kind == "signature":
                fields, sig_warnings = extract_signature_fields(block)
                front_matter.update(fields)
                warnings.extend(sig_warnings)
            elif kind == "revision":
                front_matter["revisions"] = extract_revisions(block)
            else:
                lines.append(table_to_markdown(block))
                lines.append("")
            continue

        # block is a Paragraph
        style_name = block.style.name if block.style else None
        text = block.text.strip()

        if is_toc_paragraph(style_name, text):
            continue

        rids = paragraph_image_rids(block)
        if rids:
            flush_all()
            caption_text = None
            if i + 1 < len(blocks) and isinstance(blocks[i + 1], Paragraph):
                next_style = blocks[i + 1].style.name if blocks[i + 1].style else None
                if next_style == "Caption":
                    caption_text = strip_figure_prefix(blocks[i + 1].text.strip())
            for rid in rids:
                filename = save_image(doc, rid, images_dir, image_index)
                image_index += 1
                if caption_text:
                    slug = slugify(caption_text, existing_slugs)
                    lines.append(f"![{caption_text}](images/{filename}){{#fig:{slug}}}")
                else:
                    lines.append(f"![](images/{filename})")
            lines.append("")
            continue

        if style_name == "Caption":
            prev_had_image = (
                i > 0
                and isinstance(blocks[i - 1], Paragraph)
                and paragraph_image_rids(blocks[i - 1])
            )
            if not prev_had_image:
                flush_all()
                warnings.append(f"orphan Caption paragraph with no preceding image: {text!r}")
                lines.append(text)
                lines.append("")
            continue

        if not text:
            continue

        level = word_heading_level(style_name)
        if level is not None:
            suspicious = is_suspicious_heading_text(text)
            empty_leaf = not suspicious and heading_is_empty_leaf(blocks, i, level)
            if suspicious:
                warnings.append(
                    "heading-styled paragraph reads like a numbered procedure "
                    f"step, rendered as a @@@STEPS@@@ item: {text!r}"
                )
                flush_list()
                if steps_open and level < steps_base_level:
                    flush_steps()
                if not steps_open:
                    lines.append("@@@STEPS@@@")
                    lines.append("")
                    steps_open = True
                    steps_base_level = level
                nest = level - steps_base_level
                if nest > 2:
                    warnings.append(
                        "procedure step nested deeper than the 3-level "
                        f"@@@STEPS@@@ maximum, clamped: {text!r}"
                    )
                    nest = 2
                lines.append(f"{'  ' * nest}#. {text}")
                continue
            if empty_leaf:
                flush_steps()
                warnings.append(
                    "heading-styled paragraph has no content of its own "
                    "(sole child heading immediately followed by another "
                    "heading or end of document), rendered as a list "
                    f"item instead of a heading: {text!r}"
                )
                nest = last_list_ilvl + 1 if in_list else 0
                if not in_list:
                    lines.append("")
                    in_list = True
                lines.append(f"{'  ' * nest}- {text}")
                continue
            flush_all()
            lines.append(f"{markdown_heading_prefix(level, shift)} {titlecase_heading(text)}")
            lines.append("")
            continue

        if paragraph_is_list_item(block):
            flush_steps()
            if not in_list:
                lines.append("")
                in_list = True
            last_list_ilvl = paragraph_list_ilvl(block)
            lines.append(f"{'  ' * last_list_ilvl}- {text}")
            continue

        flush_all()
        lines.append(text)
        lines.append("")

    flush_all()
    return "\n".join(lines).strip() + "\n", warnings


def extract(docx_path, output_dir):
    """Extract docx_path into output_dir/<slug>.md plus output_dir/images/.
    Returns {'markdown_path': Path, 'images_dir': Path, 'warnings': list[str]}."""
    from docx import Document

    docx_path = Path(docx_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"
    images_dir.mkdir(exist_ok=True)

    doc = Document(docx_path)

    front_matter = {
        "title": docx_path.stem,
        "author": "",
        "department": "",
        "doc_number": "",
        "current_revision": "00",
        "department_head": "",
        "signature_fields": [],
        "revisions": [],
    }
    front_matter.update(extract_header_footer_metadata(doc))

    shift = compute_heading_shift(doc)
    blocks = list(iter_block_items(doc))
    body_text, warnings = build_markdown_body(doc, blocks, shift, images_dir, front_matter)

    footer_eco_number = front_matter.pop("footer_eco_number", None)
    footer_eco_date = front_matter.pop("footer_eco_date", None)
    if footer_eco_number or footer_eco_date:
        revisions = front_matter.get("revisions") or []
        if not revisions:
            warnings.append(
                f"header/footer shows ECO {footer_eco_number!r} dated "
                f"{footer_eco_date!r} but no revision-history table was "
                "found to compare against"
            )
        else:
            latest = revisions[-1]
            mismatches = []
            if footer_eco_number and footer_eco_number != latest.get("eco_number"):
                mismatches.append(
                    f"ECO number: header/footer says {footer_eco_number!r}, "
                    f"revision table says {latest.get('eco_number')!r}"
                )
            if footer_eco_date and footer_eco_date != latest.get("eco_date"):
                mismatches.append(
                    f"ECO date: header/footer says {footer_eco_date!r}, "
                    f"revision table says {latest.get('eco_date')!r}"
                )
            if mismatches:
                warnings.append(
                    "header/footer vs. revision table disagreement - " + "; ".join(mismatches)
                )

    slug = slugify(docx_path.stem)
    md_path = output_dir / f"{slug}.md"

    yaml_text = yaml.safe_dump(front_matter, sort_keys=False, allow_unicode=True)
    parts = [f"---\n{yaml_text}---\n\n"]
    if warnings:
        parts.append("\n".join(f"<!-- EXTRACTOR: {w} -->" for w in warnings))
        parts.append("\n\n")
    parts.append(body_text)

    md_path.write_text("".join(parts), encoding="utf-8")

    return {"markdown_path": md_path, "images_dir": images_dir, "warnings": warnings}


def main():
    if len(sys.argv) != 3:
        print("Usage: python extract_docx.py <input.docx> <output_dir>")
        return 1
    result = extract(sys.argv[1], sys.argv[2])
    print(f"Wrote {result['markdown_path']}")
    for w in result["warnings"]:
        print(f"[WARN] {w}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
