"""
Test suite for the dilon-document-form-compiler and
dilon-document-form-writer skills.

Direct-invocation style, matching tests/run_tests.py: a global
passed/failed counter via check(), explicit test calls from main(), no
pytest.
"""

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

REPO_ROOT = Path(__file__).parent.parent
FORM_COMPILER_DIR = REPO_ROOT / "skills" / "dilon-document-form-compiler"
SCRIPTS_DIR = FORM_COMPILER_DIR / "scripts"
FORM_WRITER_DIR = REPO_ROOT / "skills" / "dilon-document-form-writer"
FORM_TEMPLATE_PATH = FORM_WRITER_DIR / "TEMPLATE_Form.md"
# Shared with dilon-document-compiler - both skills use the same
# header/footer/styles-only base template, no form-specific copy.
BASE_TEMPLATE = REPO_ROOT / "templates" / "TEMPLATE_Word_Base.docx"
CHECK_DEPS_SCRIPT = SCRIPTS_DIR / "check_deps.py"
FORM_COMPILER_SCRIPT = SCRIPTS_DIR / "generate_dilon_form.py"
TEST_OUTPUT_DIR = Path(__file__).parent / "form-test-output"

SHEBANG_GUARDED_SCRIPTS = [
    CHECK_DEPS_SCRIPT,
    FORM_COMPILER_SCRIPT,
    Path(__file__),
]

sys.path.insert(0, str(SCRIPTS_DIR))
sys.path.insert(0, str(REPO_ROOT / "lib"))

passed = 0
failed = 0


def check(condition, message):
    global passed, failed
    if condition:
        print(f"[PASS] {message}")
        passed += 1
    else:
        print(f"[FAIL] {message}")
        failed += 1


SAMPLE_FORM_MARKDOWN = (
    '---\n'
    'title: "Detector Head Assy Traveler"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'doc_number: "FO-99999"\n'
    'current_revision: "00"\n'
    'department_head: "Test Head"\n'
    'signature_fields:\n'
    '  - department: "Regulatory"\n'
    '    name: "Test Rep"\n'
    '  - department: "Quality"\n'
    '    name: "Test QA"\n'
    'revisions:\n'
    '  - number: "00"\n'
    '    description: "Initial test"\n'
    '    eco_number: "ECO-000"\n'
    '    eco_date: "2025-01-01"\n'
    '---\n'
    '\n'
    '| Position | Serial Number |\n'
    '|----------|----------------|\n'
    '| 1        |                |\n'
    '| 2        |                |\n'
)


FIELD_GRID_PERMUTATIONS_MARKDOWN = (
    '---\n'
    'title: "Field Grid Permutations"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'doc_number: "FO-88888"\n'
    'current_revision: "00"\n'
    'department_head: "Test Head"\n'
    'signature_fields:\n'
    '  - department: "Regulatory"\n'
    '    name: "Test Rep"\n'
    '  - department: "Quality"\n'
    '    name: "Test QA"\n'
    'revisions:\n'
    '  - number: "00"\n'
    '    description: "Initial test"\n'
    '    eco_number: "ECO-000"\n'
    '    eco_date: "2025-01-01"\n'
    '---\n'
    '\n'
    '@@@FORM_FIELD:FieldGrid@@@\n'
    'Default Pair A: | Default Pair B:\n'
    'Label Split:[label=70] | Label Default:\n'
    'Explicit Pair:[pair=70] | Inferred Remainder:\n'
    'All Explicit A:[pair=30] | All Explicit B:[pair=30] | All Explicit C:[pair=40]\n'
    'Overshoot A:[pair=70] | Overshoot B:[pair=60]\n'
    'Row Rows Default: | Row Rows Default 2: {rows=2}\n'
    'Row Rows Override:[rows=3] | Row Rows Default 3: {rows=1}\n'
    'Vertical Single: {dir=v}\n'
    'Vertical Multi A: | Vertical Multi B: {dir=v,rows=2}\n'
    'Vertical Override A:[rows=4] | Vertical Override B: {dir=v,rows=2}\n'
    'Combined A:[label=60,pair=55,rows=2] | Combined B:[pair=45]\n'
    '@@@END_FORM_FIELD@@@\n'
    '\n'
    '@@@FORM_FIELD:FieldGrid:4in@@@\n'
    'Narrow A: | Narrow B:\n'
    '@@@END_FORM_FIELD@@@\n'
    '\n'
    '@@@FORM_FIELD:FillLine@@@Default Line:@@@END_FORM_FIELD@@@\n'
    '\n'
    '@@@FORM_FIELD:FillLine@@@Custom Width Line:[width=2.5in]@@@END_FORM_FIELD@@@\n'
    '\n'
    '@@@FORM_FIELD:FillLine@@@Multi Line Notes:[lines=3]@@@END_FORM_FIELD@@@\n'
    '\n'
    '@@@FORM_FIELD:FillLine@@@Conflicting Line:[width=2in,lines=2]@@@END_FORM_FIELD@@@\n'
    '\n'
    '| Field | Value |\n'
    '|---|---|\n'
    '| Cell FillLine | @@@FORM_FIELD:FillLine@@@Cell Line:@@@END_FORM_FIELD@@@ |\n'
)


def test_field_grid_permutations_compile():
    input_md = TEST_OUTPUT_DIR / "field_grid_permutations.md"
    output_docx = TEST_OUTPUT_DIR / "field_grid_permutations.docx"
    input_md.write_text(FIELD_GRID_PERMUTATIONS_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "field grid permutations markdown compiles cleanly")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "field_grid_permutations.docx created on disk")
    if not output_docx.exists():
        return

    doc = Document(output_docx)
    grid_tables = [t for t in doc.tables if t.style is not None and t.style.name == "Table Grid"]

    def table_for(first_cell_text):
        return next(
            (t for t in grid_tables if t.rows[0].cells[0].paragraphs[0].text == first_cell_text),
            None,
        )

    default_pair_table = table_for("Default Pair A:")
    check(default_pair_table is not None, "default horizontal pair row present")
    if default_pair_table is not None:
        check(len(default_pair_table.columns) == 4, f"2 horizontal pairs -> 4 columns, got {len(default_pair_table.columns)}")
        widths = [c.width.inches for c in default_pair_table.columns]
        check(abs(widths[0] - widths[1]) < 0.05, f"default 50/50 label/blank split within the pair, got {widths[0]}in/{widths[1]}in")

    label_split_table = table_for("Label Split:")
    check(label_split_table is not None, "label= split row present")
    if label_split_table is not None:
        widths = [c.width.inches for c in label_split_table.columns]
        check(widths[0] > widths[1], f"label=70 gives the label sub-cell more width than the blank, got {widths[0]}in/{widths[1]}in")

    explicit_pair_table = table_for("Explicit Pair:")
    check(explicit_pair_table is not None, "pair= with inferred remainder row present")
    if explicit_pair_table is not None:
        widths = [c.width.inches for c in explicit_pair_table.columns]
        pair1_width = widths[0] + widths[1]
        pair2_width = widths[2] + widths[3]
        check(pair1_width > pair2_width, f"pair=70 pair wider than the inferred-remainder pair, got {pair1_width}in vs {pair2_width}in")

    all_explicit_table = table_for("All Explicit A:")
    check(all_explicit_table is not None, "all-pairs-explicit pair= row present")
    if all_explicit_table is not None:
        check(len(all_explicit_table.columns) == 6, f"3 horizontal pairs -> 6 columns, got {len(all_explicit_table.columns)}")

    overshoot_table = table_for("Overshoot A:")
    check(overshoot_table is not None, "pair= overshoot row present (compiled despite invalid spec)")
    if overshoot_table is not None:
        widths = [c.width.inches for c in overshoot_table.columns]
        pair1_width = widths[0] + widths[1]
        pair2_width = widths[2] + widths[3]
        check(abs(pair1_width - pair2_width) < 0.05, f"overshoot falls back to an even split, got {pair1_width}in vs {pair2_width}in")

    row_rows_default_table = table_for("Row Rows Default:")
    check(row_rows_default_table is not None, "row-level rows= row present")
    if row_rows_default_table is not None:
        cells = row_rows_default_table.rows[0].cells
        check(len(cells[1].paragraphs) == 2, f"row-level rows=2 applies to both pairs' blanks, first blank got {len(cells[1].paragraphs)} paragraphs")
        check(len(cells[3].paragraphs) == 2, f"row-level rows=2 applies to both pairs' blanks, second blank got {len(cells[3].paragraphs)} paragraphs")

    row_rows_override_table = table_for("Row Rows Override:")
    check(row_rows_override_table is not None, "per-pair rows= override row present")
    if row_rows_override_table is not None:
        cells = row_rows_override_table.rows[0].cells
        check(len(cells[1].paragraphs) == 3, f"per-pair rows=3 override, first blank got {len(cells[1].paragraphs)} paragraphs")
        check(len(cells[3].paragraphs) == 1, f"row-level rows=1 default on the second pair, got {len(cells[3].paragraphs)} paragraphs")

    vertical_single_table = table_for("Vertical Single:")
    check(vertical_single_table is not None, "single-pair vertical row present")
    if vertical_single_table is not None:
        check(len(vertical_single_table.columns) == 1, f"1 vertical pair -> 1 column, got {len(vertical_single_table.columns)}")

    vertical_multi_table = table_for("Vertical Multi A:")
    check(vertical_multi_table is not None, "multi-pair vertical row present")
    if vertical_multi_table is not None:
        check(len(vertical_multi_table.columns) == 2, f"2 vertical pairs -> 2 columns, got {len(vertical_multi_table.columns)}")
        cells = vertical_multi_table.rows[0].cells
        check(len(cells[0].paragraphs) == 3, f"1 label + row-level rows=2 blanks = 3 paragraphs, got {len(cells[0].paragraphs)}")
        check(len(cells[1].paragraphs) == 3, f"1 label + row-level rows=2 blanks = 3 paragraphs, got {len(cells[1].paragraphs)}")

    vertical_override_table = table_for("Vertical Override A:")
    check(vertical_override_table is not None, "vertical row with a per-pair rows= override present")
    if vertical_override_table is not None:
        cells = vertical_override_table.rows[0].cells
        check(len(cells[0].paragraphs) == 5, f"1 label + per-pair rows=4 override = 5 paragraphs, got {len(cells[0].paragraphs)}")
        check(len(cells[1].paragraphs) == 3, f"1 label + row-level rows=2 default = 3 paragraphs, got {len(cells[1].paragraphs)}")

    combined_table = table_for("Combined A:")
    check(combined_table is not None, "label=+pair=+rows= combined row present")
    if combined_table is not None:
        widths = [c.width.inches for c in combined_table.columns]
        check(widths[0] > widths[1], f"label=60 within the first pair gives the label more width, got {widths[0]}in/{widths[1]}in")
        pair1_width = widths[0] + widths[1]
        pair2_width = widths[2] + widths[3]
        check(pair1_width > pair2_width, f"pair=55 pair wider than the pair=45 pair, got {pair1_width}in vs {pair2_width}in")
        cells = combined_table.rows[0].cells
        check(len(cells[1].paragraphs) == 2, f"per-pair rows=2 on the first pair's blank, got {len(cells[1].paragraphs)} paragraphs")

    narrow_table = table_for("Narrow A:")
    check(narrow_table is not None, "block-level max-width FieldGrid present")
    if narrow_table is not None:
        total_width = sum(c.width.inches for c in narrow_table.columns)
        check(abs(total_width - 4.0) < 0.05, f"block-level max width of 4in respected, got {total_width}in")

    fillline_paragraphs = {
        p.text.split('\t')[0]: p
        for p in doc.paragraphs
        if p.text.startswith(("Default Line:", "Custom Width Line:", "Multi Line Notes:", "Conflicting Line:"))
    }

    check("Default Line:" in fillline_paragraphs, "default FillLine present")
    if "Default Line:" in fillline_paragraphs:
        tab_stops = fillline_paragraphs["Default Line:"].paragraph_format.tab_stops
        check(len(tab_stops) == 1, f"default FillLine has exactly one tab stop, found {len(tab_stops)}")

    check("Custom Width Line:" in fillline_paragraphs, "width= FillLine present")
    if "Custom Width Line:" in fillline_paragraphs:
        tab_stops = fillline_paragraphs["Custom Width Line:"].paragraph_format.tab_stops
        check(
            len(tab_stops) == 1 and abs(tab_stops[0].position.inches - 2.5) < 0.05,
            f"width=2.5in respected, got {tab_stops[0].position.inches if len(tab_stops) == 1 else 'N/A'}in",
        )

    multi_line_start = next((i for i, p in enumerate(doc.paragraphs) if p.text == "Multi Line Notes:\t"), None)
    check(multi_line_start is not None, "lines= FillLine present")
    if multi_line_start is not None:
        following = doc.paragraphs[multi_line_start + 1 : multi_line_start + 3]
        check(
            len(following) == 2 and all(p.text == "\t" for p in following),
            f"lines=3 produces 2 additional bare-tab paragraphs, got {[p.text for p in following]!r}",
        )

    check("Conflicting Line:" in fillline_paragraphs, "width=+lines= conflicting FillLine present")
    if "Conflicting Line:" in fillline_paragraphs:
        section = doc.sections[0]
        available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        tab_stops = fillline_paragraphs["Conflicting Line:"].paragraph_format.tab_stops
        check(
            len(tab_stops) == 1 and abs(tab_stops[0].position.inches - available_width) < 0.05,
            f"lines= wins over width= when both given, expected full width {available_width}in, got "
            f"{tab_stops[0].position.inches if len(tab_stops) == 1 else 'N/A'}in",
        )

    cell_fillline_row = next(
        (row for t in doc.tables for row in t.rows if row.cells[0].text == "Cell FillLine"),
        None,
    )
    check(cell_fillline_row is not None, "table row containing the cell-level FillLine is present")
    if cell_fillline_row is not None:
        value_cell_text = cell_fillline_row.cells[1].text
        check("@@@" not in value_cell_text, "table-cell FillLine marker resolved (Task 1's fix), no leftover marker text")
        check(value_cell_text == "Cell Line:\t", f"table-cell FillLine rendered correctly, got {value_cell_text!r}")


FIELD_GRID_TITLE_ROW_MARKDOWN = (
    '---\n'
    'title: "Field Grid Title Row Test"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'doc_number: "FO-77778"\n'
    'current_revision: "00"\n'
    'department_head: "Test Head"\n'
    'signature_fields:\n'
    '  - department: "Regulatory"\n'
    '    name: "Test Rep"\n'
    '  - department: "Quality"\n'
    '    name: "Test QA"\n'
    'revisions:\n'
    '  - number: "00"\n'
    '    description: "Initial test"\n'
    '    eco_number: "ECO-000"\n'
    '    eco_date: "2025-01-01"\n'
    '---\n'
    '\n'
    '@@@FORM_FIELD:FieldGrid@@@\n'
    'Work Order: | Date:\n'
    'Assembly Prep {title=true}\n'
    'Technician: | Notes:\n'
    '@@@END_FORM_FIELD@@@\n'
)


def test_field_grid_title_row_mixed_with_normal_rows_compiles():
    input_md = TEST_OUTPUT_DIR / "field_grid_title_row.md"
    output_docx = TEST_OUTPUT_DIR / "field_grid_title_row.docx"
    input_md.write_text(FIELD_GRID_TITLE_ROW_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "field grid title row markdown compiles cleanly")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "field_grid_title_row.docx created on disk")
    if not output_docx.exists():
        return

    doc = Document(output_docx)
    grid_tables = [t for t in doc.tables if t.style is not None and t.style.name == "Table Grid"]
    check(len(grid_tables) == 3, f"3 declared rows produce 3 separate row-tables, found {len(grid_tables)}")
    if len(grid_tables) != 3:
        return

    work_order_table, title_table, technician_table = grid_tables

    check(len(work_order_table.columns) == 4, f"row before the title row still renders normally, got {len(work_order_table.columns)} columns")
    check(work_order_table.rows[0].cells[0].paragraphs[0].text == "Work Order:", "row before the title row keeps its label")

    check(len(title_table.columns) == 1, f"title row renders as 1 column, found {len(title_table.columns)}")
    title_cell = title_table.rows[0].cells[0]
    check(title_cell.paragraphs[0].text == "Assembly Prep", f"title row's label text, got {title_cell.paragraphs[0].text!r}")
    check(
        len(title_cell.paragraphs[0].runs) == 1 and title_cell.paragraphs[0].runs[0].bold is True,
        "title row's label is rendered as a single bold run",
    )

    check(len(technician_table.columns) == 4, f"row after the title row still renders normally, got {len(technician_table.columns)} columns")
    check(technician_table.rows[0].cells[0].paragraphs[0].text == "Technician:", "row after the title row keeps its label")


FO_00127_REPLICA_MARKDOWN = (
    '---\n'
    'title: "Detector Head Assembly Traveler"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'doc_number: "FO-00127"\n'
    'current_revision: "01"\n'
    'department_head: "Test Head"\n'
    'signature_fields:\n'
    '  - department: "Regulatory"\n'
    '    name: "Test Rep"\n'
    '  - department: "Quality"\n'
    '    name: "Test QA"\n'
    'revisions:\n'
    '  - number: "01"\n'
    '    description: "Initial test"\n'
    '    eco_number: "ECO-000055"\n'
    '    eco_date: "2025-03-13"\n'
    '---\n'
    '\n'
    '@@@FORM_FIELD:FieldGrid@@@\n'
    'Work Order: | Date:\n'
    '5mm GAGG Crystal Lot: | Technician:\n'
    'Carrier Board Assy Lot:\n'
    'Epoxy Lot # and Expiration:\n'
    'Cure Temp:[pair=60] | Start Time:[pair=40]\n'
    'End Time:\n'
    '@@@END_FORM_FIELD@@@\n'
    '\n'
    '@@@FORM_FIELD:FillLine@@@Alignment Fixture:[width=0.5in]@@@END_FORM_FIELD@@@\n'
    '\n'
    '| Document # | Rev | Date | Initial |\n'
    '|---|---|---|---|\n'
    '| WI-00077 |  |  |  |\n'
    '\n'
    '| Position | Clean / Inspect (Initial) | Epoxy Bead Check (Initial) | Alignment Check (Initial) | Bond Check (Initial) | Serial Number | Tested (Pass / Fail) | Pulled for WO# | Pulled for WO# |\n'
    '|---|---|---|---|---|---|---|---|---|\n'
    '| 1 |  |  |  |  |  |  | WO# |  |\n'
    '| 2 |  |  |  |  |  |  | WO# |  |\n'
)


def test_fo_00127_replica_compiles_with_expected_fields():
    input_md = TEST_OUTPUT_DIR / "fo_00127_replica.md"
    output_docx = TEST_OUTPUT_DIR / "fo_00127_replica.docx"
    input_md.write_text(FO_00127_REPLICA_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "FO-00127 replica compiles cleanly")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "fo_00127_replica.docx created on disk")
    if not output_docx.exists():
        return

    doc = Document(output_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    check("@@@" not in all_text, "no leftover marker text anywhere in the compiled document")

    field_grid_labels = [
        "Work Order:", "Date:", "5mm GAGG Crystal Lot:", "Technician:",
        "Carrier Board Assy Lot:", "Epoxy Lot # and Expiration:",
        "Cure Temp:", "Start Time:", "End Time:",
    ]
    field_grid_tables = [t for t in doc.tables if t.style is not None and t.style.name == "Table Grid"]
    grid_cell_texts = {
        cell.paragraphs[0].text
        for table in field_grid_tables
        for row in table.rows
        for cell in row.cells
    }
    missing_labels = [label for label in field_grid_labels if label not in grid_cell_texts]
    check(not missing_labels, f"every FieldGrid label from FO-00127 is present, missing: {missing_labels}")

    cure_temp_table = next(
        (t for t in field_grid_tables if t.rows[0].cells[0].paragraphs[0].text == "Cure Temp:"),
        None,
    )
    check(cure_temp_table is not None, "the Cure Temp/Start Time row-table is present")
    if cure_temp_table is not None:
        widths = [col.width.inches for col in cure_temp_table.columns]
        check(len(widths) == 4, f"2 horizontal pairs produce 4 columns, found {len(widths)}")

    alignment_paragraphs = [p for p in doc.paragraphs if p.text.startswith("Alignment Fixture:")]
    check(len(alignment_paragraphs) == 1, "Alignment Fixture rendered as a single FillLine paragraph")
    if alignment_paragraphs:
        tab_stops = alignment_paragraphs[0].paragraph_format.tab_stops
        check(
            len(tab_stops) == 1 and abs(tab_stops[0].position.inches - 0.5) < 0.05,
            f"Alignment Fixture uses its width=0.5in override, got "
            f"{tab_stops[0].position.inches if len(tab_stops) == 1 else 'N/A'}in",
        )

    qc_tables = [t for t in doc.tables if t.rows and t.rows[0].cells and t.rows[0].cells[0].text == "Position"]
    check(len(qc_tables) == 1, "the 9-column QC tracking table is present")
    if qc_tables:
        check(len(qc_tables[0].columns) == 9, f"QC table has 9 columns, found {len(qc_tables[0].columns)}")


def test_generate_form_document():
    input_md = TEST_OUTPUT_DIR / "form_test.md"
    output_docx = TEST_OUTPUT_DIR / "form_test.docx"
    input_md.write_text(SAMPLE_FORM_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "form compiler exits 0 for a valid form document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "form_test.docx created on disk")

    doc = Document(output_docx)
    section = doc.sections[0]
    header_text = " ".join(
        cell.text for table in section.header.tables for row in table.rows for cell in row.cells
    )
    check("Detector Head Assy Traveler" in header_text, "header title rendered (not literal {{title}})")
    check("FO-99999" in header_text, "header doc_number rendered")
    check("{{" not in header_text, "no unrendered Jinja2 braces remain in the header")

    body_tables = [t for t in doc.tables if any("Position" in c.text for c in t.rows[0].cells)]
    check(len(body_tables) == 1, "the form's own content table (Position/Serial Number) is present")
    if body_tables:
        check(body_tables[0].style.name == "DilonTable_List", "content table gets DilonTable_List styling")

    check(len(doc.tables) < 2 or not any(
        "Preparer" in c.text for t in doc.tables for row in t.rows for c in row.cells
    ), "no signature-approval table present in the compiled form")

    # No TOC: a TOC would normally be its own paragraphs/fields near the
    # top of Part D. With no headings in the form markdown, Pandoc's --toc
    # produces no TOC content either way, but confirm no 'Table of
    # Contents' heading text was emitted (would be a stray heading if TOC
    # generation were still forced on for a form).
    check(
        not any("Table of Contents" in p.text for p in doc.paragraphs),
        "no table of contents in the compiled form",
    )


def test_compile_default_output_filename_end_to_end():
    """Invoking with only <input.md> (no output, no template arg) should
    name the compiled file '<doc_number> Rev <current_revision>.docx' next
    to the input, using SAMPLE_FORM_MARKDOWN's doc_number/current_revision."""
    input_md = TEST_OUTPUT_DIR / "form_test_default_name.md"
    input_md.write_text(SAMPLE_FORM_MARKDOWN, encoding="utf-8")
    expected_output = TEST_OUTPUT_DIR / "FO-99999 Rev 00.docx"
    if expected_output.exists():
        expected_output.unlink()

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "form compiler exits 0 with only the input arg (default output filename)")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(expected_output.exists(), f"'{expected_output.name}' created via default output filename")


MISSING_DOC_NUMBER_FORM_MARKDOWN = (
    '---\n'
    'title: "Missing Doc Number Form Test"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'department_head: "Test Head"\n'
    'signature_fields: []\n'
    'revisions: []\n'
    '---\n'
    '\n'
    'This form has no doc_number/current_revision.\n'
)


def test_compile_default_output_filename_missing_doc_number_fails_clearly():
    """With no output arg and no doc_number/current_revision in front
    matter, the form compiler can't compute a default filename - it
    should halt with a clear error rather than minting ' Rev .docx'."""
    input_md = TEST_OUTPUT_DIR / "form_test_missing_doc_number.md"
    input_md.write_text(MISSING_DOC_NUMBER_FORM_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode != 0, "form compiler reports a non-zero exit code when doc_number/current_revision are missing and no output path is given")
    check("doc_number" in result.stderr + result.stdout, "error message mentions the missing doc_number/current_revision")


def test_underscore_until_end_of_line_body_paragraph():
    import dilon_form_fields as ff
    doc = Document()
    p = doc.add_paragraph("Work Order:")
    ff.underscore_until_end_of_line(p)

    check(p.text == "Work Order:\t", "label kept, followed by a tab character")
    tab_stops = p.paragraph_format.tab_stops
    check(len(tab_stops) == 1, f"exactly one tab stop added, found {len(tab_stops)}")
    if len(tab_stops) == 1:
        from docx.enum.text import WD_TAB_ALIGNMENT
        check(tab_stops[0].alignment == WD_TAB_ALIGNMENT.RIGHT, "tab stop is right-aligned")
        section = doc.sections[0]
        available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        check(
            abs(tab_stops[0].position.inches - available_width) < 0.01,
            f"tab stop positioned at the page's available content width ({available_width}in), got {tab_stops[0].position.inches}in",
        )


def test_underscore_until_end_of_line_in_table_cell():
    import dilon_form_fields as ff
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Inches(2.0)
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.0)
    p = cell.paragraphs[0]
    p.text = "Technician:"
    ff.underscore_until_end_of_line(p)

    tab_stops = p.paragraph_format.tab_stops
    check(len(tab_stops) == 1, f"exactly one tab stop added inside a table cell, found {len(tab_stops)}")
    if len(tab_stops) == 1:
        check(
            abs(tab_stops[0].position.inches - 2.0) < 0.05,
            f"tab stop positioned at the enclosing cell's width (2.0in), got {tab_stops[0].position.inches}in",
        )


def test_apply_form_fields_marker():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_FIELD:FillLine@@@Work Order:@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "form_fields_marker.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    check(len(result_doc.paragraphs) == 1, "marker paragraph is reused (not duplicated)")
    check(result_doc.paragraphs[0].text == "Work Order:\t", f"marker replaced with label + tab, got {result_doc.paragraphs[0].text!r}")
    check("@@@" not in result_doc.paragraphs[0].text, "no marker text remains")


def test_apply_form_fields_marker_in_table_cell():
    import dilon_form_fields as ff
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Inches(2.0)
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.0)
    cell.paragraphs[0].text = "@@@FORM_FIELD:FillLine@@@WO#@@@END_FORM_FIELD@@@"
    temp_path = TEST_OUTPUT_DIR / "form_fields_marker_in_cell.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    cell_text = result_doc.tables[0].rows[0].cells[0].text
    check("@@@" not in cell_text, "no marker text remains inside the table cell")
    check(cell_text == "WO#\t", f"marker replaced with label + tab inside a table cell, got {cell_text!r}")


def test_parse_bracket_annotations():
    import dilon_form_fields as ff
    check(ff.parse_bracket_annotations("Work Order:") == ("Work Order:", {}), "no bracket -> unchanged text, empty annotations")
    check(ff.parse_bracket_annotations("Cure Temp:[pair=60]") == ("Cure Temp:", {"pair": "60"}), "single annotation parsed")
    check(
        ff.parse_bracket_annotations("Notes:[dir=v,rows=3]") == ("Notes:", {"dir": "v", "rows": "3"}),
        "multiple comma-separated annotations parsed",
    )
    check(ff.parse_bracket_annotations("Label[malformed]") == ("Label", {}), "entry without '=' is ignored, not an error")


def test_underscore_until_end_of_line_width_override():
    import dilon_form_fields as ff
    doc = Document()
    p = doc.add_paragraph("Work Order:")
    ff.underscore_until_end_of_line(p, width_override=3.0)

    tab_stops = p.paragraph_format.tab_stops
    check(len(tab_stops) == 1, f"exactly one tab stop added, found {len(tab_stops)}")
    if len(tab_stops) == 1:
        check(abs(tab_stops[0].position.inches - 3.0) < 0.01, f"tab stop positioned at the given width (3.0in), got {tab_stops[0].position.inches}in")


def test_underscore_until_end_of_line_width_override_clamped():
    import dilon_form_fields as ff
    doc = Document()
    p = doc.add_paragraph("Work Order:")
    section = doc.sections[0]
    available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    ff.underscore_until_end_of_line(p, width_override=available_width + 5.0)

    tab_stops = p.paragraph_format.tab_stops
    check(
        len(tab_stops) == 1 and abs(tab_stops[0].position.inches - available_width) < 0.01,
        f"width_override exceeding the available width is clamped to it ({available_width}in), got "
        f"{tab_stops[0].position.inches if len(tab_stops) == 1 else 'N/A'}in",
    )


def test_underscore_until_end_of_line_multiple_lines():
    import dilon_form_fields as ff
    doc = Document()
    p = doc.add_paragraph("Notes:")
    ff.underscore_until_end_of_line(p, num_lines=3)

    all_paragraphs = doc.paragraphs
    check(len(all_paragraphs) == 3, f"num_lines=3 produces 3 paragraphs, found {len(all_paragraphs)}")
    if len(all_paragraphs) == 3:
        check(all_paragraphs[0].text == "Notes:\t", f"first paragraph keeps the label, got {all_paragraphs[0].text!r}")
        check(all_paragraphs[1].text == "\t", f"extra line is a bare tab, got {all_paragraphs[1].text!r}")
        check(all_paragraphs[2].text == "\t", f"extra line is a bare tab, got {all_paragraphs[2].text!r}")
        for para in all_paragraphs:
            tab_stops = para.paragraph_format.tab_stops
            check(len(tab_stops) == 1, f"each line has exactly one tab stop, found {len(tab_stops)}")


def test_underscore_until_end_of_line_multiple_lines_ignores_width():
    import dilon_form_fields as ff
    doc = Document()
    p = doc.add_paragraph("Notes:")
    section = doc.sections[0]
    available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    ff.underscore_until_end_of_line(p, width_override=2.0, num_lines=2)

    for para in doc.paragraphs:
        tab_stops = para.paragraph_format.tab_stops
        check(
            len(tab_stops) == 1 and abs(tab_stops[0].position.inches - available_width) < 0.01,
            f"num_lines > 1 ignores width_override, expected {available_width}in, got "
            f"{tab_stops[0].position.inches if len(tab_stops) == 1 else 'N/A'}in",
        )


def test_apply_form_fields_fillline_width_annotation():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_FIELD:FillLine@@@Work Order:[width=3in]@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "form_fields_fillline_width.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    check(result_doc.paragraphs[0].text == "Work Order:\t", f"annotation stripped from label, got {result_doc.paragraphs[0].text!r}")
    tab_stops = result_doc.paragraphs[0].paragraph_format.tab_stops
    check(
        len(tab_stops) == 1 and abs(tab_stops[0].position.inches - 3.0) < 0.01,
        f"width=3in annotation applied, got {tab_stops[0].position.inches if len(tab_stops) == 1 else 'N/A'}in",
    )


def test_apply_form_fields_fillline_lines_annotation():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_FIELD:FillLine@@@Notes:[lines=3]@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "form_fields_fillline_lines.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    check(len(result_doc.paragraphs) == 3, f"lines=3 annotation produces 3 paragraphs, found {len(result_doc.paragraphs)}")
    check(result_doc.paragraphs[0].text == "Notes:\t", f"first paragraph keeps the label, got {result_doc.paragraphs[0].text!r}")


def test_form_field_re_optional_width_group():
    import dilon_form_fields as ff
    match = ff.FORM_FIELD_RE.search("@@@FORM_FIELD:FieldGrid:6.5in@@@Work Order:@@@END_FORM_FIELD@@@")
    check(match is not None, "marker with a block-level width suffix matches")
    if match:
        check(match.group(1) == "FieldGrid", f"function name captured, got {match.group(1)!r}")
        check(match.group(2) == "6.5in", f"width suffix captured, got {match.group(2)!r}")
        check(match.group(3) == "Work Order:", f"block content captured, got {match.group(3)!r}")

    match_no_width = ff.FORM_FIELD_RE.search("@@@FORM_FIELD:FillLine@@@Work Order:@@@END_FORM_FIELD@@@")
    check(match_no_width is not None, "marker without a width suffix still matches")
    if match_no_width:
        check(match_no_width.group(2) is None, "width group is None when no suffix is given")
        check(match_no_width.group(3) == "Work Order:", f"block content captured, got {match_no_width.group(3)!r}")


def test_parse_row_annotation():
    import dilon_form_fields as ff
    check(ff.parse_row_annotation("Work Order: | Date:") == ("Work Order: | Date:", {}), "no row annotation -> unchanged, empty dict")
    check(ff.parse_row_annotation("Notes: {dir=v,rows=3}") == ("Notes:", {"dir": "v", "rows": "3"}), "row annotation parsed and stripped")


def test_parse_field_grid_block_simple():
    import dilon_form_fields as ff
    block = "Work Order: | Date:\nCarrier Board Assy Lot:\n"
    rows = ff.parse_field_grid_block(block)
    check(len(rows) == 2, f"two declared rows produce two parsed rows, found {len(rows)}")
    if len(rows) == 2:
        check(rows[0]['pairs'] == [("Work Order:", {}), ("Date:", {})], f"first row's pairs, got {rows[0]['pairs']!r}")
        check(rows[0]['annotations'] == {}, "first row has no row-level annotations")
        check(rows[1]['pairs'] == [("Carrier Board Assy Lot:", {})], f"second row's single pair, got {rows[1]['pairs']!r}")


def test_parse_field_grid_block_annotations():
    import dilon_form_fields as ff
    block = "Cure Temp:[pair=60] | Start Time:[pair=40]\nNotes: {dir=v,rows=3}\n"
    rows = ff.parse_field_grid_block(block)
    check(len(rows) == 2, f"expected 2 rows, found {len(rows)}")
    if len(rows) == 2:
        check(
            rows[0]['pairs'] == [("Cure Temp:", {"pair": "60"}), ("Start Time:", {"pair": "40"})],
            f"per-pair annotations parsed, got {rows[0]['pairs']!r}",
        )
        check(rows[1]['pairs'] == [("Notes:", {})], f"row-level annotation not mistaken for a per-pair one, got {rows[1]['pairs']!r}")
        check(rows[1]['annotations'] == {"dir": "v", "rows": "3"}, f"row-level annotations parsed, got {rows[1]['annotations']!r}")


def test_parse_field_grid_block_skips_blank_and_unparseable_lines():
    import dilon_form_fields as ff
    block = "Work Order:\n\n   \n|\nDate:\n"
    rows = ff.parse_field_grid_block(block)
    check(len(rows) == 2, f"blank lines and an unparseable row are skipped, found {len(rows)} rows")
    if len(rows) == 2:
        check(rows[0]['pairs'] == [("Work Order:", {})], f"got {rows[0]['pairs']!r}")
        check(rows[1]['pairs'] == [("Date:", {})], f"got {rows[1]['pairs']!r}")


def test_resolve_row_settings_defaults_and_overrides():
    import dilon_form_fields as ff
    check(ff.resolve_row_settings({}) == ('h', 1), "defaults to horizontal, rows=1")
    check(ff.resolve_row_settings({'dir': 'v', 'rows': '3'}) == ('v', 3), "explicit dir/rows honored")
    check(ff.resolve_row_settings({'dir': 'sideways'}) == ('h', 1), "invalid dir falls back to 'h'")
    check(ff.resolve_row_settings({'rows': '0'}) == ('h', 1), "rows < 1 falls back to 1")
    check(ff.resolve_row_settings({'rows': 'abc'}) == ('h', 1), "non-numeric rows falls back to 1")


def test_resolve_pair_rows():
    import dilon_form_fields as ff
    check(ff.resolve_pair_rows({}, 2) == 2, "falls back to the row default when absent")
    check(ff.resolve_pair_rows({'rows': '5'}, 2) == 5, "per-pair override takes precedence")
    check(ff.resolve_pair_rows({'rows': '0'}, 2) == 2, "invalid override falls back to the row default")


def test_resolve_pair_widths_even_split():
    import dilon_form_fields as ff
    pairs = [("A:", {}), ("B:", {})]
    widths = ff.resolve_pair_widths(pairs, 6.0)
    check(widths == [3.0, 3.0], f"default even split across 2 pairs, got {widths}")

    pairs3 = [("A:", {}), ("B:", {}), ("C:", {})]
    widths3 = ff.resolve_pair_widths(pairs3, 6.0)
    check(widths3 == [2.0, 2.0, 2.0], f"default even split across 3 pairs, got {widths3}")


def test_resolve_pair_widths_explicit_and_remainder():
    import dilon_form_fields as ff
    pairs = [("A:", {'pair': '60'}), ("B:", {})]
    widths = ff.resolve_pair_widths(pairs, 10.0)
    check(widths == [6.0, 4.0], f"explicit pair=60 with remainder auto-filled, got {widths}")


def test_resolve_pair_widths_overshoot_falls_back():
    import dilon_form_fields as ff
    pairs = [("A:", {'pair': '60'}), ("B:", {'pair': '50'})]
    widths = ff.resolve_pair_widths(pairs, 10.0)
    check(widths == [5.0, 5.0], f"declared total > 100 falls back to an even split, got {widths}")


def test_resolve_pair_widths_nonpositive_remainder_falls_back():
    import dilon_form_fields as ff
    pairs = [("A:", {'pair': '100'}), ("B:", {})]
    widths = ff.resolve_pair_widths(pairs, 10.0)
    check(widths == [5.0, 5.0], f"a fully-declared pair leaving no room for an undeclared one falls back to an even split, got {widths}")


def test_resolve_label_width():
    import dilon_form_fields as ff
    check(ff.resolve_label_width({}, 10.0, 'h') == (5.0, 5.0), "default 50/50 split")
    check(ff.resolve_label_width({'label': '70'}, 10.0, 'h') == (7.0, 3.0), "explicit label= split")
    check(ff.resolve_label_width({'label': '70'}, 10.0, 'v') == (None, None), "dir=v ignores label=, returns (None, None)")
    check(ff.resolve_label_width({}, 10.0, 'v') == (None, None), "dir=v with no label= still returns (None, None)")


def test_resolve_title_flag():
    import dilon_form_fields as ff
    check(ff.resolve_title_flag({}) is False, "absent title= is not a title row")
    check(ff.resolve_title_flag({'title': 'true'}) is True, "title=true is a title row")
    check(ff.resolve_title_flag({'title': 'True'}) is True, "title= truthy match is case-insensitive")
    check(ff.resolve_title_flag({'title': '1'}) is True, "title=1 is a title row")
    check(ff.resolve_title_flag({'title': 'yes'}) is True, "title=yes is a title row")
    check(ff.resolve_title_flag({'title': 'YES'}) is True, "title=YES is a title row (case-insensitive)")
    check(ff.resolve_title_flag({'title': 'false'}) is False, "title=false is not a title row")
    check(ff.resolve_title_flag({'title': '0'}) is False, "title=0 is not a title row")
    check(ff.resolve_title_flag({'title': 't'}) is False, "title=t is NOT a title row (only true/1/yes are truthy)")
    check(ff.resolve_title_flag({'title': 'nope'}) is False, "unrecognized title= value is not a title row")


def test_build_field_grid_row_table_horizontal():
    import dilon_form_fields as ff
    doc = Document()
    row = {'pairs': [("Work Order:", {}), ("Date:", {})], 'annotations': {}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)

    check(table.style.name == "Table Grid", f"row-table uses the Table Grid style, got {table.style.name!r}")
    check(len(table.columns) == 4, f"2 horizontal pairs produce 4 columns, found {len(table.columns)}")
    cells = table.rows[0].cells
    check(cells[0].paragraphs[0].text == "Work Order:", f"first label cell text, got {cells[0].paragraphs[0].text!r}")
    check(cells[1].text == "", "first blank cell starts empty")
    check(cells[2].paragraphs[0].text == "Date:", f"second label cell text, got {cells[2].paragraphs[0].text!r}")
    check(cells[3].text == "", "second blank cell starts empty")
    check(len(cells[1].paragraphs) == 1, f"rows=1 default: blank cell has exactly 1 paragraph, found {len(cells[1].paragraphs)}")


def test_build_field_grid_row_table_vertical_with_rows():
    import dilon_form_fields as ff
    doc = Document()
    row = {'pairs': [("Notes:", {})], 'annotations': {'dir': 'v', 'rows': '3'}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)

    check(len(table.columns) == 1, f"1 vertical pair produces 1 column, found {len(table.columns)}")
    cell = table.rows[0].cells[0]
    check(cell.paragraphs[0].text == "Notes:", f"label paragraph text, got {cell.paragraphs[0].text!r}")
    check(len(cell.paragraphs) == 4, f"1 label paragraph + rows=3 blank paragraphs = 4 total, found {len(cell.paragraphs)}")
    for blank_paragraph in cell.paragraphs[1:]:
        check(blank_paragraph.text == "", f"blank paragraph is empty, got {blank_paragraph.text!r}")


def test_build_field_grid_row_table_is_centered():
    import dilon_form_fields as ff
    from docx.enum.table import WD_TABLE_ALIGNMENT
    doc = Document()
    row = {'pairs': [("A:", {})], 'annotations': {}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)
    check(table.alignment == WD_TABLE_ALIGNMENT.CENTER, f"row-table is centered, got {table.alignment}")


def test_build_field_grid_row_table_title_row():
    import dilon_form_fields as ff
    doc = Document()
    row = {'pairs': [("Assembly Prep", {})], 'annotations': {'title': 'true'}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)

    check(table.style.name == "Table Grid", f"title row-table uses the Table Grid style, got {table.style.name!r}")
    check(len(table.columns) == 1, f"title row produces exactly 1 column, found {len(table.columns)}")
    check(abs(table.columns[0].width.inches - 6.0) < 0.05, f"title row's single column spans the full row width, got {table.columns[0].width.inches}in")
    cell = table.rows[0].cells[0]
    check(len(cell.paragraphs) == 1, f"title row has exactly 1 paragraph (no blank cell/lines), found {len(cell.paragraphs)}")
    check(cell.paragraphs[0].text == "Assembly Prep", f"title row's label text, got {cell.paragraphs[0].text!r}")
    check(len(cell.paragraphs[0].runs) == 1, f"title row's label is a single run, found {len(cell.paragraphs[0].runs)}")
    if cell.paragraphs[0].runs:
        check(cell.paragraphs[0].runs[0].bold is True, "title row's label run is bold")


def test_build_field_grid_row_table_title_row_ignores_pair_and_label_annotations():
    import dilon_form_fields as ff
    doc = Document()
    row = {'pairs': [("Assembly Prep", {'pair': '60', 'label': '70'})], 'annotations': {'title': 'true'}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)

    check(len(table.columns) == 1, f"title row ignores pair= and stays 1 column, found {len(table.columns)}")
    check(abs(table.columns[0].width.inches - 6.0) < 0.05, f"title row ignores pair= and still spans the full row width, got {table.columns[0].width.inches}in")
    cell = table.rows[0].cells[0]
    check(cell.paragraphs[0].text == "Assembly Prep", f"title row ignores label=, single cell text is unaffected, got {cell.paragraphs[0].text!r}")


def test_build_field_grid_row_table_title_row_rows_override():
    import dilon_form_fields as ff
    doc = Document()
    row = {'pairs': [("Assembly Prep", {})], 'annotations': {'title': 'true', 'rows': '3'}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)

    cell = table.rows[0].cells[0]
    check(len(cell.paragraphs) == 3, f"title row rows=3 adds 2 blank paragraphs below the bold label, found {len(cell.paragraphs)}")
    check(cell.paragraphs[0].text == "Assembly Prep", f"first paragraph keeps the bold label, got {cell.paragraphs[0].text!r}")
    for blank_paragraph in cell.paragraphs[1:]:
        check(blank_paragraph.text == "", f"extra title-row paragraph is blank, got {blank_paragraph.text!r}")


def test_build_field_grid_row_table_title_row_multiple_tokens_uses_first():
    import dilon_form_fields as ff
    doc = Document()
    row = {'pairs': [("Assembly Prep", {}), ("Dropped:", {})], 'annotations': {'title': 'true'}}
    table = ff.build_field_grid_row_table(doc, row, 6.0)

    check(len(table.columns) == 1, f"title row with multiple pair tokens still produces 1 column, found {len(table.columns)}")
    cell = table.rows[0].cells[0]
    check(cell.paragraphs[0].text == "Assembly Prep", f"title row uses only the first pair token's label, got {cell.paragraphs[0].text!r}")


def test_build_field_grid_row_table_title_row_dir_has_no_visible_effect():
    import dilon_form_fields as ff
    doc = Document()
    row_h = {'pairs': [("Assembly Prep", {})], 'annotations': {'title': 'true', 'dir': 'h'}}
    table_h = ff.build_field_grid_row_table(doc, row_h, 6.0)
    check(len(table_h.columns) == 1, f"title row with dir=h still renders as 1 column, found {len(table_h.columns)}")

    row_v = {'pairs': [("Assembly Prep", {})], 'annotations': {'title': 'true', 'dir': 'v'}}
    table_v = ff.build_field_grid_row_table(doc, row_v, 6.0)
    check(len(table_v.columns) == 1, f"title row with dir=v still renders as 1 column, found {len(table_v.columns)}")


def test_insert_field_grid_replaces_marker_with_row_tables():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_FIELD:FieldGrid@@@\nWork Order: | Date:\nCarrier Board Assy Lot:\n@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "field_grid_marker.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    check(not any("@@@" in p.text for p in result_doc.paragraphs), "no marker text remains in the body")
    check(len(result_doc.tables) == 2, f"2 declared rows produce 2 separate row-tables, found {len(result_doc.tables)}")
    if len(result_doc.tables) == 2:
        check(result_doc.tables[0].rows[0].cells[0].paragraphs[0].text == "Work Order:", "first row's first label")
        check(result_doc.tables[1].rows[0].cells[0].paragraphs[0].text == "Carrier Board Assy Lot:", "second row's label")


def test_insert_field_grid_max_width():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_FIELD:FieldGrid:3in@@@\nA: | B:\n@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "field_grid_max_width.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    table = result_doc.tables[0]
    total_width = sum(col.width.inches for col in table.columns)
    check(abs(total_width - 3.0) < 0.05, f"block-level max width of 3in respected, got {total_width}in")


def test_field_grid_marker_inside_table_cell_warns_and_skips():
    import dilon_form_fields as ff
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Inches(2.0)
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.0)
    cell.paragraphs[0].text = "@@@FORM_FIELD:FieldGrid@@@\nA: | B:\n@@@END_FORM_FIELD@@@"
    temp_path = TEST_OUTPUT_DIR / "field_grid_in_cell.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    cell_text = result_doc.tables[0].rows[0].cells[0].text
    check("@@@FORM_FIELD:FieldGrid@@@" in cell_text, "FieldGrid marker inside a table cell is left as-is, not processed")


def test_protect_field_grid_line_breaks():
    import dilon_form_fields as ff
    source = (
        "@@@FORM_FIELD:FieldGrid@@@\n"
        "Work Order: | Date:\n"
        "Carrier Board Assy Lot:\n"
        "@@@END_FORM_FIELD@@@\n"
    )
    protected = ff.protect_field_grid_line_breaks(source)
    check(
        "Work Order: | Date:  \n" in protected,
        f"non-blank row line gets a trailing hard-break, got {protected!r}",
    )
    check(
        "Carrier Board Assy Lot:  \n" in protected,
        f"last row line also gets a trailing hard-break, got {protected!r}",
    )

    # FillLine markers and text outside FieldGrid blocks are untouched.
    unrelated = "@@@FORM_FIELD:FillLine@@@Work Order:@@@END_FORM_FIELD@@@\n\nSome body text.\n"
    check(ff.protect_field_grid_line_breaks(unrelated) == unrelated, "text with no FieldGrid block is unchanged")


def test_form_section_ranges_single_pair():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("Before - not in a section")
    doc.add_paragraph("@@@FORM_SECTION@@@")
    inside = doc.add_paragraph("Inside the section")
    doc.add_paragraph("@@@END_FORM_SECTION@@@")
    doc.add_paragraph("After - not in a section")

    in_section, mutated = ff._form_section_ranges(doc)

    check(mutated is True, "a declared section mutates the document (sentinels removed)")
    check(id(inside._p) in in_section, "the paragraph between BEGIN/END is recorded as in-section")
    check(
        not any(p.text.startswith("@@@FORM_SECTION") or p.text.startswith("@@@END_FORM_SECTION") for p in doc.paragraphs),
        "sentinel paragraphs are removed from the document",
    )
    check(len(doc.paragraphs) == 3, f"3 paragraphs remain after the 2 sentinels are removed, found {len(doc.paragraphs)}")


def test_form_section_ranges_content_outside_not_recorded():
    import dilon_form_fields as ff
    doc = Document()
    before = doc.add_paragraph("Before")
    doc.add_paragraph("@@@FORM_SECTION@@@")
    doc.add_paragraph("Inside")
    doc.add_paragraph("@@@END_FORM_SECTION@@@")
    after = doc.add_paragraph("After")

    in_section, _ = ff._form_section_ranges(doc)

    check(id(before._p) not in in_section, "content before the section is not recorded as in-section")
    check(id(after._p) not in in_section, "content after the section is not recorded as in-section")


def test_form_section_ranges_no_sections_present():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("Just a plain paragraph")

    in_section, mutated = ff._form_section_ranges(doc)

    check(in_section == set(), "no @@@FORM_SECTION@@@ markers -> empty in-section set")
    check(mutated is False, "no @@@FORM_SECTION@@@ markers -> document not mutated")


def test_form_section_ranges_unmatched_end_raises():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@END_FORM_SECTION@@@")

    raised = False
    try:
        ff._form_section_ranges(doc)
    except ff.FormSectionError:
        raised = True
    check(raised, "@@@END_FORM_SECTION@@@ with no open section raises FormSectionError")


def test_form_section_ranges_unclosed_begin_raises():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_SECTION@@@")
    doc.add_paragraph("Inside, never closed")

    raised = False
    try:
        ff._form_section_ranges(doc)
    except ff.FormSectionError:
        raised = True
    check(raised, "an unclosed @@@FORM_SECTION@@@ at end-of-document raises FormSectionError")


def test_form_section_ranges_nested_begin_raises():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_SECTION@@@")
    doc.add_paragraph("Outer content")
    doc.add_paragraph("@@@FORM_SECTION@@@")

    raised = False
    try:
        ff._form_section_ranges(doc)
    except ff.FormSectionError:
        raised = True
    check(raised, "a nested @@@FORM_SECTION@@@ before the outer one closes raises FormSectionError")


def test_form_section_ranges_multiple_nonoverlapping_pairs():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_SECTION@@@")
    first_inside = doc.add_paragraph("First section content")
    doc.add_paragraph("@@@END_FORM_SECTION@@@")
    doc.add_paragraph("Between sections")
    doc.add_paragraph("@@@FORM_SECTION@@@")
    second_inside = doc.add_paragraph("Second section content")
    doc.add_paragraph("@@@END_FORM_SECTION@@@")

    in_section, mutated = ff._form_section_ranges(doc)

    check(mutated is True, "two declared sections mutate the document")
    check(id(first_inside._p) in in_section, "first section's content recorded as in-section")
    check(id(second_inside._p) in in_section, "second section's content recorded as in-section")


def test_form_section_ranges_includes_tables():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_SECTION@@@")
    table = doc.add_table(rows=1, cols=1)
    doc.add_paragraph("@@@END_FORM_SECTION@@@")

    in_section, _ = ff._form_section_ranges(doc)

    check(id(table._element) in in_section, "a table between BEGIN/END is recorded as in-section")


def test_apply_form_fields_form_section_header_numbers_sequentially():
    import dilon_form_fields as ff
    from docx.enum.style import WD_STYLE_TYPE
    doc = Document()
    doc.styles.add_style('Form Section Header', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("@@@FORM_FIELD:Form_Section_Header@@@Assembly Prep@@@END_FORM_FIELD@@@")
    doc.add_paragraph("some body content")
    doc.add_paragraph("@@@FORM_FIELD:Form_Section_Header@@@Final Inspection@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "form_section_header_sequential.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    check(result_doc.paragraphs[0].text == "Section 1 - Assembly Prep", f"first marker numbered 1, got {result_doc.paragraphs[0].text!r}")
    check(result_doc.paragraphs[0].style.name == "Form Section Header", f"first marker gets the Form Section Header style, got {result_doc.paragraphs[0].style.name!r}")
    check(result_doc.paragraphs[2].text == "Section 2 - Final Inspection", f"second marker numbered 2, got {result_doc.paragraphs[2].text!r}")
    check(result_doc.paragraphs[2].style.name == "Form Section Header", f"second marker gets the Form Section Header style, got {result_doc.paragraphs[2].style.name!r}")


def test_apply_form_fields_form_section_header_missing_style_degrades():
    import dilon_form_fields as ff
    doc = Document()
    doc.add_paragraph("@@@FORM_FIELD:Form_Section_Header@@@Assembly Prep@@@END_FORM_FIELD@@@")
    temp_path = TEST_OUTPUT_DIR / "form_section_header_missing_style.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    check(result_doc.paragraphs[0].text == "Section 1 - Assembly Prep", f"marker still numbered and rendered without the style, got {result_doc.paragraphs[0].text!r}")
    check("@@@" not in result_doc.paragraphs[0].text, "no marker text remains")


def test_form_section_header_marker_inside_table_cell_warns_and_skips():
    import dilon_form_fields as ff
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Inches(2.0)
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.0)
    cell.paragraphs[0].text = "@@@FORM_FIELD:Form_Section_Header@@@Assembly Prep@@@END_FORM_FIELD@@@"
    temp_path = TEST_OUTPUT_DIR / "form_section_header_in_cell.docx"
    doc.save(temp_path)

    ff.apply_form_fields(temp_path)

    result_doc = Document(temp_path)
    cell_text = result_doc.tables[0].rows[0].cells[0].text
    check("@@@FORM_FIELD:Form_Section_Header@@@" in cell_text, "Form_Section_Header marker inside a table cell is left as-is, not processed")


def test_form_section_header_compiles_through_full_pipeline():
    input_md = TEST_OUTPUT_DIR / "form_section_header.md"
    output_docx = TEST_OUTPUT_DIR / "form_section_header.docx"
    input_md.write_text(
        '---\n'
        'title: "Section Header Test"\n'
        'doc_number: "FO-77777"\n'
        'current_revision: "00"\n'
        '---\n'
        '\n'
        '@@@FORM_FIELD:Form_Section_Header@@@Assembly Prep@@@END_FORM_FIELD@@@\n'
        '\n'
        '@@@FORM_FIELD:FillLine@@@Work Order:@@@END_FORM_FIELD@@@\n'
        '\n'
        '@@@FORM_FIELD:Form_Section_Header@@@Final Inspection@@@END_FORM_FIELD@@@\n',
        encoding="utf-8",
    )

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "Form_Section_Header markdown compiles cleanly")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "form_section_header.docx created on disk")
    if not output_docx.exists():
        return

    doc = Document(output_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    check("@@@" not in all_text, "no leftover marker text in the compiled document")
    check("Section 1 - Assembly Prep" in all_text, f"first section numbered and titled correctly, got {all_text!r}")
    check("Section 2 - Final Inspection" in all_text, f"second section numbered and titled correctly, got {all_text!r}")


def test_form_compile_has_no_leading_blank_paragraph():
    input_md = TEST_OUTPUT_DIR / "no_leading_blank.md"
    output_docx = TEST_OUTPUT_DIR / "no_leading_blank.docx"
    input_md.write_text(
        '---\n'
        'title: "No Leading Blank Test"\n'
        'doc_number: "FO-55555"\n'
        'current_revision: "00"\n'
        '---\n'
        '\n'
        '@@@FORM_FIELD:Form_Section_Header@@@First Section@@@END_FORM_FIELD@@@\n'
        '\n'
        '@@@FORM_FIELD:FillLine@@@Work Order:@@@END_FORM_FIELD@@@\n',
        encoding="utf-8",
    )

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "no-leading-blank markdown compiles cleanly")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "no_leading_blank.docx created on disk")
    if not output_docx.exists():
        return

    doc = Document(output_docx)
    check(len(doc.paragraphs) > 0, "compiled document has at least one paragraph")
    if doc.paragraphs:
        check(
            doc.paragraphs[0].text == "Section 1 - First Section",
            f"first paragraph is the document's real first content, no stray leading blank line carried over from the base template's Part A, got {doc.paragraphs[0].text!r}",
        )


def generate_form_stub(output_path, **overrides):
    """Python port of the dilon-document-form-writer stub-generation
    logic (lives in SKILL.md as plain instructions for Claude; ported
    here so it can be exercised by an automated test), mirroring
    tests/run_tests.py's generate_stub() for dilon-document-writer.
    Defaults doc_number to the FO- form-number convention rather than
    dilon-document-writer's narrative-doc DD_XXX_XXXXX."""
    output_path = Path(output_path)
    if output_path.exists():
        raise FileExistsError(f"Output file already exists: {output_path}")

    template = FORM_TEMPLATE_PATH.read_text(encoding="utf-8")

    values = {
        "title": overrides.get("title", "Form Title"),
        "author": overrides.get("author", "Author Name"),
        "department": overrides.get("department", "--"),
        "doc_number": overrides.get("doc_number", "FO-XXXXX"),
        "current_revision": overrides.get("current_revision", "00"),
        "department_head": overrides.get("department_head", "--"),
        "revision_description": overrides.get("revision_description", "Initial release"),
        "eco_number": overrides.get("eco_number", "ECO-TBD"),
        "eco_date": overrides.get("eco_date", "YYYY-MM-DD"),
    }

    content = template
    content = re.sub(r'title: ".*?"', f'title: "{values["title"]}"', content, count=1)
    content = re.sub(r'author: ".*?"', f'author: "{values["author"]}"', content, count=1)
    content = re.sub(r'department: ".*?"', f'department: "{values["department"]}"', content, count=1)
    content = re.sub(r'doc_number: ".*?"', f'doc_number: "{values["doc_number"]}"', content, count=1)
    content = re.sub(r'current_revision: ".*?"', f'current_revision: "{values["current_revision"]}"', content, count=1)
    content = re.sub(r'department_head: ".*?"', f'department_head: "{values["department_head"]}"', content, count=1)
    content = re.sub(
        r'- number: ".*?"\s+description: ".*?"\s+eco_number: ".*?"\s+eco_date: ".*?"',
        '- number: "{0}"\n    description: "{1}"\n    eco_number: "{2}"\n    eco_date: "{3}"'.format(
            values["current_revision"], values["revision_description"], values["eco_number"], values["eco_date"]
        ),
        content,
        count=1,
    )

    output_path.write_text(content, encoding="utf-8")


def test_form_stub_custom_params():
    path = TEST_OUTPUT_DIR / "custom_form_stub.md"
    generate_form_stub(
        path,
        title="Detector Head Assy Traveler",
        author="QA Team",
        doc_number="FO-00200",
        department="Quality",
        current_revision="01",
    )
    check(path.exists(), "custom_form_stub.md created")
    content = path.read_text(encoding="utf-8")
    check('title: "Detector Head Assy Traveler"' in content, "custom title substituted")
    check('doc_number: "FO-00200"' in content, "custom doc_number substituted")
    check('department: "Quality"' in content, "custom department substituted")
    check('current_revision: "01"' in content, "custom current_revision substituted")
    check('number: "01"' in content, "initial revision entry mirrors current_revision")


def test_form_stub_default_params():
    path = TEST_OUTPUT_DIR / "default_form_stub.md"
    generate_form_stub(path)
    check(path.exists(), "default_form_stub.md created")
    content = path.read_text(encoding="utf-8")
    check('title: "Form Title"' in content, "default title used")
    check(
        'doc_number: "FO-XXXXX"' in content,
        "default doc_number uses the FO- form-number convention, not dilon-document-writer's DD_XXX_XXXXX",
    )


def test_form_stub_refuses_overwrite():
    path = TEST_OUTPUT_DIR / "custom_form_stub.md"  # already created above
    raised = False
    try:
        generate_form_stub(path)
    except FileExistsError:
        raised = True
    check(raised, "generate_form_stub refuses to overwrite an existing file")


def test_form_template_contains_worked_example_markers():
    content = FORM_TEMPLATE_PATH.read_text(encoding="utf-8")
    check("@@@FORM_FIELD:Form_Section_Header@@@" in content, "template includes a Form_Section_Header example")
    check("@@@FORM_FIELD:FieldGrid@@@" in content, "template includes a FieldGrid example")
    check("@@@FORM_FIELD:FillLine@@@" in content, "template includes a FillLine example")


def test_form_template_compiles_through_full_pipeline():
    input_md = TEST_OUTPUT_DIR / "form_template_compile.md"
    output_docx = TEST_OUTPUT_DIR / "form_template_compile.docx"
    generate_form_stub(input_md, title="Template Compile Test", doc_number="FO-00300")

    result = subprocess.run(
        [
            sys.executable,
            str(FORM_COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(BASE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "TEMPLATE_Form.md's worked example compiles cleanly")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "form_template_compile.docx created on disk")
    if output_docx.exists():
        doc = Document(output_docx)
        all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        check("@@@" not in all_text, "no leftover marker text in the compiled worked example")


def test_no_shebang_in_form_compiler_scripts():
    def has_shebang(path):
        lines = path.read_text(encoding="utf-8").splitlines()
        return bool(lines) and lines[0].startswith("#!")

    offenders = [str(p) for p in SHEBANG_GUARDED_SCRIPTS if has_shebang(p)]
    check(not offenders, f"no shebang lines in guarded form-compiler scripts (offenders: {offenders})")


def test_check_deps_runs_and_reports():
    result = subprocess.run(
        [sys.executable, str(CHECK_DEPS_SCRIPT)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "check_deps.py exits 0 when dependencies are installed")
    check("pandoc" in result.stdout.lower(), "check_deps.py reports on pandoc")
    check("jinja2" in result.stdout.lower(), "check_deps.py reports on the jinja2 module")


def main():
    if TEST_OUTPUT_DIR.exists():
        import shutil
        shutil.rmtree(TEST_OUTPUT_DIR)
    TEST_OUTPUT_DIR.mkdir(parents=True)

    test_generate_form_document()
    test_compile_default_output_filename_end_to_end()
    test_compile_default_output_filename_missing_doc_number_fails_clearly()
    test_underscore_until_end_of_line_body_paragraph()
    test_underscore_until_end_of_line_in_table_cell()
    test_apply_form_fields_marker()
    test_apply_form_fields_marker_in_table_cell()
    test_parse_bracket_annotations()
    test_underscore_until_end_of_line_width_override()
    test_underscore_until_end_of_line_width_override_clamped()
    test_underscore_until_end_of_line_multiple_lines()
    test_underscore_until_end_of_line_multiple_lines_ignores_width()
    test_apply_form_fields_fillline_width_annotation()
    test_apply_form_fields_fillline_lines_annotation()
    test_form_field_re_optional_width_group()
    test_parse_row_annotation()
    test_parse_field_grid_block_simple()
    test_parse_field_grid_block_annotations()
    test_parse_field_grid_block_skips_blank_and_unparseable_lines()
    test_resolve_row_settings_defaults_and_overrides()
    test_resolve_pair_rows()
    test_resolve_pair_widths_even_split()
    test_resolve_pair_widths_explicit_and_remainder()
    test_resolve_pair_widths_overshoot_falls_back()
    test_resolve_pair_widths_nonpositive_remainder_falls_back()
    test_resolve_label_width()
    test_resolve_title_flag()
    test_build_field_grid_row_table_horizontal()
    test_build_field_grid_row_table_vertical_with_rows()
    test_build_field_grid_row_table_is_centered()
    test_build_field_grid_row_table_title_row()
    test_build_field_grid_row_table_title_row_ignores_pair_and_label_annotations()
    test_build_field_grid_row_table_title_row_rows_override()
    test_build_field_grid_row_table_title_row_multiple_tokens_uses_first()
    test_build_field_grid_row_table_title_row_dir_has_no_visible_effect()
    test_insert_field_grid_replaces_marker_with_row_tables()
    test_insert_field_grid_max_width()
    test_field_grid_marker_inside_table_cell_warns_and_skips()
    test_protect_field_grid_line_breaks()
    test_form_section_ranges_single_pair()
    test_form_section_ranges_content_outside_not_recorded()
    test_form_section_ranges_no_sections_present()
    test_form_section_ranges_unmatched_end_raises()
    test_form_section_ranges_unclosed_begin_raises()
    test_form_section_ranges_nested_begin_raises()
    test_form_section_ranges_multiple_nonoverlapping_pairs()
    test_form_section_ranges_includes_tables()
    test_field_grid_permutations_compile()
    test_field_grid_title_row_mixed_with_normal_rows_compiles()
    test_apply_form_fields_form_section_header_numbers_sequentially()
    test_apply_form_fields_form_section_header_missing_style_degrades()
    test_form_section_header_marker_inside_table_cell_warns_and_skips()
    test_form_section_header_compiles_through_full_pipeline()
    test_form_compile_has_no_leading_blank_paragraph()
    test_form_stub_custom_params()
    test_form_stub_default_params()
    test_form_stub_refuses_overwrite()
    test_form_template_contains_worked_example_markers()
    test_form_template_compiles_through_full_pipeline()
    test_no_shebang_in_form_compiler_scripts()
    test_check_deps_runs_and_reports()
    test_fo_00127_replica_compiles_with_expected_fields()

    print(f"\n{passed} passed, {failed} failed (dilon-document-form-compiler / dilon-document-form-writer)")
    if failed == 0:
        print("\nAll form-compiler tests passed!")
        return 0
    print("\nSome form-compiler tests failed")
    return 1


if __name__ == "__main__":
    sys.exit(main())
