"""
Test suite for the dilon-document-extractor skill.

Direct-invocation style, matching tests/run_tests.py: a global
passed/failed counter via check(), explicit test calls from main(), no
pytest.
"""

import subprocess
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches

REPO_ROOT = Path(__file__).parent.parent
EXTRACTOR_DIR = REPO_ROOT / "skills" / "dilon-document-extractor"
SCRIPTS_DIR = EXTRACTOR_DIR / "scripts"
CHECK_DEPS_SCRIPT = SCRIPTS_DIR / "check_deps.py"
EXTRACT_DOCX_SCRIPT = SCRIPTS_DIR / "extract_docx.py"
EXTRACT_PDF_SCRIPT = SCRIPTS_DIR / "extract_pdf.py"
TEST_OUTPUT_DIR = Path(__file__).parent / "extractor-test-output"

SHEBANG_GUARDED_SCRIPTS = [
    CHECK_DEPS_SCRIPT,
    EXTRACT_DOCX_SCRIPT,
    EXTRACT_PDF_SCRIPT,
    Path(__file__),
]

sys.path.insert(0, str(SCRIPTS_DIR))

passed = 0
failed = 0


def check(condition, message):
    global passed, failed
    if condition:
        print(f"[PASS] {message}")
        passed += 1
    else:
        print(f"[FAIL] {message}")
        failed += 1


def test_check_deps_runs_and_reports():
    result = subprocess.run(
        [sys.executable, str(CHECK_DEPS_SCRIPT)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "check_deps.py exits 0 when python-docx/pyyaml/pymupdf are installed")
    check("docx" in result.stdout, "check_deps.py reports on the docx module")
    check("yaml" in result.stdout, "check_deps.py reports on the yaml module")
    check("fitz" in result.stdout, "check_deps.py reports on the fitz (pymupdf) module")


def _doc_with_paragraphs(style_texts):
    """style_texts: list of (style_name_or_None, text). Returns an
    in-memory docx.Document with one paragraph per entry."""
    doc = Document()
    for style_name, text in style_texts:
        p = doc.add_paragraph(text)
        if style_name:
            p.style = doc.styles[style_name]
    return doc


def test_word_heading_level():
    import extract_docx as ex
    check(ex.word_heading_level("Heading 1") == 1, "Heading 1 -> level 1")
    check(ex.word_heading_level("Heading 3") == 3, "Heading 3 -> level 3")
    check(ex.word_heading_level("Normal") is None, "Normal -> no heading level")
    check(ex.word_heading_level(None) is None, "None style -> no heading level")


def test_compute_heading_shift_shallowest_becomes_h2():
    import extract_docx as ex
    doc = _doc_with_paragraphs([
        ("Heading 1", "Top Section"),
        ("Heading 2", "Sub Section"),
        ("Normal", "Body text."),
    ])
    check(ex.compute_heading_shift(doc) == 1, "shallowest Word level 1 shifts by +1 (1->## is level 2)")


def test_compute_heading_shift_no_headings_defaults_to_two():
    import extract_docx as ex
    doc = _doc_with_paragraphs([("Normal", "Just a paragraph.")])
    check(ex.compute_heading_shift(doc) == 2, "no headings present -> shift defaults to 2")


def test_markdown_heading_prefix():
    import extract_docx as ex
    check(ex.markdown_heading_prefix(1, 1) == "##", "level 1 + shift 1 -> ##")
    check(ex.markdown_heading_prefix(2, 1) == "###", "level 2 + shift 1 -> ###")
    check(ex.markdown_heading_prefix(1, 0) == "##", "level 1 + shift 0 clamps to a minimum of ##")


def test_is_suspicious_heading_text():
    import extract_docx as ex
    check(
        ex.is_suspicious_heading_text("Strong pressure on the Photomultiplier should be avoided."),
        "long sentence ending in a period is flagged as a suspicious heading",
    )
    check(not ex.is_suspicious_heading_text("Bonding"), "short title-case heading is not flagged")


def test_heading_is_empty_leaf():
    import extract_docx as ex
    doc = Document()
    doc.add_paragraph("Next Step", style="Heading 2")
    doc.add_paragraph("Detector Head Testing following FTP-00001", style="Heading 3")
    blocks = list(ex.iter_block_items(doc))

    check(
        ex.heading_is_empty_leaf(blocks, 0, 2) is False,
        "a heading followed by a deeper child heading is not an empty leaf",
    )
    check(
        ex.heading_is_empty_leaf(blocks, 1, 3) is True,
        "sole child heading with no content, followed by end of document, is an empty leaf",
    )

    doc2 = Document()
    doc2.add_paragraph("Next Step", style="Heading 2")
    doc2.add_paragraph("Detector Head Testing following FTP-00001", style="Heading 3")
    doc2.add_paragraph("See the referenced procedure for full test steps.", style="Normal")
    blocks2 = list(ex.iter_block_items(doc2))
    check(
        ex.heading_is_empty_leaf(blocks2, 1, 3) is False,
        "a child heading followed by real body content is not an empty leaf",
    )

    doc3 = Document()
    doc3.add_paragraph("Section A", style="Heading 2")
    doc3.add_paragraph("Section B", style="Heading 2")
    blocks3 = list(ex.iter_block_items(doc3))
    check(
        ex.heading_is_empty_leaf(blocks3, 0, 2) is False,
        "a sibling heading (same level, not a parent/child pair) does not trigger the empty-leaf rule",
    )

    doc4 = Document()
    doc4.add_paragraph("Next Step", style="Heading 2")
    doc4.add_paragraph("Detector Head Testing following FTP-00001", style="Heading 3")
    image_path = TEST_OUTPUT_DIR / "_leaf_fixture_image.png"
    image_path.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753"
        "de0000000c49444154789c63f8ffff3f0005fe02fe0def46b8000000004945"
        "4e44ae426082"
    ))
    doc4.add_picture(str(image_path))
    blocks4 = list(ex.iter_block_items(doc4))
    check(
        ex.heading_is_empty_leaf(blocks4, 1, 3) is False,
        "a heading whose only content is an embedded image (no caption text) is not an empty leaf",
    )


def test_build_markdown_body_empty_leaf_heading_becomes_bullet():
    import extract_docx as ex
    doc = Document()
    doc.add_paragraph("Next Step", style="Heading 2")
    doc.add_paragraph("Detector Head Testing following FTP-00001", style="Heading 3")

    blocks = list(ex.iter_block_items(doc))
    front_matter = {"revisions": []}
    body, warnings = ex.build_markdown_body(doc, blocks, 1, TEST_OUTPUT_DIR, front_matter)

    check("### Next Step" in body, "genuine parent heading with a child is left as a heading")
    check("#### Detector Head Testing" not in body, "childless leaf heading is not rendered as a markdown heading")
    check("- Detector Head Testing following FTP-00001" in body, "childless leaf heading rendered as a bullet instead")


def test_titlecase_heading():
    import extract_docx as ex
    check(ex.titlecase_heading("RESPONSIBILITIES") == "Responsibilities", "all-caps heading title-cased")
    check(
        ex.titlecase_heading("carrier board assembly PROCEDURE") == "Carrier Board Assembly Procedure",
        "mixed-case heading title-cased",
    )
    check(
        ex.titlecase_heading("EQUIPMENT and SUPPLIES") == "Equipment And Supplies",
        "every word's first letter is capitalized, including short words",
    )
    check(ex.titlecase_heading("PN 820-00006") == "Pn 820-00006", "digits/hyphens left untouched")


def _add_table(doc, rows):
    """rows: list of list[str]. Returns the created docx.table.Table."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.rows[r].cells[c].text = text
    return table


def test_classify_table_signature():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [
        ["Group", "Preparer", "Signature"],
        ["Engineering", "P. Gray", "Electronic"],
        ["Department", "Name", "Signature"],
        ["Regulatory", "Pedro Cruz", "Electronic"],
        ["Quality", "Rebecca Miller", "Electronic"],
        ["Engineering", "Kevin Lint", "Electronic"],
    ])
    check(ex.classify_table(table) == "signature", "canonical signature-approval table is classified as 'signature'")


def test_classify_table_revision():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [
        ["REVISION HISTORY", "", "", ""],
        ["REV #", "DESCRIPTION OF CHANGE", "ECO #", "DATE"],
        ["00", "Initial release", "ECO-000046", "4 Mar 2025"],
    ])
    check(ex.classify_table(table) == "revision", "revision-history table is classified as 'revision'")


def test_classify_table_content():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [["Setting", "Value"], ["Pressure", "75 psi"]])
    check(ex.classify_table(table) == "content", "an ordinary data table is classified as 'content'")


def test_extract_signature_fields_clean_labels():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [
        ["Group", "Preparer", "Signature"],
        ["Engineering", "P. Gray", "Electronic"],
        ["Department", "Name", "Signature"],
        ["Engineering", "Kevin Lint", "Electronic"],
        ["Regulatory", "Pedro Cruz", "Electronic"],
        ["Quality", "Rebecca Miller", "Electronic"],
    ])
    fields, warnings = ex.extract_signature_fields(table)
    check(fields["author"] == "P. Gray", "author extracted from row 1")
    check(fields["department"] == "Engineering", "department extracted from row 1")
    check(fields["department_head"] == "Kevin Lint", "department_head extracted by position")
    check(
        fields["signature_fields"] == [
            {"department": "Regulatory", "name": "Pedro Cruz"},
            {"department": "Quality", "name": "Rebecca Miller"},
        ],
        f"remaining rows extracted as signature_fields entries by position, got {fields['signature_fields']}",
    )
    check(warnings == [], "no warnings when the department head row's label matches the top department")


def test_extract_signature_fields_mismatched_labels_warns():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [
        ["Group", "Preparer", "Signature"],
        ["Engineering", "P. Gray", "Electronic"],
        ["Department", "Name", "Signature"],
        ["R&D / Eng", "K. Lint", "Electronic"],
        ["Manufacturing", "J. Jones", "Electronic"],
        ["Quality", "K. Mack", "Electronic"],
    ])
    fields, warnings = ex.extract_signature_fields(table)
    check(fields["department_head"] == "K. Lint", "department_head still assigned by position despite label mismatch")
    check(len(warnings) == 1, f"a warning is emitted for the mismatched department head label, got {len(warnings)}")


def test_extract_revisions():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [
        ["REVISION HISTORY", "", "", ""],
        ["REV #", "DESCRIPTION OF CHANGE", "ECO #", "DATE"],
        ["00", "Original product transfer to Dilon Manufacturing", "ECO-000046", "4 Mar 2025"],
    ])
    revisions = ex.extract_revisions(table)
    check(len(revisions) == 1, "one revision row extracted")
    check(revisions[0]["number"] == "00", "revision number extracted")
    check(revisions[0]["eco_number"] == "ECO-000046", "eco_number extracted")


def test_extract_header_footer_metadata():
    import extract_docx as ex
    doc = Document()
    section = doc.sections[0]
    header_table = section.header.add_table(rows=2, cols=4, width=Inches(6))
    header_table.rows[0].cells[1].text = "WI:\nNav 3, Detector Head Assembly"
    header_table.rows[0].cells[2].text = "Rev 00"
    header_table.rows[1].cells[1].text = "Number:\nWI-00077"
    section.footer.paragraphs[0].text = "WI-00077 Rev 00\tECO-000046\tRevision Date: 03/4/2025"

    fields = ex.extract_header_footer_metadata(doc)
    check(fields.get("doc_number") == "WI-00077", f"doc_number parsed from header/footer, got {fields.get('doc_number')!r}")
    check(fields.get("current_revision") == "00", f"current_revision parsed, got {fields.get('current_revision')!r}")
    check(
        fields.get("title") == "Nav 3, Detector Head Assembly",
        f"title parsed from combined label+value header cell, got {fields.get('title')!r}",
    )
    check(fields.get("footer_eco_number") == "ECO-000046", f"footer_eco_number parsed, got {fields.get('footer_eco_number')!r}")
    check(fields.get("footer_eco_date") == "03/4/2025", f"footer_eco_date parsed, got {fields.get('footer_eco_date')!r}")


def test_extract_header_footer_metadata_table_footer():
    """Newer compiler output (see populate_footer() in
    lib/dilon_docx_common.py) carries the footer ID line as the first row
    of a 3-column table instead of tab-separated text in a single
    paragraph - table cell text isn't part of footer.paragraphs, so this
    must be read from section.footer.tables instead."""
    import extract_docx as ex
    doc = Document()
    section = doc.sections[0]
    footer_table = section.footer.add_table(rows=2, cols=3, width=Inches(6))
    footer_table.rows[0].cells[0].text = "WI-00077 Rev 00"
    footer_table.rows[0].cells[1].text = "ECO-000046"
    footer_table.rows[0].cells[2].text = "Revision Date: 03/4/2025"
    footer_table.rows[1].cells[0].text = "This document is confidential."

    fields = ex.extract_header_footer_metadata(doc)
    check(fields.get("doc_number") == "WI-00077", f"doc_number parsed from table-based footer, got {fields.get('doc_number')!r}")
    check(fields.get("current_revision") == "00", f"current_revision parsed from table-based footer, got {fields.get('current_revision')!r}")
    check(fields.get("footer_eco_number") == "ECO-000046", f"footer_eco_number parsed from table-based footer, got {fields.get('footer_eco_number')!r}")
    check(fields.get("footer_eco_date") == "03/4/2025", f"footer_eco_date parsed from table-based footer, got {fields.get('footer_eco_date')!r}")


def test_extract_header_footer_metadata_split_cells():
    """Real Dilon documents (e.g. WI-00077) split each header row's label
    and value into separate table cells rather than combining them with a
    newline in one cell."""
    import extract_docx as ex
    doc = Document()
    section = doc.sections[0]
    header_table = section.header.add_table(rows=2, cols=4, width=Inches(6))
    header_table.rows[0].cells[1].text = "WI:"
    header_table.rows[0].cells[2].text = "Nav 3, Detector Head Assembly\n PN 820-00006"
    header_table.rows[0].cells[3].text = "Rev 00"
    header_table.rows[1].cells[1].text = "Number:"
    header_table.rows[1].cells[2].text = "WI-00077"

    fields = ex.extract_header_footer_metadata(doc)
    check(
        fields.get("title") == "Nav 3, Detector Head Assembly PN 820-00006",
        f"title parsed across split label/value cells, got {fields.get('title')!r}",
    )
    check(fields.get("doc_number") == "WI-00077", f"doc_number still parsed when split across cells, got {fields.get('doc_number')!r}")


def test_extract_header_footer_metadata_prototype_revision():
    """Prototype revision numbers like "02-A" (major number + alphabetic
    prototype suffix) aren't digit-only, so the header/footer Rev
    patterns must capture the full value instead of stopping at the
    leading digits."""
    import extract_docx as ex
    doc = Document()
    section = doc.sections[0]
    header_table = section.header.add_table(rows=2, cols=4, width=Inches(6))
    header_table.rows[0].cells[1].text = "WI:\nNav 3, Detector Head Assembly"
    header_table.rows[0].cells[2].text = "Rev 02-A"
    header_table.rows[1].cells[1].text = "Number:\nWI-00077"
    section.footer.paragraphs[0].text = "WI-00077 Rev 02-A\tECO-000046\tRevision Date: 03/4/2025"

    fields = ex.extract_header_footer_metadata(doc)
    check(fields.get("current_revision") == "02-A", f"prototype current_revision parsed in full, got {fields.get('current_revision')!r}")
    check(fields.get("footer_eco_number") == "ECO-000046", f"footer_eco_number still parsed alongside a prototype revision, got {fields.get('footer_eco_number')!r}")
    check(fields.get("footer_eco_date") == "03/4/2025", f"footer_eco_date still parsed alongside a prototype revision, got {fields.get('footer_eco_date')!r}")


def test_strip_figure_prefix():
    import extract_docx as ex
    check(
        ex.strip_figure_prefix("Figure 1: Crystal Ends (Polished on Left)") == "Crystal Ends (Polished on Left)",
        "colon-style figure prefix stripped",
    )
    check(
        ex.strip_figure_prefix("Figure 12.3 - Compressor Gauges") == "Compressor Gauges",
        "dash-style figure prefix with a decimal number stripped",
    )
    check(ex.strip_figure_prefix("Just a caption") == "Just a caption", "text with no prefix is unchanged")


def _add_direct_numpr(paragraph, ilvl=0):
    """Attach direct (non-style-based) w:numPr list formatting to a
    paragraph, mirroring real Dilon documents where some list items are
    left 'Normal'-styled with manual list formatting instead of the 'List
    Paragraph' style."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.makeelement(qn('w:numPr'), {})
    ilvl_el = numPr.makeelement(qn('w:ilvl'), {})
    ilvl_el.set(qn('w:val'), str(ilvl))
    numId = numPr.makeelement(qn('w:numId'), {})
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl_el)
    numPr.append(numId)
    pPr.append(numPr)


def test_paragraph_is_list_item_direct_numpr():
    import extract_docx as ex
    doc = Document()
    p1 = doc.add_paragraph("Styled as List Paragraph")
    p1.style = doc.styles["List Paragraph"]
    p2 = doc.add_paragraph("Normal style but has direct numPr")
    _add_direct_numpr(p2)
    p3 = doc.add_paragraph("Plain Normal paragraph")

    check(ex.paragraph_is_list_item(p1) is True, "'List Paragraph'-styled paragraph detected as a list item")
    check(ex.paragraph_is_list_item(p2) is True, "Normal-styled paragraph with direct numPr formatting detected as a list item")
    check(ex.paragraph_is_list_item(p3) is False, "plain Normal paragraph is not a list item")


def test_paragraph_list_ilvl():
    import extract_docx as ex
    doc = Document()
    p0 = doc.add_paragraph("Top-level item")
    p0.style = doc.styles["List Paragraph"]
    _add_direct_numpr(p0, ilvl=0)
    p1 = doc.add_paragraph("Nested item")
    p1.style = doc.styles["List Paragraph"]
    _add_direct_numpr(p1, ilvl=1)
    p2 = doc.add_paragraph("No list formatting at all")

    check(ex.paragraph_list_ilvl(p0) == 0, "top-level list item has ilvl 0")
    check(ex.paragraph_list_ilvl(p1) == 1, "nested list item has ilvl 1")
    check(ex.paragraph_list_ilvl(p2) == 0, "non-list paragraph defaults to ilvl 0")


def test_build_markdown_body_nested_list_indentation():
    import extract_docx as ex
    doc = Document()
    doc.add_paragraph("Mixing Epoxy", style="Heading 2")
    top = doc.add_paragraph("The epoxy is purchased in a pre-measured packet.")
    top.style = doc.styles["List Paragraph"]
    _add_direct_numpr(top, ilvl=0)
    nested = doc.add_paragraph("Read the outer package to confirm the expiration date.")
    nested.style = doc.styles["List Paragraph"]
    _add_direct_numpr(nested, ilvl=1)

    blocks = list(ex.iter_block_items(doc))
    front_matter = {"revisions": []}
    body, warnings = ex.build_markdown_body(doc, blocks, 1, TEST_OUTPUT_DIR, front_matter)

    check("- The epoxy is purchased in a pre-measured packet." in body, "top-level (ilvl 0) list item rendered with no indent")
    check(
        "  - Read the outer package to confirm the expiration date." in body,
        "nested (ilvl 1) list item rendered with a 2-space indent",
    )


def test_is_toc_paragraph():
    import extract_docx as ex
    check(ex.is_toc_paragraph("toc 1", "1.\tIntroduction\t3"), "'toc 1'-styled paragraph is a TOC entry")
    check(ex.is_toc_paragraph("toc 2", "1.1\tScope\t3"), "'toc 2'-styled paragraph is a TOC entry")
    check(ex.is_toc_paragraph("Normal", "TABLE OF CONTENTS"), "literal 'TABLE OF CONTENTS' text is a TOC heading")
    check(not ex.is_toc_paragraph("Normal", "Regular body text."), "ordinary body text is not a TOC paragraph")
    check(not ex.is_toc_paragraph("Heading 1", "Introduction"), "a real heading is not a TOC paragraph")


def test_build_markdown_body_skips_toc_and_converts_direct_numpr_list():
    import extract_docx as ex
    doc = Document()
    doc.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("TABLE OF CONTENTS", style="Normal")
    doc.add_paragraph("1.\tIntroduction\t3", style="toc 1")
    doc.add_paragraph("Introduction", style="Heading 1")
    p = doc.add_paragraph("Providing support as necessary.", style="Normal")
    _add_direct_numpr(p)

    blocks = list(ex.iter_block_items(doc))
    front_matter = {"revisions": []}
    body, warnings = ex.build_markdown_body(doc, blocks, 1, TEST_OUTPUT_DIR, front_matter)

    check("TABLE OF CONTENTS" not in body, "literal TOC heading text excluded from body")
    check("Introduction\t3" not in body, "toc-styled entry paragraph excluded from body")
    check("## Introduction" in body, "real heading still present after TOC removal")
    check("- Providing support as necessary." in body, "Normal-styled paragraph with direct numPr rendered as a markdown bullet")


def test_extract_flags_footer_revision_eco_mismatch():
    import extract_docx as ex
    doc = Document()
    section = doc.sections[0]
    header_table = section.header.add_table(rows=2, cols=4, width=Inches(6))
    header_table.rows[0].cells[1].text = "WI:\nMismatch Fixture"
    header_table.rows[0].cells[2].text = "Rev 00"
    header_table.rows[1].cells[1].text = "Number:\nWI-88888"
    section.footer.paragraphs[0].text = "WI-88888 Rev 00\tECO-000123\tRevision Date: 5/5/2025"

    _add_table(doc, [
        ["REVISION HISTORY", "", "", ""],
        ["REV #", "DESCRIPTION OF CHANGE", "ECO #", "DATE"],
        ["00", "Initial release", "ECO-000999", "5 May 2025"],
    ])

    fixture_path = TEST_OUTPUT_DIR / "footer-mismatch-fixture.docx"
    output_dir = TEST_OUTPUT_DIR / "footer-mismatch-extracted"
    doc.save(str(fixture_path))

    result = ex.extract(fixture_path, output_dir)
    content = result["markdown_path"].read_text(encoding="utf-8")
    front_matter = yaml.safe_load(content.split("---\n")[1])

    check("footer_eco_number" not in front_matter, "footer_eco_number not leaked into final front matter")
    check("footer_eco_date" not in front_matter, "footer_eco_date not leaked into final front matter")
    check(
        any("disagreement" in w.lower() for w in result["warnings"]),
        "ECO mismatch between footer and revision table produces a disagreement warning",
    )


def test_extract_no_warning_when_footer_matches_revision():
    import extract_docx as ex
    doc = Document()
    section = doc.sections[0]
    header_table = section.header.add_table(rows=2, cols=4, width=Inches(6))
    header_table.rows[0].cells[1].text = "WI:\nMatch Fixture"
    header_table.rows[0].cells[2].text = "Rev 00"
    header_table.rows[1].cells[1].text = "Number:\nWI-77777"
    section.footer.paragraphs[0].text = "WI-77777 Rev 00\tECO-000555\tRevision Date: 6/6/2025"

    _add_table(doc, [
        ["REVISION HISTORY", "", "", ""],
        ["REV #", "DESCRIPTION OF CHANGE", "ECO #", "DATE"],
        ["00", "Initial release", "ECO-000555", "6/6/2025"],
    ])

    fixture_path = TEST_OUTPUT_DIR / "footer-match-fixture.docx"
    output_dir = TEST_OUTPUT_DIR / "footer-match-extracted"
    doc.save(str(fixture_path))

    result = ex.extract(fixture_path, output_dir)
    check(
        not any("disagreement" in w.lower() for w in result["warnings"]),
        "matching footer/revision-table ECO data produces no disagreement warning",
    )


def test_build_markdown_body_suspicious_heading_becomes_step():
    import extract_docx as ex
    doc = Document()
    doc.add_paragraph("Preparation", style="Heading 1")
    doc.add_paragraph("Simple dirt can be blown away.", style="List Paragraph")
    doc.add_paragraph("Strong pressure on the Photomultiplier should be avoided.", style="Heading 4")

    blocks = list(ex.iter_block_items(doc))
    front_matter = {"revisions": []}
    body, warnings = ex.build_markdown_body(doc, blocks, 1, TEST_OUTPUT_DIR, front_matter)

    check("##### Strong pressure" not in body, "suspicious heading-styled paragraph not rendered as a markdown heading")
    check("@@@STEPS@@@" in body and "@@@END_STEPS@@@" in body, "suspicious heading-styled paragraph wrapped in a @@@STEPS@@@ block")
    check(
        "#. Strong pressure on the Photomultiplier should be avoided." in body,
        "suspicious heading-styled paragraph rendered as a step item instead",
    )
    check(any("rendered as a @@@STEPS@@@ item" in w for w in warnings), "conversion from heading to step item is flagged for review")


def test_build_markdown_body_step_run_nests_by_heading_level():
    import extract_docx as ex
    doc = Document()
    doc.add_paragraph("Epoxy", style="Heading 2")
    doc.add_paragraph("Mix the epoxy per the manufacturer's specs.", style="Heading 3")
    doc.add_paragraph("Use the pink needle at 10 psi.", style="Heading 4")
    doc.add_paragraph("Dispense a small amount of epoxy into a tray.", style="Heading 3")

    blocks = list(ex.iter_block_items(doc))
    front_matter = {"revisions": []}
    body, warnings = ex.build_markdown_body(doc, blocks, 1, TEST_OUTPUT_DIR, front_matter)

    check(body.count("@@@STEPS@@@") == 1, "one contiguous step run produces a single @@@STEPS@@@ block")
    check("#. Mix the epoxy per the manufacturer's specs." in body, "first sibling step rendered at top level")
    check("  #. Use the pink needle at 10 psi." in body, "deeper heading level rendered as a nested clarification")
    check("#. Dispense a small amount of epoxy into a tray." in body, "second sibling step rendered at top level")
    check(
        not any(line.startswith("  #. Dispense") for line in body.splitlines()),
        "sibling step at the run's base level is not nested under the prior step",
    )


def test_build_markdown_body_step_run_closes_around_image():
    import extract_docx as ex
    doc = Document()
    doc.add_paragraph("Epoxy", style="Heading 2")
    doc.add_paragraph("Mix the epoxy per the manufacturer's specs.", style="Heading 3")
    doc.add_paragraph("A caption-less plain paragraph interrupts the run.")
    doc.add_paragraph("Dispense a small amount of epoxy into a tray.", style="Heading 3")

    blocks = list(ex.iter_block_items(doc))
    front_matter = {"revisions": []}
    body, warnings = ex.build_markdown_body(doc, blocks, 1, TEST_OUTPUT_DIR, front_matter)

    check(body.count("@@@STEPS@@@") == 2, "an interruption closes and reopens a fresh @@@STEPS@@@ block")
    check(body.count("@@@END_STEPS@@@") == 2, "each opened @@@STEPS@@@ block is closed")


def test_slugify_dedup():
    import extract_docx as ex
    seen = set()
    check(ex.slugify("Crystal Ring", seen) == "crystal-ring", "basic slugify")
    check(ex.slugify("Crystal Ring", seen) == "crystal-ring-2", "second collision gets a -2 suffix")
    check(ex.slugify("Crystal Ring", seen) == "crystal-ring-3", "third collision gets a -3 suffix")


def _build_fixture_docx(path):
    """Builds a small synthetic .docx exercising every extraction path at
    once: header/footer metadata, a signature table with one mismatched
    role label, a revision table, two heading levels, an inline image with
    an adjacent Caption paragraph, a List Paragraph run, and a plain
    content table. Mirrors WI-00077's structure at a scale small enough to
    hand-verify in a test."""
    doc = Document()
    section = doc.sections[0]
    header_table = section.header.add_table(rows=2, cols=3, width=Inches(6))
    header_table.rows[0].cells[1].text = "WI:\nFixture Document"
    header_table.rows[0].cells[2].text = "Rev 00"
    header_table.rows[1].cells[1].text = "Number:\nWI-99999"
    section.footer.paragraphs[0].text = "WI-99999 Rev 00\tECO-000099\tRevision Date: 1/1/2026"

    doc.add_paragraph("Preparation", style="Heading 1")
    doc.add_paragraph("Clean the parts before assembly.", style="Normal")
    p1 = doc.add_paragraph("Wear clean gloves.", style="List Paragraph")
    p2 = doc.add_paragraph("Blow away loose dust.", style="List Paragraph")

    image_path = Path(__file__).parent / "extractor-test-output" / "_fixture_image.png"
    image_path.parent.mkdir(parents=True, exist_ok=True)
    # 1x1 transparent PNG, smallest valid image payload
    image_path.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753"
        "de0000000c49444154789c63f8ffff3f0005fe02fe0def46b8000000004945"
        "4e44ae426082"
    ))
    doc.add_picture(str(image_path))
    doc.add_paragraph("Figure 1: Crystal Ends (Polished on Left)", style="Caption")

    doc.add_paragraph("Bonding", style="Heading 1")
    doc.add_paragraph("Apply epoxy to the crystal.", style="Normal")

    _add_table(doc, [["Setting", "Value"], ["Pressure", "75 psi"]])

    _add_table(doc, [
        ["Group", "Preparer", "Signature"],
        ["Engineering", "P. Gray", "Electronic"],
        ["Department", "Name", "Signature"],
        ["R&D / Eng", "K. Lint", "Electronic"],
        ["Manufacturing", "J. Jones", "Electronic"],
        ["Quality", "K. Mack", "Electronic"],
    ])

    _add_table(doc, [
        ["REVISION HISTORY", "", "", ""],
        ["REV #", "DESCRIPTION OF CHANGE", "ECO #", "DATE"],
        ["00", "Initial release", "ECO-000099", "1 Jan 2026"],
    ])

    doc.save(str(path))


def test_extract_full_fixture():
    import extract_docx as ex
    fixture_path = TEST_OUTPUT_DIR / "fixture.docx"
    output_dir = TEST_OUTPUT_DIR / "fixture-extracted"
    _build_fixture_docx(fixture_path)

    result = ex.extract(fixture_path, output_dir)

    check(result["markdown_path"].exists(), "extract() writes a markdown file")
    content = result["markdown_path"].read_text(encoding="utf-8")

    front_matter_text = content.split("---\n")[1]
    front_matter = yaml.safe_load(front_matter_text)

    check(front_matter["doc_number"] == "WI-99999", f"doc_number in front matter, got {front_matter.get('doc_number')!r}")
    check(front_matter["author"] == "P. Gray", f"author in front matter, got {front_matter.get('author')!r}")
    check(front_matter["department_head"] == "K. Lint", f"department_head assigned by position, got {front_matter.get('department_head')!r}")
    check(
        front_matter["signature_fields"] == [
            {"department": "Manufacturing", "name": "J. Jones"},
            {"department": "Quality", "name": "K. Mack"},
        ],
        f"signature_fields extracted by position, got {front_matter.get('signature_fields')!r}",
    )
    check(front_matter["revisions"][0]["eco_number"] == "ECO-000099", "revisions list populated from body table")

    check("Group | Preparer | Signature" not in content, "signature table text excluded from body")
    check("REVISION HISTORY" not in content, "revision table text excluded from body")
    check("## Preparation" in content, "Heading 1 shifted to markdown H2")
    check("## Bonding" in content, "second Heading 1 also shifted to markdown H2")
    check("- Wear clean gloves." in content, "List Paragraph converted to a markdown bullet")
    check("![Crystal Ends (Polished on Left)]" in content, "figure prefix stripped, remaining caption used as alt text")
    check("{#fig:crystal-ends-polished-on-left}" in content, "figure gets a slugified id")
    check("<!-- EXTRACTOR:" in content, "at least one review comment present (mismatched signature-role labels)")

    images = list(result["images_dir"].glob("*"))
    check(len(images) == 1, f"one image extracted, got {len(images)}")


def test_table_to_markdown_pipe():
    import extract_docx as ex
    doc = Document()
    table = _add_table(doc, [["A", "B"], ["1", "2"]])
    md = ex.table_to_markdown(table)
    check(md.splitlines()[0] == "| A | B |", "pipe table header row rendered")
    check("---" in md.splitlines()[1], "pipe table separator row rendered")


def _build_fixture_pdf(path):
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Fixture PDF Document")
    page.insert_text((72, 100), "This is a body paragraph in the fixture PDF.")
    doc.save(str(path))
    doc.close()


def test_extract_pdf_banner_and_text():
    import extract_pdf as expdf
    fixture_path = TEST_OUTPUT_DIR / "fixture.pdf"
    output_dir = TEST_OUTPUT_DIR / "fixture-pdf-extracted"
    _build_fixture_pdf(fixture_path)

    result = expdf.extract(fixture_path, output_dir)

    check(result["markdown_path"].exists(), "extract_pdf.extract() writes a markdown file")
    content = result["markdown_path"].read_text(encoding="utf-8")
    check("PDF source" in content, "PDF banner comment present at the top of the draft")
    check("Fixture PDF Document" in content, "extracted text includes page content")
    check(any("PDF source" in w for w in result["warnings"]), "the banner warning is also reported in the returned warnings list")


def test_extract_docx_cli_smoke():
    fixture_path = TEST_OUTPUT_DIR / "cli-fixture.docx"
    output_dir = TEST_OUTPUT_DIR / "cli-extracted"
    _build_fixture_docx(fixture_path)

    result = subprocess.run(
        [sys.executable, str(EXTRACT_DOCX_SCRIPT), str(fixture_path), str(output_dir)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "extract_docx.py CLI exits 0 for a valid input")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(any(output_dir.glob("*.md")), "extract_docx.py CLI writes a markdown file to the output directory")


def test_no_shebang_in_extractor_scripts():
    def has_shebang(path):
        lines = path.read_text(encoding="utf-8").splitlines()
        return bool(lines) and lines[0].startswith("#!")

    offenders = [str(p) for p in SHEBANG_GUARDED_SCRIPTS if has_shebang(p)]
    check(not offenders, f"no shebang lines in guarded extractor scripts (offenders: {offenders})")


def main():
    if TEST_OUTPUT_DIR.exists():
        import shutil
        shutil.rmtree(TEST_OUTPUT_DIR)
    TEST_OUTPUT_DIR.mkdir(parents=True)

    test_check_deps_runs_and_reports()
    test_word_heading_level()
    test_compute_heading_shift_shallowest_becomes_h2()
    test_compute_heading_shift_no_headings_defaults_to_two()
    test_markdown_heading_prefix()
    test_is_suspicious_heading_text()
    test_heading_is_empty_leaf()
    test_build_markdown_body_empty_leaf_heading_becomes_bullet()
    test_titlecase_heading()
    test_classify_table_signature()
    test_classify_table_revision()
    test_classify_table_content()
    test_extract_signature_fields_clean_labels()
    test_extract_signature_fields_mismatched_labels_warns()
    test_extract_revisions()
    test_extract_header_footer_metadata()
    test_extract_header_footer_metadata_table_footer()
    test_extract_header_footer_metadata_split_cells()
    test_strip_figure_prefix()
    test_paragraph_is_list_item_direct_numpr()
    test_paragraph_list_ilvl()
    test_build_markdown_body_nested_list_indentation()
    test_is_toc_paragraph()
    test_build_markdown_body_skips_toc_and_converts_direct_numpr_list()
    test_build_markdown_body_suspicious_heading_becomes_step()
    test_build_markdown_body_step_run_nests_by_heading_level()
    test_build_markdown_body_step_run_closes_around_image()
    test_extract_flags_footer_revision_eco_mismatch()
    test_extract_no_warning_when_footer_matches_revision()
    test_extract_header_footer_metadata_prototype_revision()
    test_slugify_dedup()
    test_extract_full_fixture()
    test_table_to_markdown_pipe()
    test_extract_pdf_banner_and_text()
    test_extract_docx_cli_smoke()
    test_no_shebang_in_extractor_scripts()

    print(f"\n{passed} passed, {failed} failed (dilon-document-extractor)")
    if failed == 0:
        print("\nAll extractor tests passed!")
        return 0
    print("\nSome extractor tests failed")
    return 1


if __name__ == "__main__":
    sys.exit(main())
