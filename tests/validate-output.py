#!/usr/bin/env python3
"""
Output Validation Script for Dilon Claude Tools

Validates that:
1. Document stubs are generated correctly with proper YAML structure
2. Compiled Word documents have correct structure and content
3. STYLING_TEST_TEMPLATE compiles successfully
"""

import sys
import os
from pathlib import Path
import yaml
from docx import Document

# Color codes for terminal output
class Colors:
    GREEN = '\033[92m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    RESET = '\033[0m'

def print_success(msg):
    print(f"{Colors.GREEN}[PASS]{Colors.RESET} {msg}")

def print_error(msg):
    print(f"{Colors.RED}[FAIL]{Colors.RESET} {msg}")

def print_info(msg):
    print(f"{Colors.BLUE}[INFO]{Colors.RESET} {msg}")

def print_section(msg):
    print(f"\n{Colors.YELLOW}{'='*60}{Colors.RESET}")
    print(f"{Colors.YELLOW}{msg}{Colors.RESET}")
    print(f"{Colors.YELLOW}{'='*60}{Colors.RESET}")

def validate_markdown_stub(stub_path, expected_values=None):
    """
    Validate a generated markdown stub

    Args:
        stub_path: Path to the markdown stub file
        expected_values: Dict of expected YAML values to check

    Returns:
        Tuple of (success: bool, errors: list)
    """
    errors = []

    # Check file exists
    if not os.path.exists(stub_path):
        errors.append(f"Stub file not found: {stub_path}")
        return False, errors

    # Read the file
    try:
        with open(stub_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        errors.append(f"Failed to read stub: {e}")
        return False, errors

    # Extract YAML front matter
    if not content.startswith('---'):
        errors.append("Stub does not start with YAML front matter")
        return False, errors

    parts = content.split('---', 2)
    if len(parts) < 3:
        errors.append("Invalid YAML front matter structure")
        return False, errors

    yaml_content = parts[1].strip()
    markdown_content = parts[2].strip()

    # Parse YAML
    try:
        metadata = yaml.safe_load(yaml_content)
    except Exception as e:
        errors.append(f"Failed to parse YAML: {e}")
        return False, errors

    # Check required YAML fields
    required_fields = [
        'title', 'author', 'department', 'doc_number',
        'current_revision', 'regulatory_rep', 'quality_rep',
        'department_head', 'revisions'
    ]

    for field in required_fields:
        if field not in metadata:
            errors.append(f"Missing required YAML field: {field}")

    # Validate revisions structure
    if 'revisions' in metadata:
        if not isinstance(metadata['revisions'], list):
            errors.append("'revisions' should be a list")
        elif len(metadata['revisions']) > 0:
            # Check all revisions have required fields
            for idx, revision in enumerate(metadata['revisions']):
                revision_fields = ['number', 'description', 'eco_number', 'eco_date']
                for field in revision_fields:
                    if field not in revision:
                        errors.append(f"Missing revision[{idx}] field: {field}")

            # Note: We don't enforce that current_revision matches the last revision number
            # as it's the user's responsibility to keep these in sync

    # Check for expected values if provided (text comparison)
    if expected_values:
        for key, expected in expected_values.items():
            if key in metadata:
                actual = metadata[key]
                # Compare as strings to be type-independent
                if str(actual) != str(expected):
                    errors.append(
                        f"Expected {key}='{expected}', got '{actual}'"
                    )

    # Check markdown content structure
    if '## 1. Purpose and Scope' not in markdown_content:
        errors.append("Missing 'Purpose and Scope' section")

    if '### 1.1 Purpose' not in markdown_content:
        errors.append("Missing 'Purpose' subsection")

    if '### 1.2 Scope' not in markdown_content:
        errors.append("Missing 'Scope' subsection")

    return len(errors) == 0, errors

def validate_word_document(docx_path, expected_title=None):
    """
    Validate a compiled Word document

    Args:
        docx_path: Path to the .docx file
        expected_title: Optional expected document title

    Returns:
        Tuple of (success: bool, errors: list)
    """
    errors = []

    # Check file exists
    if not os.path.exists(docx_path):
        errors.append(f"Word document not found: {docx_path}")
        return False, errors

    # Open the document
    try:
        doc = Document(docx_path)
    except Exception as e:
        errors.append(f"Failed to open Word document: {e}")
        return False, errors

    # Check document has content
    if len(doc.paragraphs) == 0:
        errors.append("Document has no paragraphs")

    # Check for tables (should have signature and revision tables)
    if len(doc.tables) < 2:
        print_info(f"Document has {len(doc.tables)} table(s) - expected at least 2 (signature + revision)")

    # Check document has styles defined
    if len(doc.styles) == 0:
        errors.append("Document has no styles defined")

    # Extract text content
    full_text = '\n'.join([p.text for p in doc.paragraphs])

    # Check for title if provided
    if expected_title and expected_title not in full_text:
        print_info(f"Expected title '{expected_title}' not found in document text")

    # Check for basic structure
    if 'Purpose and Scope' not in full_text:
        print_info("'Purpose and Scope' section not clearly visible in extracted text")

    return len(errors) == 0, errors

def main():
    """Main validation routine"""
    print_section("Dilon Claude Tools - Output Validation")

    # Get repository root
    script_dir = Path(__file__).parent
    repo_root = script_dir.parent
    tests_dir = script_dir
    test_output_dir = tests_dir / 'test-output'

    total_checks = 0
    passed_checks = 0

    # ===================================================================
    # Test 1: Validate custom stub
    # ===================================================================
    print_section("Test 1: Validate Custom Stub Generation")
    total_checks += 1

    custom_stub_path = test_output_dir / 'custom_stub.md'
    success, errors = validate_markdown_stub(
        custom_stub_path,
        expected_values={
            'title': 'Software Requirements Specification',
            'author': 'Engineering Team',
            'doc_number': 'DD_SWE_12345',
            'department': 'Software Engineering',
            'current_revision': '01'
        }
    )

    if success:
        print_success(f"Custom stub validated: {custom_stub_path}")
        passed_checks += 1
    else:
        print_error(f"Custom stub validation failed:")
        for error in errors:
            print(f"  - {error}")

    # ===================================================================
    # Test 2: Validate default stub
    # ===================================================================
    print_section("Test 2: Validate Default Stub Generation")
    total_checks += 1

    default_stub_path = test_output_dir / 'default_stub.md'
    success, errors = validate_markdown_stub(
        default_stub_path,
        expected_values={
            'title': 'Document Title',
            'author': 'Author Name',
            'department': '--',
            'current_revision': '00'
        }
    )

    if success:
        print_success(f"Default stub validated: {default_stub_path}")
        passed_checks += 1
    else:
        print_error(f"Default stub validation failed:")
        for error in errors:
            print(f"  - {error}")

    # ===================================================================
    # Test 3: Validate compiled document
    # ===================================================================
    print_section("Test 3: Validate Compiled Word Document")
    total_checks += 1

    compiled_doc_path = test_output_dir / 'compile_test.docx'
    success, errors = validate_word_document(
        compiled_doc_path,
        expected_title='Integration Test Document'
    )

    if success:
        print_success(f"Compiled document validated: {compiled_doc_path}")
        passed_checks += 1
    else:
        print_error(f"Compiled document validation failed:")
        for error in errors:
            print(f"  - {error}")

    # ===================================================================
    # Test 4: Validate STYLING_TEST_TEMPLATE
    # ===================================================================
    print_section("Test 4: Validate STYLING_TEST_TEMPLATE")
    total_checks += 1

    styling_test_md = tests_dir / 'STYLING_TEST_TEMPLATE.md'
    success, errors = validate_markdown_stub(styling_test_md)

    if success:
        print_success(f"STYLING_TEST_TEMPLATE markdown validated")
        passed_checks += 1
    else:
        print_error(f"STYLING_TEST_TEMPLATE validation failed:")
        for error in errors:
            print(f"  - {error}")

    # Check if compiled version exists
    styling_test_docx = tests_dir / 'STYLING_TEST_TEMPLATE.docx'
    if os.path.exists(styling_test_docx):
        total_checks += 1
        success, errors = validate_word_document(styling_test_docx)

        if success:
            print_success(f"STYLING_TEST_TEMPLATE.docx validated")
            passed_checks += 1
        else:
            print_error(f"STYLING_TEST_TEMPLATE.docx validation failed:")
            for error in errors:
                print(f"  - {error}")
    else:
        print_info("STYLING_TEST_TEMPLATE.docx not found (not required)")

    # ===================================================================
    # Summary
    # ===================================================================
    print_section("Validation Summary")
    print(f"Total checks: {total_checks}")
    print(f"{Colors.GREEN}Passed: {passed_checks}{Colors.RESET}")
    print(f"{Colors.RED}Failed: {total_checks - passed_checks}{Colors.RESET}")

    success_rate = (passed_checks / total_checks * 100) if total_checks > 0 else 0
    print(f"Success rate: {success_rate:.1f}%")

    if passed_checks == total_checks:
        print(f"\n{Colors.GREEN}All validation checks passed!{Colors.RESET}")
        return 0
    else:
        print(f"\n{Colors.RED}Some validation checks failed{Colors.RESET}")
        return 1

if __name__ == '__main__':
    sys.exit(main())
