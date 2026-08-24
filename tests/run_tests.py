"""
Test suite for Dilon Claude Tools skills.

Exercises the dilon-document-writer stub-generation logic (ported here in
Python, since the live version is plain Claude behavior described in
SKILL.md, not a script) and the dilon-document-compiler script directly,
then runs the existing output validator.
"""

import re
import shutil
import struct
import subprocess
import sys
import zipfile
import zlib
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

sys.path.insert(0, str(Path(__file__).parent.parent / "skills" / "dilon-document-compiler" / "scripts"))
import generate_dilon_doc as compiler
import step_numbering
import dilon_docx_common
from dilon_docx_common import (
    validate_list_nesting_depth,
    ListNestingError,
    remap_ordered_lists_to_dilon_step_list,
    ensure_blank_line_after_list_continue_markers,
    resolve_list_continuations,
    ListContinuationError,
    _paragraph_num_id_and_ilvl,
)

REPO_ROOT = Path(__file__).parent.parent
TEST_OUTPUT_DIR = Path(__file__).parent / "test-output"
WRITER_DIR = REPO_ROOT / "skills" / "dilon-document-writer"
COMPILER_DIR = REPO_ROOT / "skills" / "dilon-document-compiler"
TEMPLATE_PATH = WRITER_DIR / "TEMPLATE_Document.md"
COMPILER_SCRIPT = COMPILER_DIR / "scripts" / "generate_dilon_doc.py"
CHECK_DEPS_SCRIPT = COMPILER_DIR / "scripts" / "check_deps.py"
SIGNATURE_TEMPLATE = REPO_ROOT / "templates" / "TEMPLATE_Word_Base.docx"

# Scripts that must not carry a `#!/usr/bin/env python3` shebang: Windows'
# py launcher parses that line and can re-dispatch to a different,
# dependency-less python3.exe (e.g. the Microsoft Store WindowsApps stub)
# instead of the real interpreter. This project is Windows-only.
SHEBANG_GUARDED_SCRIPTS = [
    COMPILER_SCRIPT,
    CHECK_DEPS_SCRIPT,
    Path(__file__).parent / "run_tests.py",
    Path(__file__).parent / "validate-output.py",
]

SAMPLE_MARKDOWN = (
    '---\n'
    'title: "Integration Test Document"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'doc_number: "DD_TST_99999"\n'
    'current_revision: "00"\n'
    'regulatory_rep: "Test Rep"\n'
    'quality_rep: "Test QA"\n'
    'department_head: "Test Head"\n'
    'revisions:\n'
    '  - number: "00"\n'
    '    description: "Initial test"\n'
    '    eco_number: "ECO-000"\n'
    '    eco_date: "2025-01-01"\n'
    '---\n'
    '\n'
    '## 1. Purpose and Scope\n'
    '\n'
    '### 1.1 Purpose\n'
    'This document tests the compilation process.\n'
    '\n'
    '### 1.2 Scope\n'
    'Comprehensive integration testing.\n'
)

passed = 0
failed = 0


def check(condition, message):
    global passed, failed
    if condition:
        print(f"[PASS] {message}")
        passed += 1
    else:
        print(f"[FAIL] {message}")
        failed += 1


def generate_stub(output_path, **overrides):
    """Python port of the dilon-document-writer stub-generation logic
    (lives in SKILL.md as plain instructions for Claude; ported here so
    it can be exercised by an automated test)."""
    output_path = Path(output_path)
    if output_path.exists():
        raise FileExistsError(f"Output file already exists: {output_path}")

    template = TEMPLATE_PATH.read_text(encoding="utf-8")

    values = {
        "title": overrides.get("title", "Document Title"),
        "author": overrides.get("author", "Author Name"),
        "department": overrides.get("department", "--"),
        "doc_number": overrides.get("doc_number", "DD_XXX_XXXXX"),
        "current_revision": overrides.get("current_revision", "00"),
        "regulatory_rep": overrides.get("regulatory_rep", "--"),
        "quality_rep": overrides.get("quality_rep", "--"),
        "department_head": overrides.get("department_head", "--"),
        "revision_description": overrides.get("revision_description", "Initial release"),
        "eco_number": overrides.get("eco_number", "ECO-TBD"),
        "eco_date": overrides.get("eco_date", "YYYY-MM-DD"),
    }

    content = template
    content = re.sub(r'title: ".*?"', f'title: "{values["title"]}"', content, count=1)
    content = re.sub(r'author: ".*?"', f'author: "{values["author"]}"', content, count=1)
    content = re.sub(r'department: ".*?"', f'department: "{values["department"]}"', content, count=1)
    content = re.sub(r'doc_number: ".*?"', f'doc_number: "{values["doc_number"]}"', content, count=1)
    content = re.sub(r'current_revision: ".*?"', f'current_revision: "{values["current_revision"]}"', content, count=1)
    content = re.sub(r'regulatory_rep: ".*?"', f'regulatory_rep: "{values["regulatory_rep"]}"', content, count=1)
    content = re.sub(r'quality_rep: ".*?"', f'quality_rep: "{values["quality_rep"]}"', content, count=1)
    content = re.sub(r'department_head: ".*?"', f'department_head: "{values["department_head"]}"', content, count=1)
    content = re.sub(
        r'- number: ".*?"\s+description: ".*?"\s+eco_number: ".*?"\s+eco_date: ".*?"',
        '- number: "{0}"\n    description: "{1}"\n    eco_number: "{2}"\n    eco_date: "{3}"'.format(
            values["current_revision"], values["revision_description"], values["eco_number"], values["eco_date"]
        ),
        content,
        count=1,
    )

    output_path.write_text(content, encoding="utf-8")


def test_stub_custom_params():
    path = TEST_OUTPUT_DIR / "custom_stub.md"
    generate_stub(
        path,
        title="Software Requirements Specification",
        author="Engineering Team",
        doc_number="DD_SWE_12345",
        department="Software Engineering",
        current_revision="01",
    )
    check(path.exists(), "custom_stub.md created")


def test_stub_default_params():
    path = TEST_OUTPUT_DIR / "default_stub.md"
    generate_stub(path)
    check(path.exists(), "default_stub.md created")


def test_stub_duplicate_file_error():
    path = TEST_OUTPUT_DIR / "custom_stub.md"  # already created above
    raised = False
    try:
        generate_stub(path)
    except FileExistsError:
        raised = True
    check(raised, "generate_stub refuses to overwrite an existing file")


def test_ensure_blank_line_single_marker():
    result = compiler.ensure_blank_line_after_table_markers(
        '@@@TABLE_STYLE:DilonTable_Chart@@@\n| a |\n'
    )
    check(result == '@@@TABLE_STYLE:DilonTable_Chart@@@\n\n| a |\n',
          "single TABLE_STYLE marker gets a blank line inserted before the table")


def test_ensure_blank_line_stacked_style_then_columns():
    result = compiler.ensure_blank_line_after_table_markers(
        '@@@TABLE_STYLE:DilonTable_Chart@@@\n@@@TABLE_COLUMNS:1,x@@@\n| a |\n'
    )
    check(result == '@@@TABLE_STYLE:DilonTable_Chart@@@\n@@@TABLE_COLUMNS:1,x@@@\n\n| a |\n',
          "stacked STYLE+COLUMNS markers get exactly one blank line inserted after the last marker line")


def test_ensure_blank_line_stacked_columns_then_style():
    result = compiler.ensure_blank_line_after_table_markers(
        '@@@TABLE_COLUMNS:1,x@@@\n@@@TABLE_STYLE:DilonTable_Chart@@@\n| a |\n'
    )
    check(result == '@@@TABLE_COLUMNS:1,x@@@\n@@@TABLE_STYLE:DilonTable_Chart@@@\n\n| a |\n',
          "stacked COLUMNS+STYLE markers (reverse order) get exactly one blank line inserted after the last marker line")


def test_ensure_blank_line_idempotent_when_already_blank():
    already_blank = '@@@TABLE_STYLE:DilonTable_Chart@@@\n@@@TABLE_COLUMNS:1,x@@@\n\n| a |\n'
    result = compiler.ensure_blank_line_after_table_markers(already_blank)
    check(result == already_blank,
          "already-blank-line marker stack is left unchanged (no extra blank line inserted between stacked markers)")


def test_ensure_blank_line_between_images_inserts_when_missing():
    result = compiler.ensure_blank_line_between_images(
        '![First.](a.png){#fig:a}\n![Second.](b.png){#fig:b}\n'
    )
    check(result == '![First.](a.png){#fig:a}\n\n![Second.](b.png){#fig:b}\n',
          "two back-to-back image-only lines get a blank line inserted between them")


def test_ensure_blank_line_between_images_three_in_a_row():
    result = compiler.ensure_blank_line_between_images(
        '![A.](a.png)\n![B.](b.png)\n![C.](c.png)\n'
    )
    check(result == '![A.](a.png)\n\n![B.](b.png)\n\n![C.](c.png)\n',
          "a run of 3+ back-to-back images gets a blank line inserted between every pair")


def test_ensure_blank_line_between_images_idempotent_when_already_blank():
    already_blank = '![First.](a.png)\n\n![Second.](b.png)\n'
    result = compiler.ensure_blank_line_between_images(already_blank)
    check(result == already_blank,
          "already-blank-line-separated images are left unchanged")


def test_ensure_blank_line_between_images_leaves_non_image_lines_alone():
    text = '![Figure.](a.png)\nSome prose right after it.\n'
    result = compiler.ensure_blank_line_between_images(text)
    check(result == text,
          "an image line followed by ordinary prose (not another image-only line) is left unchanged")


def test_ensure_blank_line_between_images_ignores_inline_image_with_text():
    text = '![Figure.](a.png) plus trailing text on the same line\n![Second.](b.png)\n'
    result = compiler.ensure_blank_line_between_images(text)
    check(result == text,
          "a line with an image PLUS other text isn't image-only, so no blank line is forced in")


def test_parse_column_widths_valid_with_flex():
    check(compiler.parse_column_widths('1.5,x,1,1', 4) == [1.5, 'x', 1.0, 1.0],
          "parse_column_widths accepts one 'x' entry among numeric widths")


def test_parse_column_widths_valid_all_numeric():
    check(compiler.parse_column_widths('1,2,3', 3) == [1.0, 2.0, 3.0],
          "parse_column_widths accepts an all-numeric spec with no flex column")


def test_parse_column_widths_case_insensitive_flex():
    check(compiler.parse_column_widths('1,X', 2) == [1.0, 'x'],
          "parse_column_widths treats 'X' the same as 'x'")


def test_parse_column_widths_rejects_count_mismatch():
    check(compiler.parse_column_widths('1,2', 3) is None,
          "parse_column_widths rejects a spec with fewer entries than columns")


def test_parse_column_widths_rejects_multiple_flex():
    check(compiler.parse_column_widths('x,x,1', 3) is None,
          "parse_column_widths rejects a spec with two 'x' entries")


def test_parse_column_widths_rejects_non_numeric():
    check(compiler.parse_column_widths('1,abc', 2) is None,
          "parse_column_widths rejects a non-numeric, non-'x' entry")


def test_parse_column_widths_rejects_zero_or_negative():
    check(compiler.parse_column_widths('0,1', 2) is None,
          "parse_column_widths rejects a zero-or-negative width")


def test_parse_column_widths_rejects_non_finite():
    check(compiler.parse_column_widths('1,nan', 2) is None,
          "parse_column_widths rejects a NaN width")
    check(compiler.parse_column_widths('1,inf', 2) is None,
          "parse_column_widths rejects an Infinity width")


def test_apply_table_column_widths_fixed_and_flex():
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    compiler.apply_table_column_widths(table, [1.5, 'x', 1.0, 1.0], available_width=6.27)
    widths = [round(col.width.inches, 2) for col in table.columns]
    check(widths == [1.5, 2.77, 1.0, 1.0],
          f"apply_table_column_widths sets fixed widths and computes the flex column (got {widths})")
    cell_widths = [round(cell.width.inches, 2) for cell in table.rows[0].cells]
    check(cell_widths == [1.5, 2.77, 1.0, 1.0],
          f"apply_table_column_widths also sets per-cell width, not just column width (got {cell_widths})")


def test_apply_table_column_widths_all_fixed():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    compiler.apply_table_column_widths(table, [2.0, 3.0], available_width=6.0)
    widths = [round(col.width.inches, 2) for col in table.columns]
    check(widths == [2.0, 3.0], f"apply_table_column_widths sets all-fixed widths as given (got {widths})")
    cell_widths = [round(cell.width.inches, 2) for cell in table.rows[0].cells]
    check(cell_widths == [2.0, 3.0],
          f"apply_table_column_widths also sets per-cell width, not just column width (got {cell_widths})")


def test_apply_table_column_widths_raises_when_flex_overflows():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    raised = False
    try:
        compiler.apply_table_column_widths(table, [5.0, 'x'], available_width=4.0)
    except ValueError:
        raised = True
    check(raised, "apply_table_column_widths raises ValueError when fixed widths leave no room for the flex column")


def test_apply_table_column_widths_raises_when_all_fixed_overflows():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    raised = False
    try:
        compiler.apply_table_column_widths(table, [4.0, 4.0], available_width=6.0)
    except ValueError:
        raised = True
    check(raised, "apply_table_column_widths raises ValueError when all-fixed widths exceed the available width")


def test_compile_missing_input_error():
    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(TEST_OUTPUT_DIR / "nonexistent.md"),
            str(TEST_OUTPUT_DIR / "nonexistent.docx"),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode != 0, "compiler reports a non-zero exit code for a missing input file")


def test_compile_valid_document():
    input_md = TEST_OUTPUT_DIR / "compile_test.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a valid document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "compile_test.docx created on disk")


def test_create_signature_table_structure():
    """create_signature_table() builds the signature-approval table
    programmatically (it used to be baked into TEMPLATE_Word_Base.docx
    as a docxtpl-rendered table)."""
    metadata = {
        "department": "Engineering",
        "author": "Jane Author",
        "regulatory_rep": "Reg Rep",
        "quality_rep": "QA Rep",
        "department_head": "Dept Head",
    }
    available_width = Inches(6.768055555555556)
    table = compiler.create_signature_table(metadata, available_width)

    check(len(table.rows) == 6, f"signature table has 6 rows, got {len(table.rows)}")
    check(len(table.columns) == 3, f"signature table has 3 columns, got {len(table.columns)}")

    rows_text = [[c.text for c in row.cells] for row in table.rows]
    check(rows_text[0] == ["Group", "Preparer", "Signature"], f"row 0 is the Group/Preparer/Signature header, got {rows_text[0]}")
    check(rows_text[1] == ["Engineering", "Jane Author", "Electronic"], f"row 1 has the preparer's department/author, got {rows_text[1]}")
    check(rows_text[2] == ["Department", "Name", "Signature"], f"row 2 is the Department/Name/Signature header, got {rows_text[2]}")
    check(rows_text[3] == ["Regulatory", "Reg Rep", "Electronic"], f"row 3 has the regulatory rep, got {rows_text[3]}")
    check(rows_text[4] == ["Quality", "QA Rep", "Electronic"], f"row 4 has the quality rep, got {rows_text[4]}")
    check(rows_text[5] == ["Engineering", "Dept Head", "Electronic"], f"row 5 has the department head, got {rows_text[5]}")

    check(table.rows[0].cells[0].paragraphs[0].runs[0].font.bold is True, "header row 0 is bold")
    check(table.rows[1].cells[0].paragraphs[0].runs[0].font.bold is not True, "data row 1 is not bold")


def test_compile_signature_table_generated_programmatically():
    """The signature-approval table is now built by create_signature_table()
    and inserted into Part A directly, instead of being pre-baked into
    TEMPLATE_Word_Base.docx as Jinja fields."""
    input_md = TEST_OUTPUT_DIR / "compile_test_signature.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_signature.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a document with a signature table")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)

    doc = Document(output_docx)

    def header_row(table):
        return [c.text for c in table.rows[0].cells]

    generated_tables = [t for t in doc.tables if header_row(t) == ["Group", "Preparer", "Signature"]]
    check(len(generated_tables) == 1, f"exactly one Group/Preparer/Signature table is present in the output, found {len(generated_tables)}")

    rows_text = [[c.text for c in row.cells] for t in generated_tables for row in t.rows]
    check(["Engineering", "Test Suite", "Electronic"] in rows_text, "preparer row (department/author) rendered from front matter")
    check(["Regulatory", "Test Rep", "Electronic"] in rows_text, "regulatory rep row rendered from front matter")
    check(["Quality", "Test QA", "Electronic"] in rows_text, "quality rep row rendered from front matter")
    check(["Engineering", "Test Head", "Electronic"] in rows_text, "department head row rendered from front matter")


def test_compile_has_no_leading_blank_paragraph():
    """TEMPLATE_Word_Base.docx's body always carries one empty paragraph
    before its sectPr (every real Word file needs at least one paragraph
    if it has no other body content). Part A used to keep that paragraph
    ahead of the signature table it inserts, so the compiled document's
    body started with a stray blank paragraph instead of the table
    (same root cause dilon-document-form-compiler had - see
    strip_leading_empty_paragraphs() in lib/dilon_docx_common.py).

    The header-to-body gap on every page (including this one) is governed
    by TEMPLATE_Word_Base.docx's top margin instead, so no spacer
    paragraph belongs here - the table should be the body's first
    content, immediately after the header/footer/signature-table setup."""
    input_md = TEST_OUTPUT_DIR / "compile_test_no_leading_blank.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_no_leading_blank.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for the no-leading-blank check document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "compile_test_no_leading_blank.docx created on disk")
    if not output_docx.exists():
        return

    doc = Document(output_docx)
    first_child = doc.element.body[0]
    check(
        first_child.tag == qn('w:tbl'),
        f"body's first content element is the signature table, no stray leading blank paragraph, got {first_child.tag!r}",
    )


def test_compile_has_no_title_page():
    input_md = TEST_OUTPUT_DIR / "compile_test_no_title_page.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_no_title_page.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a valid document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    doc = Document(output_docx)
    check(not any(p.style and p.style.name == 'Title' for p in doc.paragraphs),
          "no paragraph carries the 'Title' style - the title page is gone")
    check(not any('Master Document' in p.text for p in doc.paragraphs),
          "the title page's 'Master Document' boilerplate is gone")
    check(not any('Effectivity and Location' in p.text for p in doc.paragraphs),
          "the title page's 'Effectivity and Location' boilerplate is gone")


def test_compile_header_signature_revision_widths():
    """Regression test for a formatting request: the running header,
    signature-approval table, and revision-history table must all extend
    to the page's full content width (previously narrower than the
    margins, leaving unused space on the right), with their fixed
    columns matching a hand-tuned reference document's widths exactly and
    their remaining column absorbing whatever width is left."""
    input_md = TEST_OUTPUT_DIR / "compile_test_column_geometry.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_column_geometry.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for the column-geometry check document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    doc = Document(output_docx)
    section = doc.sections[0]
    available_width_twips = section.page_width.twips - section.left_margin.twips - section.right_margin.twips

    def grid_dxa(table):
        grid = table._element.find(qn('w:tblGrid'))
        return [int(g.get(qn('w:w'))) for g in grid.findall(qn('w:gridCol'))]

    def cell_dxa(table, row_idx):
        """Per-cell tcW width, not just the tblGrid definition - this is
        what Word's AutoFit actually renders from when a table isn't
        pinned to a fixed layout, so a passing grid_dxa() check alone
        does NOT prove the widths render correctly (see
        test_apply_table_column_widths_fixed_and_flex for the same
        distinction on @@@TABLE_COLUMNS@@@-marked tables)."""
        widths = []
        for cell in table.rows[row_idx].cells:
            tc_pr = cell._element.tcPr
            tc_w = tc_pr.find(qn('w:tcW')) if tc_pr is not None else None
            widths.append(int(tc_w.get(qn('w:w'))) if tc_w is not None else None)
        return widths

    def is_fixed_layout(table):
        tbl_layout = table._element.tblPr.find(qn('w:tblLayout'))
        return tbl_layout is not None and tbl_layout.get(qn('w:type')) == 'fixed'

    header_table = section.header.tables[0]
    header_widths = grid_dxa(header_table)
    check(header_widths[0] == 1525 and header_widths[2] == 1255 and header_widths[3] == 1625,
          f"header logo/Rev/Page columns match the reference document's widths (got {header_widths})")
    check(sum(header_widths) == available_width_twips,
          f"header table fills the full page content width (got {sum(header_widths)}, expected {available_width_twips})")
    check(cell_dxa(header_table, 0) == header_widths,
          f"header row's actual per-cell widths match the tblGrid definition, not just the grid (got {cell_dxa(header_table, 0)})")
    check(is_fixed_layout(header_table), "header table uses a fixed layout, so Word can't AutoFit its columns away")

    def header_row(table):
        return [c.text for c in table.rows[0].cells]

    signature_table = next(t for t in doc.tables if header_row(t) == ["Group", "Preparer", "Signature"])
    sig_widths = grid_dxa(signature_table)
    check(sig_widths[0] == 2330 and sig_widths[2] == 2440,
          f"signature table Group/Signature columns match the reference document's widths (got {sig_widths})")
    check(sum(sig_widths) == available_width_twips,
          f"signature table fills the full page content width (got {sum(sig_widths)}, expected {available_width_twips})")
    check(cell_dxa(signature_table, 1) == sig_widths,
          f"signature table's actual per-cell widths match the tblGrid definition, not just the grid (got {cell_dxa(signature_table, 1)})")
    check(is_fixed_layout(signature_table), "signature table uses a fixed layout, so Word can't AutoFit its columns away")

    revision_table = next(t for t in doc.tables if header_row(t)[0] == "REVISION HISTORY")
    rev_widths = grid_dxa(revision_table)
    check(rev_widths[0] == 805 and rev_widths[2] == 1620 and rev_widths[3] == 1535,
          f"revision table REV#/ECO#/DATE columns match the reference document's widths (got {rev_widths})")
    check(sum(rev_widths) == available_width_twips,
          f"revision table fills the full page content width (got {sum(rev_widths)}, expected {available_width_twips})")
    check(cell_dxa(revision_table, 2) == rev_widths,
          f"revision table's actual per-cell widths match the tblGrid definition, not just the grid (got {cell_dxa(revision_table, 2)})")
    check(is_fixed_layout(revision_table), "revision table uses a fixed layout, so Word can't AutoFit its columns away")


def test_compile_footer_table_layout():
    """Regression test: the running footer is a 3-column/2-row table -
    row 1 is doc_number/rev (left) | ECO # (center) | revision date
    (right) in three equal-width columns, row 2 is a single cell spanning
    all 3 columns for the confidentiality notice. The table must fill the
    full page content width and use a fixed layout (see
    test_compile_header_signature_revision_widths_and_author_centering's
    is_fixed_layout for why that matters)."""
    input_md = TEST_OUTPUT_DIR / "compile_test_footer_table.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_footer_table.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for the footer-table check document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    doc = Document(output_docx)
    section = doc.sections[0]
    available_width_twips = section.page_width.twips - section.left_margin.twips - section.right_margin.twips

    footer_tables = section.footer.tables
    check(len(footer_tables) == 1, f"footer contains exactly one table (got {len(footer_tables)})")
    if not footer_tables:
        return
    table = footer_tables[0]

    check(len(table.rows) == 2 and len(table.columns) == 3,
          f"footer table is 3 columns x 2 rows (got {len(table.columns)}x{len(table.rows)})")

    grid = table._element.find(qn('w:tblGrid'))
    grid_widths = [int(g.get(qn('w:w'))) for g in grid.findall(qn('w:gridCol'))]
    check(sum(grid_widths) == available_width_twips,
          f"footer table fills the full page content width (got {sum(grid_widths)}, expected {available_width_twips})")
    check(max(grid_widths) - min(grid_widths) <= 2,
          f"footer table's 3 columns are equal width, up to a rounding twip or two (got {grid_widths})")

    tbl_layout = table._element.tblPr.find(qn('w:tblLayout'))
    check(tbl_layout is not None and tbl_layout.get(qn('w:type')) == 'fixed',
          "footer table uses a fixed layout, so Word can't AutoFit its columns away")

    row1_texts = [c.text for c in table.rows[0].cells]
    check("DD_TST_99999 Rev 00" in row1_texts[0], f"footer row 1 col 1 has the doc_number/rev (got {row1_texts[0]!r})")
    check(row1_texts[1] == "ECO-000", f"footer row 1 col 2 has the ECO number (got {row1_texts[1]!r})")
    check("2025-01-01" in row1_texts[2], f"footer row 1 col 3 has the revision date (got {row1_texts[2]!r})")

    row1_alignments = [c.paragraphs[0].alignment for c in table.rows[0].cells]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    check(row1_alignments == [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT],
          f"footer row 1 is left/center/right justified across its 3 columns (got {row1_alignments})")

    row2_cells = table.rows[1].cells
    check(row2_cells[0]._element is row2_cells[1]._element is row2_cells[2]._element,
          "footer row 2's three grid positions resolve to one merged cell spanning the full width")
    check("confidential" in row2_cells[0].text and "prohibited" in row2_cells[0].text,
          f"footer row 2's merged cell carries the confidentiality notice (got {row2_cells[0].text!r})")


def test_compile_bom_front_matter():
    """Regression test: a UTF-8 BOM at the start of the input markdown
    (e.g. from PowerShell's Set-Content -Encoding UTF8) must not cause the
    YAML front matter to be silently dropped."""
    input_md = TEST_OUTPUT_DIR / "compile_test_bom.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_bom.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8-sig")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a BOM-prefixed input file")
    check("Metadata extracted: []" not in result.stdout, "BOM-prefixed front matter is not silently dropped")

    if output_docx.exists():
        doc = Document(output_docx)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # "No revision history available" is the Part B fallback text used only
        # when 'revisions' is missing from metadata - i.e. only when the YAML
        # front matter failed to parse.
        check("No revision history available" not in full_text, "revision table from BOM-prefixed file's metadata reaches the output document")
    else:
        check(False, "revision table from BOM-prefixed file's metadata reaches the output document")


TABLE_MARKER_MARKDOWN = SAMPLE_MARKDOWN + (
    '\n### 1.3 Table Test\n'
    '@@@TABLE_STYLE:DilonTable_Chart@@@\n'
    '| Thread | Zephyr Name | Priority |\n'
    '|---|---|---|\n'
    '| A | B | C |\n'
    '\n'
    '| Default | Table |\n'
    '|---|---|\n'
    '| X | Y |\n'
)


def test_compile_table_marker_no_blank_line():
    """Regression test: a @@@TABLE_STYLE@@@ marker immediately followed by
    a pipe table with no blank line (the documented convention) must
    produce a real, styled table - not garbled literal text."""
    input_md = TEST_OUTPUT_DIR / "compile_test_table_marker.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_table_marker.docx"
    input_md.write_text(TABLE_MARKER_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a marker-adjacent table")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "marker-adjacent table renders as a real, styled table (skipped: compile failed)")
        return

    doc = Document(output_docx)

    def header_row(table):
        return [c.text for c in table.rows[0].cells]

    marked_tables = [t for t in doc.tables if header_row(t) == ['Thread', 'Zephyr Name', 'Priority']]
    check(len(marked_tables) == 1, "marker-adjacent table survives conversion as a real table")
    if marked_tables:
        check(marked_tables[0].style is not None and marked_tables[0].style.name == 'DilonTable_Chart',
              "marker-adjacent table receives the DilonTable_Chart style")

    default_tables = [t for t in doc.tables if header_row(t) == ['Default', 'Table']]
    check(len(default_tables) == 1, "unmarked table survives conversion as a real table")
    if default_tables:
        check(default_tables[0].style is not None and default_tables[0].style.name == 'DilonTable_List',
              "unmarked table still receives the default DilonTable_List style")

    garbled = [p.text for p in doc.paragraphs if '@@@' in p.text or '|---' in p.text]
    check(not garbled, f"no leftover marker/pipe-table text in output (found: {garbled})")


ADJACENT_TABLES_MARKDOWN = SAMPLE_MARKDOWN + (
    '\n### 1.3 Adjacent Tables Test\n'
    '@@@TABLE_COLUMNS:1,1@@@\n'
    '| First | Table |\n'
    '|---|---|\n'
    '| A | B |\n'
    '\n'
    '@@@TABLE_COLUMNS:1,1@@@\n'
    '| Second | Table |\n'
    '|---|---|\n'
    '| C | D |\n'
)


def test_compile_adjacent_tables_no_merge():
    """Regression test: two tables back-to-back with only a
    @@@TABLE_COLUMNS@@@ marker (no other body text) between them must not
    be merged by Word. Deleting the marker paragraph entirely would leave
    the two <w:tbl> elements directly adjacent in document.xml, which Word
    renders as a single merged table on open - so the marker paragraph
    must be emptied, not removed, when it is the only separator."""
    input_md = TEST_OUTPUT_DIR / "compile_test_adjacent_tables.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_adjacent_tables.docx"
    input_md.write_text(ADJACENT_TABLES_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for back-to-back marker-separated tables")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "adjacent tables stay separate (skipped: compile failed)")
        return

    doc = Document(output_docx)

    def header_row(table):
        return [c.text for c in table.rows[0].cells]

    first_tables = [t for t in doc.tables if header_row(t) == ['First', 'Table']]
    second_tables = [t for t in doc.tables if header_row(t) == ['Second', 'Table']]
    check(len(first_tables) == 1, "first table survives as its own distinct table")
    check(len(second_tables) == 1, "second table survives as its own distinct table")

    # Walk the raw body children and confirm no <w:tbl> is immediately
    # followed by another <w:tbl> with zero paragraphs in between - that
    # adjacency is exactly what Word merges into one visual table.
    body_children = list(doc.element.body)
    merged_pairs = [
        i for i in range(len(body_children) - 1)
        if body_children[i].tag.endswith('tbl') and body_children[i + 1].tag.endswith('tbl')
    ]
    check(not merged_pairs,
          f"no <w:tbl> is directly adjacent to another <w:tbl> in document.xml (found {len(merged_pairs)} such pair(s))")

    garbled = [p.text for p in doc.paragraphs if '@@@' in p.text or '|---' in p.text]
    check(not garbled, f"no leftover marker/pipe-table text in output (found: {garbled})")


COLUMN_WIDTH_MARKDOWN = SAMPLE_MARKDOWN + (
    '\n### 1.3 Column Width Test\n'
    '@@@TABLE_STYLE:DilonTable_Chart@@@\n'
    '@@@TABLE_COLUMNS:1.5,x,1,1@@@\n'
    '| Register | Address | Bit 7 | Bit 6 |\n'
    '|---|---|---|---|\n'
    '| CTRL_REG1 | 0x20 | ODR3 | ODR2 |\n'
    '\n'
    '@@@TABLE_COLUMNS:2,3@@@\n'
    '| Name | Value |\n'
    '|---|---|\n'
    '| A | B |\n'
    '\n'
    '@@@TABLE_COLUMNS:1,2@@@\n'
    '| Mismatched | Columns | Table |\n'
    '|---|---|---|\n'
    '| A | B | C |\n'
    '\n'
    '@@@TABLE_COLUMNS:10,x@@@\n'
    '| Overflow | Table |\n'
    '|---|---|\n'
    '| A | B |\n'
    '\n'
    '@@@TABLE_STYLE:DilonTable_Chart@@@\n'
    '@@@TABLE_COLUMNS:1.5, x, 1, 1@@@\n'
    '| Spaced | Spec | Bit A | Bit B |\n'
    '|---|---|---|---|\n'
    '| SPACED_REG | 0x21 | ODR1 | ODR0 |\n'
    '\n'
    '@@@TABLE_COLUMNS:2,x@@@\n'
    '@@@TABLE_STYLE:DilonTable_Chart@@@\n'
    '| Order | Test |\n'
    '|---|---|\n'
    '| A | B |\n'
)


def test_compile_table_column_widths():
    """Render test: compile a document with @@@TABLE_COLUMNS@@@-marked
    tables (stacked with @@@TABLE_STYLE@@@, standalone, and two invalid
    specs) and verify the ACTUAL rendered column widths in the output
    docx match the marker's specified values - not just that
    compilation succeeded."""
    input_md = TEST_OUTPUT_DIR / "compile_test_column_widths.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_column_widths.docx"
    input_md.write_text(COLUMN_WIDTH_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for column-width-marked tables")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "column-width tables render with correct widths (skipped: compile failed)")
        return

    doc = Document(output_docx)

    def header_row(table):
        return [c.text for c in table.rows[0].cells]

    section = doc.sections[0]
    available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches

    # Table 1: stacked @@@TABLE_STYLE@@@ + @@@TABLE_COLUMNS:1.5,x,1,1@@@
    reg_tables = [t for t in doc.tables if header_row(t) == ['Register', 'Address', 'Bit 7', 'Bit 6']]
    check(len(reg_tables) == 1, "stacked-marker table survives conversion as a real table")
    if reg_tables:
        table = reg_tables[0]
        check(table.style is not None and table.style.name == 'DilonTable_Chart',
              "stacked-marker table still receives its DilonTable_Chart style")
        actual_widths = [round(col.width.inches, 2) for col in table.columns]
        expected_flex = round(available_width - 1.5 - 1.0 - 1.0, 2)
        expected_widths = [1.5, expected_flex, 1.0, 1.0]
        check(actual_widths == expected_widths,
              f"stacked-marker table's rendered widths match the marker spec (expected {expected_widths}, got {actual_widths})")

    # Table 2: standalone @@@TABLE_COLUMNS:2,3@@@, no @@@TABLE_STYLE@@@
    name_tables = [t for t in doc.tables if header_row(t) == ['Name', 'Value']]
    check(len(name_tables) == 1, "standalone-column-width table survives conversion as a real table")
    if name_tables:
        table = name_tables[0]
        actual_widths = [round(col.width.inches, 2) for col in table.columns]
        check(actual_widths == [2.0, 3.0],
              f"standalone all-numeric column widths render exactly as specified (got {actual_widths})")
        check(table.style is not None and table.style.name == 'DilonTable_List',
              "table with only a @@@TABLE_COLUMNS@@@ marker still gets the default DilonTable_List style")

    # Table 3: @@@TABLE_COLUMNS:1,2@@@ on a 3-column table - entry count mismatch, must warn and skip
    mismatched_tables = [t for t in doc.tables if header_row(t) == ['Mismatched', 'Columns', 'Table']]
    check(len(mismatched_tables) == 1, "entry-count-mismatch table still survives conversion as a real table")
    check("Invalid @@@TABLE_COLUMNS@@@ spec" in result.stdout,
          "entry-count-mismatch spec prints a warning instead of failing compilation")

    # Table 4: @@@TABLE_COLUMNS:10,x@@@ - fixed width alone exceeds available content width
    overflow_tables = [t for t in doc.tables if header_row(t) == ['Overflow', 'Table']]
    check(len(overflow_tables) == 1, "overflowing-width table still survives conversion as a real table")
    check("Could not apply column widths" in result.stdout,
          "overflowing column-width spec prints a warning instead of failing compilation")

    # Table 5: stacked @@@TABLE_STYLE@@@ + @@@TABLE_COLUMNS:1.5, x, 1, 1@@@ (spaced spec)
    spaced_tables = [t for t in doc.tables if header_row(t) == ['Spaced', 'Spec', 'Bit A', 'Bit B']]
    check(len(spaced_tables) == 1, "spaced-spec table survives conversion as a real table")
    if spaced_tables:
        table = spaced_tables[0]
        check(table.style is not None and table.style.name == 'DilonTable_Chart',
              "spaced-spec table still receives its DilonTable_Chart style")
        actual_widths = [round(col.width.inches, 2) for col in table.columns]
        expected_flex = round(available_width - 1.5 - 1.0 - 1.0, 2)
        expected_widths = [1.5, expected_flex, 1.0, 1.0]
        check(actual_widths == expected_widths,
              f"spaced @@@TABLE_COLUMNS:1.5, x, 1, 1@@@ spec renders with correct widths, not garbled "
              f"(expected {expected_widths}, got {actual_widths})")

    # Table 6: @@@TABLE_COLUMNS:2,x@@@ followed by @@@TABLE_STYLE@@@ (COLUMNS-then-STYLE order)
    order_tables = [t for t in doc.tables if header_row(t) == ['Order', 'Test']]
    check(len(order_tables) == 1, "COLUMNS-then-STYLE table survives conversion as a real table")
    if order_tables:
        table = order_tables[0]
        check(table.style is not None and table.style.name == 'DilonTable_Chart',
              "COLUMNS-then-STYLE table (reverse marker order) still receives its DilonTable_Chart style")
        actual_widths = [round(col.width.inches, 2) for col in table.columns]
        expected_flex = round(available_width - 2.0, 2)
        expected_widths = [2.0, expected_flex]
        check(actual_widths == expected_widths,
              f"COLUMNS-then-STYLE table's rendered widths match the marker spec (expected {expected_widths}, got {actual_widths})")


def test_compile_with_default_templates():
    """Regression test for a bug where the compiler's default template
    lookup pointed at the wrong directory instead of the repo-root
    templates/ directory. Invokes with only <input> <output> (no template
    arg) so the script must resolve its own default, rather than the
    explicit three-argument form SKILL.md always uses."""
    input_md = TEST_OUTPUT_DIR / "compile_test_defaults.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_defaults.docx"
    input_md.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 with only 2 args (default template lookup)")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
    check(output_docx.exists(), "compile_test_defaults.docx created via default template lookup")


HEADING_NUMBERING_MARKDOWN = (
    '---\n'
    'title: "Heading Numbering Test Document"\n'
    'author: "Test Suite"\n'
    'department: "Engineering"\n'
    'doc_number: "DD_TST_88888"\n'
    'current_revision: "00"\n'
    'regulatory_rep: "Test Rep"\n'
    'quality_rep: "Test QA"\n'
    'department_head: "Test Head"\n'
    'revisions:\n'
    '  - number: "00"\n'
    '    description: "Initial test"\n'
    '    eco_number: "ECO-000"\n'
    '    eco_date: "2025-01-01"\n'
    '---\n'
    '\n'
    '## First Section\n'
    '\n'
    '### First Subsection\n'
    'Content.\n'
    '\n'
    '#### First Nested Item\n'
    'Content.\n'
    '\n'
    '## Second Section\n'
    '\n'
    '### Second Subsection\n'
    'Content.\n'
)


def make_test_png():
    """Build a minimal valid 1x1 PNG in-memory (no Pillow dependency)."""
    def chunk(tag, data):
        return struct.pack('>I', len(data)) + tag + data + struct.pack('>I', zlib.crc32(tag + data))
    ihdr = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)  # 1x1, 8-bit RGB
    idat = zlib.compress(b'\x00\xff\x00\x00')  # filter byte + one red pixel
    return b'\x89PNG\r\n\x1a\n' + chunk(b'IHDR', ihdr) + chunk(b'IDAT', idat) + chunk(b'IEND', b'')


IMAGE_PATH_MARKDOWN = SAMPLE_MARKDOWN + (
    '\n## Image Test\n\n'
    '![A tiny red test image.](images_subdir/test.png)\n'
)


def test_compile_resolves_relative_image_paths():
    """Regression test: relative image paths in the markdown (documented
    as 'relative to the markdown file' in MARKDOWN_STYLING_GUIDE.md) must
    resolve against the INPUT MARKDOWN's directory, not whatever cwd the
    compiler happens to be invoked from. Before the --resource-path fix,
    running the compiler from a different cwd than the markdown file
    silently dropped the image - Pandoc warns and substitutes a
    placeholder, but still exits 0, so nothing failed loudly."""
    images_dir = TEST_OUTPUT_DIR / "images_subdir"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "test.png").write_bytes(make_test_png())

    input_md = TEST_OUTPUT_DIR / "compile_test_images.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_images.docx"
    input_md.write_text(IMAGE_PATH_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        cwd=str(REPO_ROOT),  # deliberately NOT the markdown file's own directory
    )
    check(result.returncode == 0, "compiler exits 0 for a document with a relative image path")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "relative image path resolves and embeds (skipped: compile failed)")
        return

    with zipfile.ZipFile(output_docx) as z:
        rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
        check('relationships/image' in rels,
              "compiled document's body (not just the header) embeds an image relationship")


def test_render_jinja_substitutes_body_fields():
    text = compiler.render_jinja("This document is {{doc_number}}, rev {{current_revision}}.", {
        "doc_number": "WI-00077",
        "current_revision": "01",
    })
    check(text == "This document is WI-00077, rev 01.", f"body Jinja2 fields resolved, got {text!r}")


def test_render_jinja_raw_block_escapes_literal_braces():
    text = compiler.render_jinja(
        "Use {% raw %}{{doc_number}}{% endraw %} to reference the doc number.",
        {"doc_number": "WI-00077"},
    )
    check(
        text == "Use {{doc_number}} to reference the doc number.",
        f"raw block preserves literal braces, got {text!r}",
    )


def test_render_jinja_noop_without_braces():
    text = compiler.render_jinja("Plain text with no template fields.", {"doc_number": "WI-00077"})
    check(text == "Plain text with no template fields.", "body with no {{...}} is returned unchanged")


def test_compile_body_jinja_substitution():
    input_md = TEST_OUTPUT_DIR / "compile_test_jinja.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_jinja.docx"
    markdown = SAMPLE_MARKDOWN.replace(
        "This document tests the compilation process.",
        "This document, {{doc_number}}, tests the compilation process.",
    )
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a document with body-level {{doc_number}}")

    doc = Document(output_docx)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    check("DD_TST_99999" in body_text, "body {{doc_number}} resolved to the front-matter value in the compiled output")
    check("{{doc_number}}" not in body_text, "no literal Jinja2 braces remain in the compiled output")


def test_figure_auto_numbering():
    """Render test: a figure caption written as
    ![Description.](path.png){#fig:label} (per the updated
    MARKDOWN_STYLING_GUIDE.md convention) must come out of Pandoc as a
    distinct 'Image Caption'-styled paragraph (thanks to the
    'Captioned Figure'/'Image Caption' styles added to
    TEMPLATE_Word_Base.docx), which apply_figure_captions() then
    rewrites into a 'Caption'-styled paragraph carrying live
    STYLEREF/SEQ fields - not static numbered text. Also verifies the
    {#fig:label} bookmark and a [text](#fig:label) cross-reference
    resolve to a real hyperlink, and that Word is told to recalculate
    fields on open."""
    images_dir = TEST_OUTPUT_DIR / "figure_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "labeled.png").write_bytes(make_test_png())
    (images_dir / "unlabeled.png").write_bytes(make_test_png())

    markdown = SAMPLE_MARKDOWN + (
        '\n## First Section\n\n'
        'See [the figure](#fig:test-figure) below.\n\n'
        '![A labeled test figure.](figure_images/labeled.png){#fig:test-figure}\n\n'
        '## Second Section\n\n'
        '![An unlabeled test figure.](figure_images/unlabeled.png)\n'
    )

    input_md = TEST_OUTPUT_DIR / "compile_test_figures.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_figures.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a document with figure captions")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "figure captions render as auto-numbered Word captions (skipped: compile failed)")
        return

    doc = Document(output_docx)
    caption_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == 'Caption']
    check(len(caption_paragraphs) == 2, f"exactly 2 paragraphs end up styled 'Caption' (got {len(caption_paragraphs)})")

    with zipfile.ZipFile(output_docx) as z:
        xml = z.read('word/document.xml').decode('utf-8')
        settings = z.read('word/settings.xml').decode('utf-8')

    instrs = re.findall(r'w:instr="([^"]*)"', xml)
    styleref_count = sum(1 for i in instrs if i.strip() == 'STYLEREF 2 \\s')
    seq_count = sum(1 for i in instrs if i.strip() == 'SEQ Figure \\* ARABIC \\s 2')
    check(styleref_count == 2, f"both captions get a 'STYLEREF 2 \\\\s' chapter-number field (got {styleref_count})")
    check(seq_count == 2, f"both captions get a 'SEQ Figure \\\\* ARABIC \\\\s 2' running-count field (got {seq_count})")

    check('w:name="fig:test-figure"' in xml, "the labeled figure gets a 'fig:test-figure' bookmark")
    check('w:anchor="fig:test-figure"' in xml, "the [the figure](#fig:test-figure) link becomes a real hyperlink to that bookmark")

    check('<w:updateFields w:val="true"/>' in settings, "document is set to recalculate fields (figure numbers) on open")

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    image_paragraphs = [p for p in doc.paragraphs if p._p.findall('.//' + qn('a:blip'))]
    check(len(image_paragraphs) == 2, f"both the labeled and unlabeled image end up in their own paragraph (got {len(image_paragraphs)})")
    check(all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in image_paragraphs),
          f"every image paragraph is centered, captioned or not (got {[p.alignment for p in image_paragraphs]})")


def test_figure_auto_numbering_consecutive_images_no_blank_line():
    """Regression test: two figures declared back-to-back with no blank
    line between them (a natural way to author consecutive figures) used
    to silently lose ALL figure treatment for BOTH images - Markdown
    merges adjacent non-blank lines into one paragraph, so neither image
    was alone in its paragraph, and Pandoc's implicit-figures extension
    only promotes a solo image to a captioned figure. Fixed by
    ensure_blank_line_between_images() in lib/dilon_docx_common.py."""
    images_dir = TEST_OUTPUT_DIR / "consecutive_figure_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "first.png").write_bytes(make_test_png())
    (images_dir / "second.png").write_bytes(make_test_png())

    markdown = SAMPLE_MARKDOWN + (
        '\n## First Section\n\n'
        '![First figure.](consecutive_figure_images/first.png){#fig:consecutive-first}\n'
        '![Second figure.](consecutive_figure_images/second.png){#fig:consecutive-second}\n'
    )

    input_md = TEST_OUTPUT_DIR / "compile_test_consecutive_figures.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_consecutive_figures.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for back-to-back (no blank line) figures")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "back-to-back figures both get caption treatment (skipped: compile failed)")
        return

    doc = Document(output_docx)
    caption_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == 'Caption']
    check(len(caption_paragraphs) == 2,
          f"both back-to-back figures end up styled 'Caption', not merged into one plain paragraph (got {len(caption_paragraphs)})")

    with zipfile.ZipFile(output_docx) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    check('w:name="fig:consecutive-first"' in xml, "the first of the two back-to-back figures still gets its {#fig:...} bookmark")
    check('w:name="fig:consecutive-second"' in xml, "the second of the two back-to-back figures still gets its {#fig:...} bookmark")
    check(xml.count('<a:blip') == 2, f"both images are still embedded (got {xml.count('<a:blip')})")


def test_apply_figure_captions_narrows_bookmark_to_number_only():
    images_dir = TEST_OUTPUT_DIR / "fig_bookmark_narrow_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "test.png").write_bytes(make_test_png())

    md = "![A test caption.](fig_bookmark_narrow_images/test.png){#fig:my-label}\n"
    docx_path = TEST_OUTPUT_DIR / "fig_bookmark_narrow_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE, resource_dir=TEST_OUTPUT_DIR)
    dilon_docx_common.apply_figure_captions(docx_path)

    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml').decode('utf-8')

    check('w:name="fig:my-label"' in xml, "the fig:my-label bookmark still exists after narrowing")
    drawing_pos = xml.find('<w:drawing>')
    bookmark_pos = xml.find('w:name="fig:my-label"')
    check(bookmark_pos > drawing_pos, "the bookmark now starts AFTER the image, not before it")
    caption_pos = xml.find('Figure ')
    check(bookmark_pos != -1 and caption_pos != -1 and abs(bookmark_pos - caption_pos) < 60,
          "the bookmark sits immediately around the 'Figure ' text, not far from it")


def test_apply_figure_captions_narrowing_is_noop_without_fig_id():
    images_dir = TEST_OUTPUT_DIR / "fig_bookmark_no_id_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "test.png").write_bytes(make_test_png())

    md = "![A caption with no id.](fig_bookmark_no_id_images/test.png)\n"
    docx_path = TEST_OUTPUT_DIR / "fig_bookmark_no_id_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE, resource_dir=TEST_OUTPUT_DIR)
    dilon_docx_common.apply_figure_captions(docx_path)  # should not raise

    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    check('bookmarkStart' not in xml, "no bookmark is fabricated for a figure with no {#fig:...} id")


def test_narrow_section_bookmarks_shrinks_to_heading_only():
    md = "## Section One {#sec:one}\n\nSome body text.\n\n## Section Two {#sec:two}\n\nMore body text.\n"
    docx_path = TEST_OUTPUT_DIR / "sec_bookmark_narrow_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    count = dilon_docx_common.narrow_section_bookmarks(docx_path)
    check(count == 2, f"both section bookmarks are narrowed (got {count})")

    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    one_start = xml.find('w:name="sec:one"')
    body_text_pos = xml.find('Some body text.')
    check(one_start != -1 and body_text_pos != -1, "both markers present")
    one_end_search_region = xml[one_start:body_text_pos]
    check('bookmarkEnd' in one_end_search_region, "sec:one's bookmarkEnd now closes before the section's body text")


def test_narrow_section_bookmarks_ignores_headings_without_sec_id():
    md = "## Plain Heading\n\nBody text.\n"
    docx_path = TEST_OUTPUT_DIR / "sec_bookmark_no_id_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    count = dilon_docx_common.narrow_section_bookmarks(docx_path)
    check(count == 0, f"a heading with no {{#sec:...}} id is left alone (got count={count})")


def test_preprocess_reference_markers_converts_recognized_types():
    md = "See [](#fig:a) and [](#sec:b) and [](#step:c)."
    result = dilon_docx_common.preprocess_reference_markers(md)
    check(result == "See XREF:fig:a and XREF:sec:b and XREF:step:c.", f"all three types convert to sentinels (got {result!r})")


def test_preprocess_reference_markers_leaves_real_link_text_and_unknown_types_untouched():
    md = "[see the figure](#fig:a) and [](#other:x)"
    result = dilon_docx_common.preprocess_reference_markers(md)
    check(result == md, "non-empty link text and an unrecognized type prefix are both left alone")


def test_resolve_reference_markers_dispatches_by_type():
    images_dir = TEST_OUTPUT_DIR / "xref_dispatch_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "test.png").write_bytes(make_test_png())

    md = (
        "## Section One {#sec:intro}\n\n"
        "![A caption.](xref_dispatch_images/test.png){#fig:pic}\n\n"
        "See XREF:fig:pic and XREF:sec:intro.\n"
    )
    docx_path = TEST_OUTPUT_DIR / "xref_dispatch_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE, resource_dir=TEST_OUTPUT_DIR)
    dilon_docx_common.apply_figure_captions(docx_path)
    dilon_docx_common.narrow_section_bookmarks(docx_path)

    resolved = dilon_docx_common.resolve_reference_markers(docx_path, {
        'fig': dilon_docx_common.resolve_fig_reference,
        'sec': dilon_docx_common.resolve_sec_reference,
    })
    check(resolved == 2, f"both sentinels resolved (got {resolved})")

    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    check('REF fig:pic \\h' in xml, "fig resolves to a plain hyperlinked REF (no \\r - it's not a native list item)")
    check('Section ' in xml and 'REF sec:intro \\r \\h' in xml, "sec resolves to literal 'Section ' + a hyperlinked REF \\r")
    check('XREF' not in xml, "no sentinel remains")


def test_resolve_reference_markers_missing_anchor_raises():
    md = "See XREF:fig:does-not-exist."
    docx_path = TEST_OUTPUT_DIR / "xref_missing_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    try:
        dilon_docx_common.resolve_reference_markers(docx_path, {'fig': dilon_docx_common.resolve_fig_reference})
        check(False, "a reference to a nonexistent anchor raises ReferenceResolutionError")
    except dilon_docx_common.ReferenceResolutionError as exc:
        check("does-not-exist" in str(exc), f"the error names the missing label (got: {exc})")


def test_resolve_reference_markers_duplicate_anchor_raises():
    md = "## One {#sec:dup}\n\nText.\n\n## Two {#sec:dup}\n\nSee XREF:sec:dup.\n"
    docx_path = TEST_OUTPUT_DIR / "xref_duplicate_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    try:
        dilon_docx_common.resolve_reference_markers(docx_path, {'sec': dilon_docx_common.resolve_sec_reference})
        check(False, "two {#sec:dup} anchors raises ReferenceResolutionError")
    except dilon_docx_common.ReferenceResolutionError as exc:
        check("dup" in str(exc), f"the error names the duplicated label (got: {exc})")


def test_heading_auto_numbering():
    """Render test: headings written WITHOUT manual numbers (per the
    updated MARKDOWN_STYLING_GUIDE.md convention) must come out of Pandoc
    styled as Heading 2/3/4, and those styles must be linked (via the
    signature template's --reference-doc numbering) to a single shared
    multilevel list so Word auto-numbers them as 1./1.1/1.1.1, etc.
    A missing/broken link here is exactly the failure mode that produces
    unnumbered or doubled section numbers in compiled documents."""
    input_md = TEST_OUTPUT_DIR / "compile_test_heading_numbering.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_heading_numbering.docx"
    input_md.write_text(HEADING_NUMBERING_MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [
            sys.executable,
            str(COMPILER_SCRIPT),
            str(input_md),
            str(output_docx),
            str(SIGNATURE_TEMPLATE),
        ],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for unnumbered headings")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        check(False, "Heading 2/3/4 styles share one auto-numbering list (skipped: compile failed)")
        return

    with zipfile.ZipFile(output_docx) as z:
        check("word/numbering.xml" in z.namelist(),
              "compiled document contains a word/numbering.xml part")

    doc = Document(output_docx)

    heading_text = [p.text for p in doc.paragraphs if p.style and p.style.name in ("Heading 2", "Heading 3", "Heading 4")]
    check(heading_text == ["First Section", "First Subsection", "First Nested Item", "Second Section", "Second Subsection"],
          f"heading paragraph text carries no manually-typed numbers (got {heading_text})")

    num_ids = {}
    for style_name, expected_ilvl in (("Heading 2", "0"), ("Heading 3", "1"), ("Heading 4", "2")):
        style = doc.styles[style_name]
        num_pr = style.element.find('.//' + qn('w:numPr'))
        check(num_pr is not None, f"'{style_name}' style is linked to a numbering list")
        if num_pr is None:
            continue
        num_id_el = num_pr.find(qn('w:numId'))
        ilvl_el = num_pr.find(qn('w:ilvl'))
        num_id = num_id_el.get(qn('w:val')) if num_id_el is not None else None
        # w:ilvl is commonly omitted for level 0 - it's OOXML's implicit default.
        ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
        num_ids[style_name] = num_id
        check(ilvl == expected_ilvl, f"'{style_name}' is linked at list level {expected_ilvl} (got {ilvl})")

    check(len(set(num_ids.values())) == 1 and None not in num_ids.values(),
          f"Heading 2/3/4 all link to the SAME numbering list, not separate ones (got {num_ids})")


def test_heading2_has_no_automatic_page_break():
    doc = Document(SIGNATURE_TEMPLATE)
    heading2 = doc.styles['Heading 2']
    check(heading2.paragraph_format.page_break_before is not True,
          f"Heading 2 no longer forces a page break (got {heading2.paragraph_format.page_break_before!r})")


def test_get_step_list_abstract_num_id_found():
    """SIGNATURE_TEMPLATE must have the 'Dilon Step List' style + sample
    paragraph built per the spec's Template Requirement section."""
    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    check(abstract_id is not None, "finds an abstractNumId for 'Dilon Step List' in the real template")


def test_get_step_list_abstract_num_id_missing_style():
    empty_template = TEST_OUTPUT_DIR / "step_numbering_no_style_template.docx"
    from docx import Document
    Document().save(empty_template)
    abstract_id = step_numbering.get_step_list_abstract_num_id(empty_template)
    check(abstract_id is None, "returns None (not an exception) when the template has no 'Dilon Step List' style")


def test_create_num_instance_first_allocation():
    from docx import Document
    doc = Document(SIGNATURE_TEMPLATE)
    numbering_element = doc.part.numbering_part.element
    existing_ids = {int(n.get(qn('w:numId'))) for n in numbering_element.findall(qn('w:num'))}
    new_id = step_numbering.create_num_instance(numbering_element, "1")
    check(new_id not in existing_ids, f"allocates a numId ({new_id}) that didn't already exist")
    check(new_id == max(existing_ids, default=0) + 1, f"allocates max(existing)+1 (got {new_id}, existing max {max(existing_ids, default=0)})")


def test_create_num_instance_sequential_allocations_dont_collide():
    from docx import Document
    doc = Document(SIGNATURE_TEMPLATE)
    numbering_element = doc.part.numbering_part.element
    first = step_numbering.create_num_instance(numbering_element, "1")
    second = step_numbering.create_num_instance(numbering_element, "1")
    check(first != second, f"two sequential allocations never collide (got {first}, {second})")


def test_create_num_instance_writes_start_override():
    """Word continues a level's counter across separate numId instances
    that share the same abstractNumId, unless a startOverride forces a
    restart (confirmed against the real template: our throwaway sample
    paragraph's numId=67 bled its count of 1 into a freshly-allocated
    numId sharing abstractNumId=50, rendering 2/3/4 instead of 1/2/3). A
    new step-list sequence must always start at 1 regardless of what else
    used the same abstract list earlier in the document."""
    from docx import Document
    doc = Document(SIGNATURE_TEMPLATE)
    numbering_element = doc.part.numbering_part.element
    new_id = step_numbering.create_num_instance(numbering_element, "1")

    num_el = next(n for n in numbering_element.findall(qn('w:num')) if n.get(qn('w:numId')) == str(new_id))
    lvl_override = num_el.find(qn('w:lvlOverride'))
    check(lvl_override is not None, "new numId carries a lvlOverride")
    if lvl_override is not None:
        start_override = lvl_override.find(qn('w:startOverride'))
        check(start_override is not None and start_override.get(qn('w:val')) == '1',
              "lvlOverride forces the level to start at 1")


def test_ensure_blank_line_around_steps_markers_inserts_both_sides():
    md = "@@@STEPS@@@\n#. First\n#. Second\n@@@END_STEPS@@@\n"
    result = step_numbering.ensure_blank_line_around_steps_markers(md)
    check(result == "@@@STEPS@@@\n\n#. First\n#. Second\n\n@@@END_STEPS@@@\n",
          f"blank lines inserted after @@@STEPS@@@ and before @@@END_STEPS@@@ (got {result!r})")


def test_ensure_blank_line_around_steps_markers_idempotent():
    md = "@@@STEPS@@@\n\n#. First\n\n@@@END_STEPS@@@\n"
    result = step_numbering.ensure_blank_line_around_steps_markers(md)
    check(result == md, "already-blank-line case is left unchanged")


def test_apply_section_scoped_step_numbering_single_section():
    md = (
        "## Section One\n\n"
        "@@@STEPS@@@\n\n"
        "#. First\n"
        "#. Second\n"
        "    #. Sub of second\n"
        "\n@@@END_STEPS@@@\n"
    )
    docx_path = TEST_OUTPUT_DIR / "step_numbering_single_section_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    count = step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)
    check(count == 3, f"all three step paragraphs get numbered (got {count})")

    doc = Document(docx_path)
    check(all('@@@STEPS' not in p.text and '@@@END_STEPS' not in p.text for p in doc.paragraphs),
          "both wrapper marker paragraphs are removed")
    step_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Dilon Step List']
    check(len(step_paras) == 3, f"all three paragraphs carry the 'Dilon Step List' style (got {len(step_paras)})")


def test_apply_section_scoped_step_numbering_reuses_numid_within_section():
    md = (
        "## Section One\n\n"
        "@@@STEPS@@@\n\n#. First\n\n@@@END_STEPS@@@\n\n"
        "An interrupting paragraph.\n\n"
        "@@@STEPS@@@\n\n#. Second\n\n@@@END_STEPS@@@\n"
    )
    docx_path = TEST_OUTPUT_DIR / "step_numbering_two_blocks_one_section_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)

    doc = Document(docx_path)
    num_ids = set()
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Dilon Step List':
            num_id, _ = dilon_docx_common._paragraph_num_id_and_ilvl(p._p)
            num_ids.add(num_id)
    check(len(num_ids) == 1, f"two separate @@@STEPS@@@ blocks in the same section share one numId (got {num_ids})")


def test_apply_section_scoped_step_numbering_restarts_across_sections():
    md = (
        "## Section One\n\n@@@STEPS@@@\n\n#. First\n\n@@@END_STEPS@@@\n\n"
        "## Section Two\n\n@@@STEPS@@@\n\n#. Second\n\n@@@END_STEPS@@@\n"
    )
    docx_path = TEST_OUTPUT_DIR / "step_numbering_two_sections_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)

    doc = Document(docx_path)
    num_ids = []
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Dilon Step List':
            num_id, _ = dilon_docx_common._paragraph_num_id_and_ilvl(p._p)
            num_ids.append(num_id)
    check(len(set(num_ids)) == 2, f"a new section gets a fresh numId (got {num_ids})")


def test_apply_section_scoped_step_numbering_unclosed_block_raises():
    md = "@@@STEPS@@@\n\n#. First\n"
    docx_path = TEST_OUTPUT_DIR / "step_numbering_unclosed_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    try:
        step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)
        check(False, "an @@@STEPS@@@ with no matching @@@END_STEPS@@@ raises StepBlockError")
    except step_numbering.StepBlockError as exc:
        check("END_STEPS" in str(exc), f"the error mentions the missing closing marker (got: {exc})")


def test_apply_section_scoped_step_numbering_open_block_across_section_boundary_raises():
    """Regression test: a @@@STEPS@@@ left open when a new ## heading
    starts used to silently split into two separately-numbered
    sequences (the paragraphs after the heading getting a fresh numId
    keyed to the new section) instead of halting - a plausible
    authoring slip (forgetting @@@END_STEPS@@@) that must not produce
    silently wrong step numbers."""
    md = (
        "## Section One\n\n@@@STEPS@@@\n\n#. First\n#. Second\n\n"
        "## Section Two\n\n#. Third\n\n@@@END_STEPS@@@\n"
    )
    docx_path = TEST_OUTPUT_DIR / "step_numbering_open_across_section_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    try:
        step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)
        check(False, "an @@@STEPS@@@ left open across a ## section boundary raises StepBlockError")
    except step_numbering.StepBlockError as exc:
        check("section heading" in str(exc), f"the error mentions the section boundary (got: {exc})")


def test_apply_section_scoped_step_numbering_skips_gracefully_without_abstract_id():
    md = "@@@STEPS@@@\n\n#. First\n\n@@@END_STEPS@@@\n"
    docx_path = TEST_OUTPUT_DIR / "step_numbering_no_abstract_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)
    step_numbering.apply_section_scoped_step_numbering(docx_path, None)  # should not raise

    doc = Document(docx_path)
    check(all('@@@STEPS' not in p.text and '@@@END_STEPS' not in p.text for p in doc.paragraphs),
          "wrapper markers are still stripped even with numbering skipped")


def test_apply_section_scoped_step_numbering_preserves_inline_formatting():
    md = "@@@STEPS@@@\n\n#. Use **IPA** and a lint-free cloth.\n\n@@@END_STEPS@@@\n"
    docx_path = TEST_OUTPUT_DIR / "step_numbering_formatting_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)

    doc = Document(docx_path)
    step_para = [p for p in doc.paragraphs if 'IPA' in p.text][0]
    bold_runs = [r for r in step_para.runs if r.bold]
    check(len(bold_runs) == 1 and bold_runs[0].text == 'IPA', "bold formatting on 'IPA' survives")


def test_resolve_step_reference_builds_composite_field():
    md = "## Section One\n\n@@@STEPS@@@\n\n#. Hold the board. []{#step:x}\n\n@@@END_STEPS@@@\n\nSee [](#step:x).\n"
    md = dilon_docx_common.preprocess_reference_markers(md)
    docx_path = TEST_OUTPUT_DIR / "step_resolver_callback_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    abstract_id = step_numbering.get_step_list_abstract_num_id(SIGNATURE_TEMPLATE)
    step_numbering.apply_section_scoped_step_numbering(docx_path, abstract_id)
    dilon_docx_common.resolve_reference_markers(docx_path, {'step': step_numbering.resolve_step_reference})

    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    check('STYLEREF 2 \\s' in xml, "the section-number half is present")
    check('REF step:x \\r \\h' in xml, "the step-number half is present")
    check('Step ' in xml, "the literal 'Step ' prefix is present")


STEP_REDESIGN_MARKDOWN = (
    '\n## Carrier Board Assembly Procedure\n\n'
    '@@@STEPS@@@\n\n'
    '#. Wear clean gloves.\n'
    '#. Simple dirt such as lint or light dust can be blown away before wiping.\n'
    '    #. Hold the board by the edges. []{#step:hold-board-by-edges}\n\n'
    '@@@END_STEPS@@@\n\n'
    'NOTE: Clean the entire crystal but give special attention to the polished end.\n\n'
    '@@@STEPS@@@\n\n'
    '#. Visually inspect both the crystal and the photomultiplier for defects.\n'
    '#. Set the cleaned crystals aside on a clean lint free cloth.\n\n'
    '@@@END_STEPS@@@\n\n'
    'As described in [](#step:hold-board-by-edges), always support the board by its edges.\n'
)


def test_compile_section_scoped_step_numbering_end_to_end():
    """Integration test: two @@@STEPS@@@ blocks in one section
    (interrupted by a NOTE), a nested sub-step, and a cross-reference,
    compiled through the real pipeline."""
    markdown = SAMPLE_MARKDOWN + STEP_REDESIGN_MARKDOWN
    input_md = TEST_OUTPUT_DIR / "compile_test_section_step_numbering.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_section_step_numbering.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a document with section-scoped @@@STEPS@@@ blocks")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    doc = Document(output_docx)
    check(all('@@@STEPS' not in p.text and 'STEP' not in p.text for p in doc.paragraphs
              if p.style is None or p.style.name != 'Dilon Step List' or '@@@' not in p.text),
          "no wrapper marker or sentinel text remains anywhere")
    step_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == 'Dilon Step List']
    check(len(step_paragraphs) == 5, f"all 5 steps across both blocks get the 'Dilon Step List' style (got {len(step_paragraphs)})")

    num_ids = set()
    for p in step_paragraphs:
        num_id_el = p._p.find('.//' + qn('w:numPr') + '/' + qn('w:numId'))
        if num_id_el is not None:
            num_ids.add(num_id_el.get(qn('w:val')))
    check(len(num_ids) == 1, f"both blocks in the same section share one numId (got {num_ids})")

    with zipfile.ZipFile(output_docx) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    check('w:name="step:hold-board-by-edges"' in xml, "the step's anchor survives as a real bookmark")
    check('REF step:hold-board-by-edges \\r \\h' in xml, "the cross-reference resolves to a live REF field")
    check('STYLEREF 2 \\s' in xml, "the cross-reference includes a live section-number field")


def test_compile_duplicate_step_anchor_fails_clearly():
    markdown = SAMPLE_MARKDOWN + (
        '\n## Section\n\n@@@STEPS@@@\n\n'
        '#. First. []{#step:dup}\n#. Second. []{#step:dup}\n\n@@@END_STEPS@@@\n'
    )
    input_md = TEST_OUTPUT_DIR / "compile_test_duplicate_step_anchor.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_duplicate_step_anchor.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode != 0, "compiler exits non-zero for a duplicate {#step:x} anchor")
    check("dup" in result.stderr.lower() or "dup" in result.stdout.lower(),
          "the failure message names the duplicated label")


FULL_XREF_MARKDOWN = (
    '\n## Assembly Section {#sec:assembly}\n\n'
    '![A widget.](diagrams/example.png){#fig:widget}\n\n'
    '@@@STEPS@@@\n\n'
    '#. Install the widget. []{#step:install-widget}\n\n'
    '@@@END_STEPS@@@\n\n'
    'See [](#fig:widget), [](#sec:assembly), and [](#step:install-widget) for full context.\n'
)


def test_compile_full_cross_reference_set_end_to_end():
    """Integration test: a figure, a section, and a step, each
    referenced via [](#TYPE:label), compiled through the real
    pipeline."""
    images_dir = TEST_OUTPUT_DIR / "diagrams"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "example.png").write_bytes(make_test_png())

    markdown = SAMPLE_MARKDOWN + FULL_XREF_MARKDOWN
    input_md = TEST_OUTPUT_DIR / "compile_test_full_xref.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_full_xref.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a document exercising all three reference types")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    with zipfile.ZipFile(output_docx) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    check('REF fig:widget \\h' in xml, "the figure reference resolved")
    check('REF sec:assembly \\r \\h' in xml, "the section reference resolved")
    check('REF step:install-widget \\r \\h' in xml, "the step reference resolved")
    check('XREF' not in xml, "no sentinel remains")


def test_compile_broken_reference_fails_clearly():
    markdown = SAMPLE_MARKDOWN + '\nSee [](#fig:does-not-exist) for details.\n'
    input_md = TEST_OUTPUT_DIR / "compile_test_broken_xref.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_broken_xref.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode != 0, "compiler exits non-zero for a reference to a nonexistent figure")
    check("does-not-exist" in result.stderr.lower() or "does-not-exist" in result.stdout.lower(),
          "the failure message names the missing label")


def test_compile_duplicate_sec_anchor_fails_clearly():
    markdown = SAMPLE_MARKDOWN + (
        '\n## One {#sec:dup}\n\nText.\n\n## Two {#sec:dup}\n\nSee [](#sec:dup).\n'
    )
    input_md = TEST_OUTPUT_DIR / "compile_test_duplicate_sec_anchor.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_duplicate_sec_anchor.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode != 0, "compiler exits non-zero for a duplicate {#sec:x} anchor")
    check("dup" in result.stderr.lower() or "dup" in result.stdout.lower(),
          "the failure message names the duplicated label")


def test_compile_duplicate_fig_anchor_fails_clearly():
    images_dir = TEST_OUTPUT_DIR / "dup_fig_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "a.png").write_bytes(make_test_png())
    (images_dir / "b.png").write_bytes(make_test_png())

    markdown = SAMPLE_MARKDOWN + (
        '\n![First.](dup_fig_images/a.png){#fig:dup}\n\n'
        '![Second.](dup_fig_images/b.png){#fig:dup}\n\n'
        'See [](#fig:dup).\n'
    )
    input_md = TEST_OUTPUT_DIR / "compile_test_duplicate_fig_anchor.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_duplicate_fig_anchor.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode != 0, "compiler exits non-zero for a duplicate {#fig:x} anchor")
    check("dup" in result.stderr.lower() or "dup" in result.stdout.lower(),
          "the failure message names the duplicated label")


def _bookmark_start_end_mismatches(doc):
    """Walks a compiled document's body in order, treating bookmarkStart/
    bookmarkEnd as a LIFO stack (proper nesting). Returns a list of
    (end_id, expected_id_or_None) tuples for every bookmarkEnd whose id
    doesn't close the innermost still-open bookmarkStart - i.e. every
    place the document's bookmark structure isn't well-formed.

    docxcompose's Composer.renumber_bookmarks() reassigns every
    bookmarkStart's id sequentially in document order, then separately
    reassigns every bookmarkEnd's id sequentially in document order,
    assuming the Nth start always pairs with the Nth end. That's only
    true when bookmarks never nest. Our documents nest constantly (an
    un-narrowed heading bookmark stays open around a {#fig:x} bookmark
    that opens and closes inside it), so the naive renumbering hands a
    figure's bookmarkEnd the wrong id - the figure's real end lands on
    an unrelated, later bookmark instead."""
    stack = []
    mismatches = []
    for el in doc.element.body.iter():
        if el.tag == qn('w:bookmarkStart'):
            stack.append(el.get(qn('w:id')))
        elif el.tag == qn('w:bookmarkEnd'):
            end_id = el.get(qn('w:id'))
            if stack and stack[-1] == end_id:
                stack.pop()
            else:
                mismatches.append((end_id, stack[-1] if stack else None))
    return mismatches


def _bookmark_span_contains_drawing(doc, bookmark_name):
    """True if the named bookmark's start...end span (in document order)
    contains a <w:drawing> (embedded picture) element - i.e. the
    bookmark isn't narrowly wrapping just caption text, it's swallowing
    an image. This is the literal Word-visible symptom of the
    docxcompose renumbering bug: a REF field against a bookmark whose
    end got reassigned to some later, unrelated position spans - and so
    inserts a copy of - everything in between, images included."""
    body = doc.element.body
    all_els = list(body.iter())
    start_idx = end_idx = None
    for i, el in enumerate(all_els):
        if el.tag == qn('w:bookmarkStart') and el.get(qn('w:name')) == bookmark_name:
            start_idx = i
            start_id = el.get(qn('w:id'))
            break
    if start_idx is None:
        raise AssertionError(f"no bookmarkStart named {bookmark_name!r} found")
    for i in range(start_idx + 1, len(all_els)):
        el = all_els[i]
        if el.tag == qn('w:bookmarkEnd') and el.get(qn('w:id')) == start_id:
            end_idx = i
            break
    if end_idx is None:
        raise AssertionError(f"no matching bookmarkEnd for {bookmark_name!r} (id={start_id})")
    return any(el.tag == qn('w:drawing') for el in all_els[start_idx:end_idx])


MERGE_BUG_MARKDOWN = (
    '\n## Widget Assembly\n\n'
    '![First widget.](merge_bug_images/one.png){#fig:widget-one}\n\n'
    'Reference the first widget here ([](#fig:widget-one)).\n\n'
    '![Second widget.](merge_bug_images/two.png){#fig:widget-two}\n\n'
    '![Third widget.](merge_bug_images/three.png){#fig:widget-three}\n'
)


def test_compile_figure_reference_bookmark_not_corrupted_by_merge():
    """Regression test for a bug report: a [](#fig:x) cross-reference
    rendered as a live copy of the figure's image pasted into the text,
    instead of hyperlinked "Figure N.M" text. Root cause: Part D's own
    bookmarks are correctly paired (verified separately by tracing every
    pipeline stage), but compose_documents()'s merge step
    (docxcompose's Composer.append() -> renumber_bookmarks()) reassigns
    bookmarkStart/bookmarkEnd ids independently and positionally,
    corrupting any nested bookmark pair - which an un-narrowed heading
    bookmark wrapping a {#fig:x} bookmark always produces. Reproduces
    with an ordinary (unnamed) heading followed by three back-to-back
    figures, the first one referenced - the same shape as the real
    document that surfaced the bug."""
    images_dir = TEST_OUTPUT_DIR / "merge_bug_images"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / "one.png").write_bytes(make_test_png())
    (images_dir / "two.png").write_bytes(make_test_png())
    (images_dir / "three.png").write_bytes(make_test_png())

    markdown = SAMPLE_MARKDOWN + MERGE_BUG_MARKDOWN
    input_md = TEST_OUTPUT_DIR / "compile_test_merge_bug.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_merge_bug.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for the merge-bug repro document")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    doc = Document(output_docx)

    mismatches = _bookmark_start_end_mismatches(doc)
    check(not mismatches,
          f"every bookmarkEnd in the merged document closes its correct bookmarkStart "
          f"(found {len(mismatches)} mismatch(es): {mismatches[:5]}{'...' if len(mismatches) > 5 else ''})")

    check(not _bookmark_span_contains_drawing(doc, 'fig:widget-one'),
          "the fig:widget-one bookmark wraps only its caption text, not an embedded picture "
          "(a REF field against it would otherwise paste the image, not hyperlinked text)")


def test_validate_list_nesting_depth_passes_at_three_levels():
    md = (
        "#. Top\n"
        "    #. Second\n"
        "        #. Third\n"
    )
    docx_path = TEST_OUTPUT_DIR / "nesting_ok_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)
    validate_list_nesting_depth(docx_path)  # should not raise


def test_validate_list_nesting_depth_rejects_four_levels():
    md = (
        "#. Top\n"
        "    #. Second\n"
        "        #. Third\n"
        "            #. Fourth\n"
    )
    docx_path = TEST_OUTPUT_DIR / "nesting_bad_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)
    try:
        validate_list_nesting_depth(docx_path)
        check(False, "a 4-level-deep ordered list raises ListNestingError")
    except ListNestingError as exc:
        check("Fourth" in str(exc), f"the error names the offending item's text (got: {exc})")


def test_validate_list_nesting_depth_ignores_bullet_lists():
    md = (
        "- Top\n"
        "    - Second\n"
        "        - Third\n"
        "            - Fourth\n"
    )
    docx_path = TEST_OUTPUT_DIR / "nesting_bullets_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)
    validate_list_nesting_depth(docx_path)  # bullets aren't capped - should not raise


def test_remap_ordered_lists_restyles_decimal_paragraphs():
    md = "#. First\n#. Second\n"
    docx_path = TEST_OUTPUT_DIR / "remap_ordered_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    count = remap_ordered_lists_to_dilon_step_list(docx_path)
    check(count == 2, f"both ordered-list paragraphs get restyled (got {count})")

    doc = Document(docx_path)
    step_styled = [p for p in doc.paragraphs if p.style and p.style.name == 'Dilon Step List']
    check(len(step_styled) == 2, "both paragraphs now carry the 'Dilon Step List' style")


def test_remap_ordered_lists_leaves_bullets_alone():
    md = "- First\n- Second\n"
    docx_path = TEST_OUTPUT_DIR / "remap_bullets_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    count = remap_ordered_lists_to_dilon_step_list(docx_path)
    check(count == 0, f"a bullet list is not restyled (got count={count})")

    doc = Document(docx_path)
    step_styled = [p for p in doc.paragraphs if p.style and p.style.name == 'Dilon Step List']
    check(len(step_styled) == 0, "no paragraph carries 'Dilon Step List' style")


def test_remap_ordered_lists_preserves_native_numbering():
    """The remap must only change *style*, never numId/ilvl - Pandoc's
    own numbering is already numerically correct and must keep driving
    the rendered number."""
    md = "#. First\n#. Second\n    #. Nested\n#. Third\n"
    docx_path = TEST_OUTPUT_DIR / "remap_preserves_numbering_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    before = [_paragraph_num_id_and_ilvl(p._p) for p in Document(docx_path).paragraphs]
    remap_ordered_lists_to_dilon_step_list(docx_path)
    after = [_paragraph_num_id_and_ilvl(p._p) for p in Document(docx_path).paragraphs]
    check(before == after, f"numId/ilvl unchanged by the style remap (before={before}, after={after})")


def test_ensure_blank_line_after_list_continue_marker_inserts_when_missing():
    md = "@@@CONTINUE:#list:x@@@\n#. Third\n"
    result = ensure_blank_line_after_list_continue_markers(md)
    check(result == "@@@CONTINUE:#list:x@@@\n\n#. Third\n", f"a blank line is inserted (got {result!r})")


def test_ensure_blank_line_after_list_continue_marker_idempotent():
    md = "@@@CONTINUE:#list:x@@@\n\n#. Third\n"
    result = ensure_blank_line_after_list_continue_markers(md)
    check(result == md, "already-blank-line case is left unchanged")


def test_resolve_list_continuations_reuses_numid():
    md = (
        "#. First\n"
        "#. Second []{#list:cleaning-procedure}\n\n"
        "Some interrupting paragraph.\n\n"
    )
    md = ensure_blank_line_after_list_continue_markers(
        md + "@@@CONTINUE:#list:cleaning-procedure@@@\n#. Third\n#. Fourth\n"
    )
    docx_path = TEST_OUTPUT_DIR / "list_continue_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    resolved = resolve_list_continuations(docx_path)
    check(resolved == 1, f"exactly one continuation marker resolved (got {resolved})")

    doc = Document(docx_path)
    check(all('@@@CONTINUE' not in p.text for p in doc.paragraphs), "the marker paragraph is removed")

    num_ids = set()
    for p in doc.paragraphs:
        num_id, _ = _paragraph_num_id_and_ilvl(p._p)
        if num_id is not None:
            num_ids.add(num_id)
    check(len(num_ids) == 1, f"both blocks now share exactly one numId (got {num_ids})")


def test_resolve_list_continuations_missing_anchor_raises():
    md = ensure_blank_line_after_list_continue_markers(
        "@@@CONTINUE:#list:does-not-exist@@@\n#. Third\n"
    )
    docx_path = TEST_OUTPUT_DIR / "list_continue_missing_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    try:
        resolve_list_continuations(docx_path)
        check(False, "a @@@CONTINUE@@@ with no matching {#list:name} anchor raises ListContinuationError")
    except ListContinuationError as exc:
        check("does-not-exist" in str(exc), f"the error names the missing label (got: {exc})")


def test_resolve_list_continuations_duplicate_anchor_raises():
    md = (
        "#. First []{#list:dup}\n\n"
        "#. Second []{#list:dup}\n\n"
    )
    md = ensure_blank_line_after_list_continue_markers(
        md + "@@@CONTINUE:#list:dup@@@\n#. Third\n"
    )
    docx_path = TEST_OUTPUT_DIR / "list_continue_duplicate_test.docx"
    compiler.markdown_to_docx(md, docx_path, reference_doc=SIGNATURE_TEMPLATE)

    try:
        resolve_list_continuations(docx_path)
        check(False, "two {#list:dup} anchors raises ListContinuationError")
    except ListContinuationError as exc:
        check("dup" in str(exc), f"the error names the duplicated label (got: {exc})")


def test_compile_ordered_list_and_continuation_end_to_end():
    """Integration test: a #. list interrupted by a paragraph and
    resumed via {#list:name}/@@@CONTINUE@@@, compiled through the real
    pipeline - numbering.xml and the shared numId must survive the
    full A/B/C/D docxcompose merge."""
    markdown = SAMPLE_MARKDOWN + (
        "\n## Ordered List Continuation Example\n\n"
        "#. First item\n"
        "#. Second item []{#list:demo-list}\n\n"
        "An interrupting paragraph.\n\n"
        "@@@CONTINUE:#list:demo-list@@@\n"
        "#. Third item\n"
        "#. Fourth item\n"
    )
    input_md = TEST_OUTPUT_DIR / "compile_test_ordered_list.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_ordered_list.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0, "compiler exits 0 for a document with a continued #. list")
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return

    doc = Document(output_docx)
    check(all('@@@CONTINUE' not in p.text for p in doc.paragraphs), "no marker text remains")
    step_styled = [p for p in doc.paragraphs if p.style and p.style.name == 'Dilon Step List'
                   and p.text.strip() in ('First item', 'Second item', 'Third item', 'Fourth item')]
    check(len(step_styled) == 4, f"all four items got the 'Dilon Step List' style (got {len(step_styled)})")

    num_ids = set()
    for p in step_styled:
        num_id, _ = _paragraph_num_id_and_ilvl(p._p)
        if num_id is not None:
            num_ids.add(num_id)
    check(len(num_ids) == 1, f"all four items share one numId across the interruption (got {num_ids})")


def test_compile_four_level_nested_list_fails_clearly():
    markdown = SAMPLE_MARKDOWN + (
        "\n## Over-Nested List Example\n\n"
        "#. Top\n"
        "    #. Second\n"
        "        #. Third\n"
        "            #. Fourth\n"
    )
    input_md = TEST_OUTPUT_DIR / "compile_test_over_nested.md"
    output_docx = TEST_OUTPUT_DIR / "compile_test_over_nested.docx"
    input_md.write_text(markdown, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(COMPILER_SCRIPT), str(input_md), str(output_docx), str(SIGNATURE_TEMPLATE)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode != 0, "compiler exits non-zero for a 4-level-deep ordered list")
    check("nested" in result.stderr.lower() or "nested" in result.stdout.lower(),
          "the failure message mentions nesting, not a raw traceback only")


def test_no_shebang_in_python_scripts():
    def has_shebang(path):
        lines = path.read_text(encoding="utf-8").splitlines()
        return bool(lines) and lines[0].startswith("#!")

    offenders = [str(p) for p in SHEBANG_GUARDED_SCRIPTS if has_shebang(p)]
    check(not offenders, f"no shebang lines in guarded scripts (offenders: {offenders})")


def run_validator():
    result = subprocess.run(
        [sys.executable, "validate-output.py"],
        cwd=str(Path(__file__).parent),
    )
    return result.returncode == 0


def main():
    if TEST_OUTPUT_DIR.exists():
        shutil.rmtree(TEST_OUTPUT_DIR)
    TEST_OUTPUT_DIR.mkdir(parents=True)

    test_stub_custom_params()
    test_stub_default_params()
    test_stub_duplicate_file_error()
    test_ensure_blank_line_single_marker()
    test_ensure_blank_line_stacked_style_then_columns()
    test_ensure_blank_line_stacked_columns_then_style()
    test_ensure_blank_line_idempotent_when_already_blank()
    test_ensure_blank_line_between_images_inserts_when_missing()
    test_ensure_blank_line_between_images_three_in_a_row()
    test_ensure_blank_line_between_images_idempotent_when_already_blank()
    test_ensure_blank_line_between_images_leaves_non_image_lines_alone()
    test_ensure_blank_line_between_images_ignores_inline_image_with_text()
    test_parse_column_widths_valid_with_flex()
    test_parse_column_widths_valid_all_numeric()
    test_parse_column_widths_case_insensitive_flex()
    test_parse_column_widths_rejects_count_mismatch()
    test_parse_column_widths_rejects_multiple_flex()
    test_parse_column_widths_rejects_non_numeric()
    test_parse_column_widths_rejects_zero_or_negative()
    test_parse_column_widths_rejects_non_finite()
    test_apply_table_column_widths_fixed_and_flex()
    test_apply_table_column_widths_all_fixed()
    test_apply_table_column_widths_raises_when_flex_overflows()
    test_apply_table_column_widths_raises_when_all_fixed_overflows()
    test_compile_missing_input_error()
    test_compile_valid_document()
    test_create_signature_table_structure()
    test_compile_signature_table_generated_programmatically()
    test_compile_has_no_leading_blank_paragraph()
    test_compile_has_no_title_page()
    test_compile_header_signature_revision_widths()
    test_compile_footer_table_layout()
    test_compile_bom_front_matter()
    test_compile_table_marker_no_blank_line()
    test_compile_adjacent_tables_no_merge()
    test_compile_table_column_widths()
    test_compile_with_default_templates()
    test_compile_resolves_relative_image_paths()
    test_render_jinja_substitutes_body_fields()
    test_render_jinja_raw_block_escapes_literal_braces()
    test_render_jinja_noop_without_braces()
    test_compile_body_jinja_substitution()
    test_heading_auto_numbering()
    test_figure_auto_numbering()
    test_figure_auto_numbering_consecutive_images_no_blank_line()
    test_apply_figure_captions_narrows_bookmark_to_number_only()
    test_apply_figure_captions_narrowing_is_noop_without_fig_id()
    test_narrow_section_bookmarks_shrinks_to_heading_only()
    test_narrow_section_bookmarks_ignores_headings_without_sec_id()
    test_preprocess_reference_markers_converts_recognized_types()
    test_preprocess_reference_markers_leaves_real_link_text_and_unknown_types_untouched()
    test_resolve_reference_markers_dispatches_by_type()
    test_resolve_reference_markers_missing_anchor_raises()
    test_resolve_reference_markers_duplicate_anchor_raises()
    test_heading2_has_no_automatic_page_break()
    test_get_step_list_abstract_num_id_found()
    test_get_step_list_abstract_num_id_missing_style()
    test_create_num_instance_first_allocation()
    test_create_num_instance_sequential_allocations_dont_collide()
    test_create_num_instance_writes_start_override()
    test_ensure_blank_line_around_steps_markers_inserts_both_sides()
    test_ensure_blank_line_around_steps_markers_idempotent()
    test_apply_section_scoped_step_numbering_single_section()
    test_apply_section_scoped_step_numbering_reuses_numid_within_section()
    test_apply_section_scoped_step_numbering_restarts_across_sections()
    test_apply_section_scoped_step_numbering_unclosed_block_raises()
    test_apply_section_scoped_step_numbering_open_block_across_section_boundary_raises()
    test_apply_section_scoped_step_numbering_skips_gracefully_without_abstract_id()
    test_apply_section_scoped_step_numbering_preserves_inline_formatting()
    test_resolve_step_reference_builds_composite_field()
    test_compile_section_scoped_step_numbering_end_to_end()
    test_compile_duplicate_step_anchor_fails_clearly()
    test_compile_full_cross_reference_set_end_to_end()
    test_compile_broken_reference_fails_clearly()
    test_compile_duplicate_sec_anchor_fails_clearly()
    test_compile_duplicate_fig_anchor_fails_clearly()
    test_compile_figure_reference_bookmark_not_corrupted_by_merge()
    test_validate_list_nesting_depth_passes_at_three_levels()
    test_validate_list_nesting_depth_rejects_four_levels()
    test_validate_list_nesting_depth_ignores_bullet_lists()
    test_remap_ordered_lists_restyles_decimal_paragraphs()
    test_remap_ordered_lists_leaves_bullets_alone()
    test_remap_ordered_lists_preserves_native_numbering()
    test_ensure_blank_line_after_list_continue_marker_inserts_when_missing()
    test_ensure_blank_line_after_list_continue_marker_idempotent()
    test_resolve_list_continuations_reuses_numid()
    test_resolve_list_continuations_missing_anchor_raises()
    test_resolve_list_continuations_duplicate_anchor_raises()
    test_compile_ordered_list_and_continuation_end_to_end()
    test_compile_four_level_nested_list_fails_clearly()
    test_no_shebang_in_python_scripts()

    print(f"\n{passed} passed, {failed} failed (direct-invocation checks)")

    validator_ok = run_validator()

    if failed == 0 and validator_ok:
        print("\nAll tests passed!")
        return 0
    print("\nSome tests failed")
    return 1


if __name__ == "__main__":
    sys.exit(main())
